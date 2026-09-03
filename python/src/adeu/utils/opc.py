# FILE: src/adeu/utils/opc.py
"""Opening WordprocessingML packages that `python-docx` refuses on sight.

`docx.Document()` accepts exactly one main-part content type,
`...wordprocessingml.document.main+xml`, and raises `ValueError: ... is not a
Word file` for everything else. That rejects **templates** (`.dotx`) and
**macro-enabled** documents (`.docm`, `.dotm`) — all of which are ordinary
WordprocessingML packages whose main part differs only in its declared content
type. `@adeu/core` reads them all, so the refusal was also a dual-engine parity
break (CC-11).

Two ways to fix this, and the difference matters:

1. **Rewrite the content type on load.** One line, and wrong: `python-docx`
   serialises `[Content_Types].xml` from the parts' own `content_type`, so the
   rewrite follows the document all the way to `save()` and silently converts
   the user's template into a document. Fidelity loss on a read-only operation.
2. **Teach `python-docx` the other content types.** `PartFactory.part_type_for`
   is the mapping it consults to decide which class a part becomes; registering
   the template and macro types against `DocumentPart` makes the part a real
   `DocumentPart` while it keeps its own content type. `save()` then round-trips
   `template.main+xml` untouched — a `.dotx` in is a `.dotx` out.

This module does (2). Verified: `odot_uic_drywell.dotx` opens with 106
paragraphs / 5 tables, and a save preserves `template.main+xml`.
"""

from __future__ import annotations

from typing import IO, Union

from docx.document import Document as DocumentObject
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.package import OpcPackage
from docx.opc.part import PartFactory
from docx.parts.document import DocumentPart

#: Main-part content types that are WordprocessingML documents in everything but
#: name. Values are the literal strings from ECMA-376 / the MS-word extensions;
#: `python-docx`'s `CONTENT_TYPE` enum only defines the first.
WML_DOCUMENT_MAIN = str(CT.WML_DOCUMENT_MAIN)
WML_TEMPLATE_MAIN = "application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml"
WML_DOCUMENT_MACRO_MAIN = "application/vnd.ms-word.document.macroEnabled.main+xml"
WML_TEMPLATE_MACRO_MAIN = "application/vnd.ms-word.template.macroEnabledTemplate.main+xml"

#: Everything `load_document` will open. Order is documentation, not logic.
MAIN_DOCUMENT_CONTENT_TYPES: tuple[str, ...] = (
    WML_DOCUMENT_MAIN,  # .docx
    WML_TEMPLATE_MAIN,  # .dotx
    WML_DOCUMENT_MACRO_MAIN,  # .docm
    WML_TEMPLATE_MACRO_MAIN,  # .dotm
)

_EXTENSION_BY_CONTENT_TYPE = {
    WML_DOCUMENT_MAIN: ".docx",
    WML_TEMPLATE_MAIN: ".dotx",
    WML_DOCUMENT_MACRO_MAIN: ".docm",
    WML_TEMPLATE_MACRO_MAIN: ".dotm",
}


def _register_part_types() -> None:
    """Map the extra main-part content types onto `DocumentPart`.

    Idempotent, and deliberately additive: the entry `python-docx` already
    installs for `WML_DOCUMENT_MAIN` is never touched. Without this the part
    falls back to the generic `Part`, which has no `.document`.
    """
    for content_type in (
        WML_TEMPLATE_MAIN,
        WML_DOCUMENT_MACRO_MAIN,
        WML_TEMPLATE_MACRO_MAIN,
    ):
        PartFactory.part_type_for.setdefault(content_type, DocumentPart)


_register_part_types()


def is_template(document: DocumentObject) -> bool:
    """True when `document` came from a `.dotx`/`.dotm` package.

    Callers that write a *new* file next to the original (sanitize, CLI apply)
    can use this to pick the right extension; the content type itself is
    preserved by `save()` either way.
    """
    return document.part.content_type in (WML_TEMPLATE_MAIN, WML_TEMPLATE_MACRO_MAIN)


def suggested_extension(document: DocumentObject) -> str:
    """The file extension matching `document`'s main-part content type."""
    return _EXTENSION_BY_CONTENT_TYPE.get(document.part.content_type, ".docx")


def load_document(source: Union[str, IO[bytes]]) -> DocumentObject:
    """Open a WordprocessingML package: `.docx`, `.dotx`, `.docm` or `.dotm`.

    Drop-in replacement for `docx.Document()`. Same return type, same failure
    for genuinely wrong input — but the error names the content type it found
    and the ones it accepts, instead of `python-docx`'s bare "is not a Word
    file", because the CLI used to surface that as an unhandled traceback.

    Note the argument is required here. `docx.Document()` defaults to opening
    its bundled empty template, which is never what Adeu wants and turns a
    forgotten argument into a silently empty document.
    """
    part = OpcPackage.open(source).main_document_part

    if part.content_type not in MAIN_DOCUMENT_CONTENT_TYPES:
        accepted = "\n  ".join(
            f"{content_type}  ({_EXTENSION_BY_CONTENT_TYPE[content_type]})"
            for content_type in MAIN_DOCUMENT_CONTENT_TYPES
        )
        raise ValueError(
            "not a Word document: the package's main part declares content type\n"
            f"  {part.content_type}\n"
            "Adeu reads WordprocessingML packages whose main part is one of:\n"
            f"  {accepted}\n"
            "A .doc (pre-2007 binary), .rtf or .odt is a different format entirely — "
            "convert it to .docx in Word first."
        )

    # Guaranteed by the registration above; asserted because a future
    # python-docx could change PartFactory and the failure would otherwise be an
    # AttributeError deep inside a caller.
    if not isinstance(part, DocumentPart):
        raise ValueError(
            f"main part for {part.content_type} loaded as {type(part).__name__}, "
            "not DocumentPart — adeu.utils.opc's PartFactory registration did not take "
            "effect (python-docx internals changed?)"
        )

    return part.document
