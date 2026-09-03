"""
Repro: the Python engine projected a manual page break as 22 characters of
literal OOXML markup.

    <w:p><w:r><w:t>A</w:t><w:br w:type="page"/><w:t>B</w:t></w:r></w:p>

projected as `A<w:br w:type="page"/>B`, while the Node engine projected `A\nB`.

This was not an accidental serialization: `_PAGE_BREAK_TOKEN` was a deliberate
in-band sentinel, and it is load-bearing — `pagination._tokenize_into_atomic_blocks`
splits on it so that manual page breaks start new virtual pages. The cost was
that an LLM read the markup as prose, a `target_text` spanning a break had to
include the XML, and the two engines silently disagreed (`docs/FIDELITY.md`
documents a newline, which only Node honoured).

CC-10 keeps the signal but moves it out of markup: both engines now project
U+000C FORM FEED, the conventional plain-text page separator. Pagination splits
on that instead, so the capability survives while the character stream stays
free of markup.

The two properties below are the ones that regressed against each other, so
both are pinned: **no markup in the projection** AND **manual breaks still
paginate**.
"""

import io

import pytest
from docx import Document
from docx.oxml import parse_xml

from adeu.ingest import extract_text_from_stream
from adeu.pagination import PAGE_BREAK_TOKEN, paginate
from adeu.redline.mapper import DocumentMapper
from adeu.utils.docx import get_run_text

NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'


def _docx(*body_xml: str) -> bytes:
    doc = Document()
    for xml in body_xml:
        doc.element.body.append(parse_xml(xml))
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _project(data: bytes, clean_view: bool = True) -> str:
    return extract_text_from_stream(io.BytesIO(data), clean_view=clean_view, include_appendix=False)


PAGE_BREAK_RUN = f'<w:p {NS}><w:r><w:t>A</w:t><w:br w:type="page"/><w:t>B</w:t></w:r></w:p>'
LINE_BREAK_RUN = f"<w:p {NS}><w:r><w:t>C</w:t><w:br/><w:t>D</w:t></w:r></w:p>"


# ---------------------------------------------------------------------------
# No markup may reach the character stream
# ---------------------------------------------------------------------------
@pytest.mark.parametrize("clean_view", [False, True])
def test_page_break_does_not_project_as_markup(clean_view):
    text = _project(_docx(PAGE_BREAK_RUN), clean_view)
    assert "<w:" not in text, f"raw OOXML reached the projection: {text!r}"
    assert "w:br" not in text
    assert text.strip() == f"A{PAGE_BREAK_TOKEN}B"


def test_page_break_projects_as_form_feed():
    assert PAGE_BREAK_TOKEN == "\f", "the page-break signal must stay U+000C"
    assert _project(_docx(PAGE_BREAK_RUN)).strip() == "A\fB"


def test_plain_line_break_still_projects_as_newline():
    """Only w:type="page" is special; a soft break stays a newline."""
    assert _project(_docx(LINE_BREAK_RUN)).strip() == "C\nD"


def test_ingest_and_mapper_agree_on_the_break():
    data = _docx(PAGE_BREAK_RUN, LINE_BREAK_RUN)
    assert DocumentMapper(Document(io.BytesIO(data)), clean_view=True).full_text == _project(data)


def test_public_run_helper_agrees_with_the_projection():
    """get_run_text inlined the literal markup independently of the constant."""
    doc = Document(io.BytesIO(_docx(PAGE_BREAK_RUN)))
    run_texts = [get_run_text(r) for p in doc.paragraphs for r in p.runs]
    joined = "".join(run_texts)
    assert "<w:" not in joined, f"get_run_text still emits markup: {joined!r}"
    assert joined == f"A{PAGE_BREAK_TOKEN}B"


# ---------------------------------------------------------------------------
# ...while the capability the markup was carrying survives
# ---------------------------------------------------------------------------
def test_manual_page_break_still_forces_a_new_page():
    """The whole reason the sentinel existed. Removing it must not cost this."""
    result = paginate(f"First page body.{PAGE_BREAK_TOKEN}Second page body.", "")
    assert result.pages[0].total_pages == 2, (
        f"manual page break no longer splits: {[p.page_content for p in result.pages]}"
    )
    assert "First page body." in result.pages[0].page_content
    assert "Second page body." in result.pages[1].page_content


def test_pagination_offsets_survive_the_shorter_token():
    """
    pagination advanced its cursor by len(token). That was 22; it is now 1, and
    the offsets it reports must still land on the real text.
    """
    body = f"Alpha.{PAGE_BREAK_TOKEN}Beta."
    result = paginate(body, "")
    assert result.body_page_offsets[0] == 0
    for content, offset in zip(result.body_pages, result.body_page_offsets, strict=True):
        assert body[offset : offset + len(content)] == content, (
            f"page offset {offset} does not point at its own content {content!r}"
        )
    # The second page must start after the 1-char token, not 22 chars later.
    assert result.body_page_offsets[1] == len("Alpha.") + len(PAGE_BREAK_TOKEN)
