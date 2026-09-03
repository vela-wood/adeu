"""
Repro: the Python engine silently dropped table rows and cells wrapped in
structured document tags (content controls).

Word wraps table content in `w:sdt` elements whenever a template uses content
controls. Two shapes matter for tables:

    <w:tbl>                              <w:tr>
      <w:sdt>                              <w:sdt>
        <w:sdtContent>                       <w:sdtContent>
          <w:tr>...</w:tr>      and            <w:tc>...</w:tc>
        </w:sdtContent>                      </w:sdtContent>
      </w:sdt>                             </w:sdt>
    </w:tbl>                             </w:tr>

`python-docx` exposes `Table.rows` via `CT_Tbl.tr_lst` (`./w:tr`) and
`_Row.cells` via `CT_Row.tc_lst` (`./w:tc`) — both DIRECT-CHILD lookups. Any
row or cell behind an sdt wrapper was therefore invisible to every Python
projection: the row vanished entirely, and a wrapped cell vanished along with
its `" | "` separator, silently misaligning the remaining columns.

The Node engine has always traversed these (its primitives use a descendant
search), so this was an engine-parity gap. Real-world blast radius: the
FedRAMP SSP Moderate rev4 template carries 371 cell-level SDTs.

Both the ingest projection (`extract_text_from_stream`) and the redline
mapper (`DocumentMapper.full_text`) must agree — they are two halves of the
Virtual Text contract, so every assertion below is made against both.
"""

import io
import re

import pytest
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn

from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText
from adeu.outline import extract_outline
from adeu.redline.engine import RedlineEngine
from adeu.redline.mapper import DocumentMapper
from tests.sdt_fixtures import load_shared_fixture_xml
from tests.utils import assert_word_readable_ids, corpus_path

NS = (
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
    'xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"'
)


def _p(text: str) -> str:
    return f'<w:p {NS}><w:r><w:t xml:space="preserve">{text}</w:t></w:r></w:p>'


TABLE_XML = load_shared_fixture_xml("sdt_table.xml")


def _build_sdt_table_docx() -> bytes:
    doc = Document()
    doc.add_paragraph("Intro paragraph.")
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    for xml in (TABLE_XML, _p("Outro paragraph.")):
        el = parse_xml(xml)
        if sectPr is not None:
            sectPr.addprevious(el)
        else:
            body.append(el)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


@pytest.fixture(scope="module")
def sdt_table_bytes() -> bytes:
    return _build_sdt_table_docx()


def _project(data: bytes, clean_view: bool) -> str:
    return extract_text_from_stream(io.BytesIO(data), clean_view=clean_view, include_appendix=False)


def _mapper_text(data: bytes, clean_view: bool) -> str:
    return DocumentMapper(Document(io.BytesIO(data)), clean_view=clean_view).full_text


_CC_TOKEN_RE = re.compile(r"\{#/?cc:\d+[^}]*\}")


def _strip_cc(text: str) -> str:
    """Drop content-control anchors so CC-0's assertions stay about ROWS.

    CC-1b started projecting `{#cc:N}` pairs around these very controls. That
    is correct and asserted separately below; this suite's subject is whether
    the wrapped rows and cells are VISIBLE AT ALL, which must keep reading the
    same way regardless of how much chrome CC-1 adds around them.
    """
    return _CC_TOKEN_RE.sub("", text)


def _row_line(text: str, first_cell: str) -> str:
    for line in text.splitlines():
        if _strip_cc(line).startswith(first_cell):
            return _strip_cc(line)
    raise AssertionError(f"no projected row starting with {first_cell!r} in:\n{text}")


# ---------------------------------------------------------------------------
# Cell-level sdt: <w:tr><w:sdt><w:sdtContent><w:tc>
# ---------------------------------------------------------------------------
@pytest.mark.parametrize("clean_view", [False, True])
def test_cell_level_sdt_content_is_projected(sdt_table_bytes, clean_view):
    text = _project(sdt_table_bytes, clean_view)
    assert "Contracting Officer" in text, (
        f"cell-level sdt content dropped from the {'clean' if clean_view else 'raw'} projection:\n{text}"
    )


@pytest.mark.parametrize("clean_view", [False, True])
def test_cell_level_sdt_row_keeps_its_column_separator(sdt_table_bytes, clean_view):
    """The dropped cell also dropped its ' | ', misaligning the whole column."""
    line = _row_line(_project(sdt_table_bytes, clean_view), "Role")
    assert " | " in line, f"row 1 rendered a single cell (column misalignment): {line!r}"
    assert line.split(" | ")[0].strip() == "Role"
    assert "Contracting Officer" in line.split(" | ")[1]


def test_first_row_column_count_drives_the_gfm_divider(sdt_table_bytes):
    """The divider is built from row 1's cell count — it must say two columns."""
    text = _project(sdt_table_bytes, clean_view=False)
    assert "--- | ---" in text, f"divider reports the wrong column count:\n{text}"


# ---------------------------------------------------------------------------
# Row-level sdt: <w:tbl><w:sdt><w:sdtContent><w:tr>
# ---------------------------------------------------------------------------
@pytest.mark.parametrize("clean_view", [False, True])
def test_row_level_sdt_row_is_projected(sdt_table_bytes, clean_view):
    text = _project(sdt_table_bytes, clean_view)
    assert "Approver" in text, f"row-level sdt row vanished entirely:\n{text}"
    assert "Jane Roe" in text
    assert _row_line(text, "Approver").startswith("Approver | Jane Roe")


# ---------------------------------------------------------------------------
# Nested sdt-in-sdt rows (FedRAMP w15:repeatingSectionItem shape)
# ---------------------------------------------------------------------------
@pytest.mark.parametrize("clean_view", [False, True])
def test_nested_repeating_section_row_is_projected(sdt_table_bytes, clean_view):
    text = _project(sdt_table_bytes, clean_view)
    assert "Repeated" in text, f"nested sdt (repeatingSectionItem) row vanished:\n{text}"
    assert _row_line(text, "Repeated").startswith("Repeated | Item One")


# ---------------------------------------------------------------------------
# Ordering, no duplication, and no bleed into neighbouring blocks
# ---------------------------------------------------------------------------
def test_rows_project_in_document_order_exactly_once(sdt_table_bytes):
    text = _project(sdt_table_bytes, clean_view=True)
    lines = [ln for ln in text.splitlines() if ln.strip()]
    lines = [_strip_cc(ln) for ln in lines]
    assert lines[0] == "Intro paragraph."
    assert lines[1] == "Role | Contracting Officer"
    assert lines[2] == "--- | ---"
    assert lines[3] == "Approver | Jane Roe"
    assert lines[4].startswith("Notes | ")
    assert lines[5] == "Repeated | Item One"
    assert lines[6] == "Outro paragraph."

    for token in ("Contracting Officer", "Approver", "Jane Roe", "Repeated", "Item One"):
        assert text.count(token) == 1, f"{token!r} projected {text.count(token)} times:\n{text}"


@pytest.mark.parametrize("clean_view", [False, True])
def test_block_level_sdt_inside_a_cell_is_projected(sdt_table_bytes, clean_view):
    """
    A BLOCK-level w:sdt (a content control wrapping a w:p rather than a w:tr or
    w:tc) is handled separately from the row/cell walk, by the block-level
    content-control descent in utils.docx._iter_block_children (Python) and
    iter_block_items (Node). Word emits this shape for rich-text controls and
    template placeholders.
    """
    line = _row_line(_project(sdt_table_bytes, clean_view), "Notes")
    assert "Approved without conditions." in line, f"block-level sdt still dropped: {line}"


# ---------------------------------------------------------------------------
# Virtual Text contract: ingest and mapper must produce identical text
# ---------------------------------------------------------------------------
@pytest.mark.parametrize("clean_view", [False, True])
def test_ingest_and_mapper_stay_synchronized(sdt_table_bytes, clean_view):
    projected = _project(sdt_table_bytes, clean_view)
    mapped = _mapper_text(sdt_table_bytes, clean_view)
    assert mapped == projected, (
        "DocumentMapper drifted from the ingest projection — offsets are now wrong.\n"
        f"ingest:\n{projected!r}\nmapper:\n{mapped!r}"
    )


@pytest.mark.parametrize("target", ["Contracting Officer", "Jane Roe", "Item One"])
def test_mapper_offsets_resolve_sdt_wrapped_content(sdt_table_bytes, target):
    """Text that is visible must also be addressable — a real run has to back it."""
    mapper = DocumentMapper(Document(io.BytesIO(sdt_table_bytes)))
    start = mapper.full_text.index(target)
    end = start + len(target)
    covering = [s for s in mapper.spans if s.run is not None and s.start < end and s.end > start]
    assert covering, f"no run-backed virtual-text span covers {target!r} at offset {start}"
    assert "".join(s.text for s in covering) == target


# ---------------------------------------------------------------------------
# A0.3 (apply half): visible must also mean editable. An edit whose target sits
# inside a row-level content control has to resolve, and its tracked change has
# to land INSIDE the wrapper rather than being hoisted out of it — otherwise
# accepting the revision would detach the text from the control.
# ---------------------------------------------------------------------------
def _tr_count(data: bytes) -> int:
    """Every w:tr in the package, including sdt-wrapped ones.

    python-docx's ``len(table.rows)`` reads ``./w:tr`` (direct children only),
    so it cannot see a row behind an sdt wrapper — the very rows under test.
    """
    return len(Document(io.BytesIO(data)).element.xpath("//w:tr"))


def test_modify_inside_row_level_sdt_applies_as_a_tracked_change(sdt_table_bytes):
    engine = RedlineEngine(io.BytesIO(sdt_table_bytes), author="A0 Reviewer")
    stats = engine.process_batch([ModifyText(target_text="Jane Roe", new_text="John Roe")])
    assert stats["edits_applied"] == 1, f"edit inside a row-level sdt did not resolve: {stats}"
    assert stats["edits_skipped"] == 0, f"edit inside a row-level sdt was skipped: {stats}"

    out = engine.save_to_stream().getvalue()
    root = Document(io.BytesIO(out)).element

    # Structure: the revision marks must be descendants of the wrapped w:tr.
    wrapped_row = root.xpath("//w:sdt/w:sdtContent/w:tr[.//w:t[contains(text(),'Approver')]]")
    assert wrapped_row, "the row-level content control no longer wraps a w:tr"
    # The engine diffs at token level: "Jane Roe" -> "John Roe" shares " Roe",
    # so only the differing token is redlined.
    ins_text = "".join(t.text or "" for t in wrapped_row[0].xpath(".//w:ins//w:t"))
    del_text = "".join(t.text or "" for t in wrapped_row[0].xpath(".//w:del//w:delText"))
    assert "John" in ins_text, f"insertion did not land inside the control (got {ins_text!r})"
    assert "Jane" in del_text, f"deletion did not land inside the control (got {del_text!r})"

    # Geometry: tracked edits must not add or drop rows.
    assert _tr_count(out) == _tr_count(sdt_table_bytes), "table row count changed"

    # Projection: raw view shows the CriticMarkup pair, clean view the result.
    assert "{--Jane--}{++John++}" in _project(out, clean_view=False)
    clean_line = _row_line(_project(out, clean_view=True), "Approver")
    assert clean_line.startswith("Approver | John Roe"), clean_line

    assert_word_readable_ids(out, "after editing inside a row-level content control: ")


def test_accepting_the_row_level_sdt_edit_keeps_the_control(sdt_table_bytes):
    """Accept must leave the control in place with the new value inside it."""
    engine = RedlineEngine(io.BytesIO(sdt_table_bytes), author="A0 Reviewer")
    engine.process_batch([ModifyText(target_text="Jane Roe", new_text="John Roe")])

    accept_engine = RedlineEngine(engine.save_to_stream(), author="A0 Reviewer")
    accept_engine.accept_all_revisions()
    accepted = accept_engine.save_to_stream().getvalue()

    root = Document(io.BytesIO(accepted)).element
    assert root.xpath("//w:sdt/w:sdtContent/w:tr"), "accept dissolved the row-level content control"
    assert _tr_count(accepted) == _tr_count(sdt_table_bytes), "accept changed the table row count"

    text = _project(accepted, clean_view=True)
    assert _row_line(text, "Approver").startswith("Approver | John Roe")
    assert "Jane Roe" not in text


# ---------------------------------------------------------------------------
# A5.0 — corpus scale check (skips cleanly when the document is absent).
#
# Was A0.5; moved into A5 on 2026-08-21 because it needs `corpus_path()`, a
# CC-3 deliverable, while CC-3 depends on CC-0. Now on that shared helper (the
# local resolver this file carried has been removed, as its author asked).
#
# Measured 2026-08-21 on the real template: 498,800 chars, floor comfortably
# cleared. Do NOT read a green run as proof that sdt traversal works — with
# row/cell sdt descent disabled the same document still projects 490,345 chars,
# so this floor passes with the bug present. The discriminating check is CC-3's
# tests/test_corpus_validation.py::test_a5_1_cell_level_sdt_content_is_visible
# _at_scale, which asserts on text reachable ONLY through a cell-level sdt and
# derives it from the document rather than hardcoding it. Keep both: this one
# guards scale, that one guards the repair. See PROGRESS.md.
# ---------------------------------------------------------------------------
def test_fedramp_corpus_projects_at_full_scale():
    text = _project(corpus_path("fedramp_ssp_rev4").read_bytes(), clean_view=True)
    assert len(text) > 400_000, f"clean view projected only {len(text):,} chars"

    # The shape CC-0 actually repaired: 371 cell-level SDTs in this template.
    assert "Information System Abbreviation" in text


# ---------------------------------------------------------------------------
# Outline builder walks the same tree (its offset math replays the projection)
# ---------------------------------------------------------------------------
def test_outline_finds_a_heading_inside_a_row_level_sdt():
    heading = f'<w:p {NS}><w:pPr><w:pStyle w:val="Heading1"/></w:pPr><w:r><w:t>Hidden Section Heading</w:t></w:r></w:p>'
    table_xml = f"""<w:tbl {NS}>
      <w:tblPr><w:tblW w:w="0" w:type="auto"/></w:tblPr>
      <w:tblGrid><w:gridCol w:w="4000"/></w:tblGrid>
      <w:sdt>
        <w:sdtPr><w:id w:val="201"/></w:sdtPr>
        <w:sdtContent>
          <w:tr><w:tc><w:tcPr><w:tcW w:w="4000" w:type="dxa"/></w:tcPr>{heading}</w:tc></w:tr>
        </w:sdtContent>
      </w:sdt>
    </w:tbl>"""

    doc = Document()
    doc.add_paragraph("Before.")
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    el = parse_xml(table_xml)
    if sectPr is not None:
        sectPr.addprevious(el)
    else:
        body.append(el)
    stream = io.BytesIO()
    doc.save(stream)
    data = stream.getvalue()

    text = _project(data, clean_view=False)
    assert "Hidden Section Heading" in text

    nodes = extract_outline(Document(io.BytesIO(data)), text, [text], [0])
    titles = [n.text for n in nodes]
    assert "Hidden Section Heading" in titles, f"outline missed the sdt-wrapped heading: {titles}"


# ---------------------------------------------------------------------------
# Regression guard: plain tables must be byte-identical to before the fix
# ---------------------------------------------------------------------------
def test_plain_table_projection_is_unchanged():
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            cell.text = f"R{r}C{c}"
    stream = io.BytesIO()
    doc.save(stream)
    data = stream.getvalue()

    text = _project(data, clean_view=True)
    assert "R0C0 | R0C1 | R0C2\n--- | --- | ---\nR1C0 | R1C1 | R1C2" in text
    assert _mapper_text(data, clean_view=True) == text


def test_merged_cells_still_deduplicate():
    """gridSpan/vMerge dedup relies on cell identity — the fix must preserve it."""
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            cell.text = f"R{r}C{c}"
    table.cell(0, 0).merge(table.cell(0, 1))
    stream = io.BytesIO()
    doc.save(stream)
    data = stream.getvalue()

    text = _project(data, clean_view=True)
    # The horizontally merged cell covers two grid columns but must project once.
    for token in ("R0C0", "R0C1", "R0C2"):
        assert text.count(token) == 1, f"{token!r} projected {text.count(token)} times:\n{text}"
    assert "R0C2" in text.split("\n---")[0], f"merged row lost its trailing cell:\n{text}"
    assert _mapper_text(data, clean_view=True) == text


# ---------------------------------------------------------------------------
# CC-1b — the same controls now carry anchors. Asserted here, next to the
# visibility guarantees, so a future change cannot restore visibility while
# silently dropping the anchors (or vice versa).
# ---------------------------------------------------------------------------
def test_row_and_cell_controls_carry_anchors(sdt_table_bytes):
    text = _project(sdt_table_bytes, clean_view=False)
    assert "Role | {#cc:1}Contracting Officer{#/cc:1}" in text, "cell-level control lost its inline anchors"
    assert "{#cc:2}Approver | Jane Roe{#/cc:2}" in text, "row-level control must bracket the whole row line"
    assert "Notes | {#cc:3}Approved without conditions.{#/cc:3}" in text, (
        "a block-level control inside a cell must anchor INLINE, not on its "
        "own lines — token lines would break the '|' row grammar"
    )
    assert _mapper_text(sdt_table_bytes, clean_view=False) == text
