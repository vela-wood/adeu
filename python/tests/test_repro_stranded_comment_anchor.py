"""
Demo run "Asteria v Northstar" (2026-08-12), defect B: the processed protective
order contained the fragment

    ... reasonably necessary for this litigationAttorney's Eyes Only;

i.e. a reviewer's commented phrase left standing alone in a sentence whose
surrounding words had been deleted around it.

This is NOT an engine bug. The output XML shows two SEPARATE tracked deletions
(w:id=4 and w:id=6) bracketing the comment anchor, so the agent deliberately
made two edits on either side of "Attorney's Eyes Only" and — presumably to
preserve Sarah Chen's comment anchor — kept the phrase itself. Each deletion is
individually legal and each was individually reported as applied. Nothing ever
cross-checked the two against each other, so the agent got no signal that the
sentence it left behind reads as gibberish once accepted.

The gap is therefore advisory, and the fix is an ADVISORY: a batch-level
warning, never a rejection. Editing around a foreign comment to keep its anchor
alive is a legitimate (if delicate) move, so the engine must still apply the
batch — it just has to say what it noticed.

The Node twin is node/packages/core/src/repro.stranded-comment-anchor.test.ts.
"""

import io
import re

from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine

CLAUSE_LEFT = ", unless the parties agree that a particular document or material produced is for "
ANCHORED = "Attorney's Eyes Only"
CLAUSE_RIGHT = " and is so designated"
SENTENCE = (
    "the officers, directors, and employees of the receiving party to whom disclosure "
    "is reasonably necessary for this litigation" + CLAUSE_LEFT + ANCHORED + CLAUSE_RIGHT + ";"
)


def _protective_order() -> io.BytesIO:
    """
    The demo's shape: one clause with Sarah Chen's comment anchored to the
    middle phrase and ordinary body text on both sides of it.
    """
    from docx import Document

    doc = Document()
    doc.add_paragraph("PROTECTIVE ORDER")
    doc.add_paragraph(SENTENCE)
    doc.add_paragraph("Nothing in this Order abridges any party's rights.")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Seed the foreign comment the way a reviewer would: a pure annotation,
    # no tracked change (new_text == target_text).
    eng = RedlineEngine(stream, author="Sarah Chen")
    eng.apply_edits(
        [
            ModifyText(
                target_text=ANCHORED,
                new_text=ANCHORED,
                comment="Should this tier survive the meet-and-confer?",
            )
        ]
    )
    out = eng.save_to_stream()
    out.seek(0)
    return out


def _protective_order_with_prior_left_deletion() -> io.BytesIO:
    """Same clause, but the text LEFT of the anchor is already deleted by
    another author in an earlier round."""
    src = _protective_order()
    eng = RedlineEngine(src, author="Opposing Counsel")
    eng.process_batch([ModifyText(target_text=CLAUSE_LEFT, new_text="")])
    out = eng.save_to_stream()
    out.seek(0)
    return out


def _protective_order_already_stranded() -> io.BytesIO:
    """Both sides already deleted by someone else, before this batch runs."""
    src = _protective_order()
    eng = RedlineEngine(src, author="Opposing Counsel")
    eng.process_batch(
        [
            ModifyText(target_text=CLAUSE_LEFT, new_text=""),
            ModifyText(target_text=CLAUSE_RIGHT, new_text=""),
        ]
    )
    out = eng.save_to_stream()
    out.seek(0)
    return out


def _stranded_warnings(res) -> list:
    """All batch-level advisory lines about a stranded anchor."""
    return [d for d in (res.get("skipped_details") or []) if re.search(r"stands alone|anchored text", d, re.I)]


# ────────────────────────────────────────────────────────────────────────────
# RED — the demo defect. Two legal deletions, one stranded phrase, silence.
# ────────────────────────────────────────────────────────────────────────────
def test_batch_deleting_both_sides_of_a_comment_anchor_warns():
    engine = RedlineEngine(_protective_order(), author="Agent")
    res = engine.process_batch(
        [
            ModifyText(target_text=CLAUSE_LEFT, new_text=""),
            ModifyText(target_text=CLAUSE_RIGHT, new_text=""),
        ]
    )

    # An advisory, NEVER a rejection: both edits still apply.
    assert res["edits_applied"] == 2
    assert res["edits_skipped"] == 0

    warnings = _stranded_warnings(res)
    assert len(warnings) == 1, f"expected one stranded-anchor advisory, got {res.get('skipped_details')}"

    # Actionable: which comment, whose, and what text was left behind. An
    # anonymous "a comment was stranded" reads like engine bookkeeping, which
    # is exactly how the demo run rationalised the broken sentence as success.
    assert "Com:1" in warnings[0]
    assert "Sarah Chen" in warnings[0]
    assert ANCHORED in warnings[0]
    assert warnings[0].lstrip().startswith("- Warning:")


def test_stranded_anchor_advisory_is_not_a_skip():
    engine = RedlineEngine(_protective_order(), author="Agent")
    res = engine.process_batch(
        [
            ModifyText(target_text=CLAUSE_LEFT, new_text=""),
            ModifyText(target_text=CLAUSE_RIGHT, new_text=""),
        ]
    )
    for report in res["edits"]:
        assert report["status"] == "applied"
        assert report.get("error") is None


# ────────────────────────────────────────────────────────────────────────────
# GREEN controls — the boundary. A false positive here trains the caller to
# ignore the warning, which is worse than not having it.
# ────────────────────────────────────────────────────────────────────────────
def test_single_edit_removing_the_anchored_text_too_is_silent():
    engine = RedlineEngine(_protective_order(), author="Agent")
    res = engine.process_batch([ModifyText(target_text=CLAUSE_LEFT + ANCHORED + CLAUSE_RIGHT, new_text="")])
    assert res["edits_applied"] == 1
    # The anchored text goes away WITH the clause, so nothing is stranded.
    assert _stranded_warnings(res) == []


def test_deleting_on_only_one_side_is_silent():
    engine = RedlineEngine(_protective_order(), author="Agent")
    res = engine.process_batch([ModifyText(target_text=CLAUSE_RIGHT, new_text="")])
    assert res["edits_applied"] == 1
    assert _stranded_warnings(res) == []


def test_ordinary_modify_under_a_foreign_comment_is_silent():
    engine = RedlineEngine(_protective_order(), author="Agent")
    res = engine.process_batch([ModifyText(target_text=ANCHORED, new_text="Outside Counsel Only")])
    assert res["edits_applied"] == 1
    assert _stranded_warnings(res) == []


# ────────────────────────────────────────────────────────────────────────────
# Attribution — the advisory is about what THIS batch caused.
# ────────────────────────────────────────────────────────────────────────────
def test_supplying_the_second_bracket_to_a_prior_deletion_warns():
    # The left side was already gone when the document arrived; deleting the
    # right side is the act that strands the phrase, so this batch owns it.
    engine = RedlineEngine(_protective_order_with_prior_left_deletion(), author="Agent")
    res = engine.process_batch([ModifyText(target_text=CLAUSE_RIGHT, new_text="")])

    assert res["edits_applied"] == 1
    warnings = _stranded_warnings(res)
    assert len(warnings) == 1
    assert ANCHORED in warnings[0]


def test_anchor_already_stranded_before_the_batch_is_silent():
    # Nagging about a condition the caller did not cause — and may not be
    # allowed to fix — is noise on every subsequent batch.
    engine = RedlineEngine(_protective_order_already_stranded(), author="Agent")
    res = engine.process_batch(
        [
            ModifyText(
                target_text="Nothing in this Order abridges any party's rights.",
                new_text="Nothing in this Order abridges any party's appellate rights.",
            )
        ]
    )
    assert res["edits_applied"] == 1
    assert _stranded_warnings(res) == []


def test_document_without_comments_is_silent():
    from docx import Document

    doc = Document()
    doc.add_paragraph("PROTECTIVE ORDER")
    doc.add_paragraph(SENTENCE)
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="Agent")
    res = engine.process_batch(
        [
            ModifyText(target_text=CLAUSE_LEFT, new_text=""),
            ModifyText(target_text=CLAUSE_RIGHT, new_text=""),
        ]
    )
    assert res["edits_applied"] == 2
    assert _stranded_warnings(res) == []
