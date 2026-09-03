import re
from io import BytesIO

import pytest
from docx import Document
from docx.oxml import parse_xml

from adeu.cli import handle_apply
from adeu.markup import format_ambiguity_error
from adeu.models import AcceptChange
from adeu.redline.engine import RedlineEngine
from tests.utils import approx_tokens


def test_terse_ambiguity_error_is_much_smaller():
    haystack = "An ambiguous phrase " * 20
    match_positions = [(m.start(), m.end()) for m in re.finditer(re.escape("phrase"), haystack)]

    msg_full = format_ambiguity_error(
        edit_index=1,
        target_text="phrase",
        haystack=haystack,
        match_positions=match_positions,
        terse=False,
    )

    msg_terse = format_ambiguity_error(
        edit_index=1,
        target_text="phrase",
        haystack=haystack,
        match_positions=match_positions,
        terse=True,
    )

    assert approx_tokens(msg_terse) <= 150
    assert approx_tokens(msg_terse) < approx_tokens(msg_full)
    assert "1." in msg_terse
    assert "2." in msg_terse
    assert "3." not in msg_terse


def test_full_ambiguity_error_is_unchanged_by_default():
    haystack = "An ambiguous phrase " * 10
    match_positions = [(m.start(), m.end()) for m in re.finditer(re.escape("phrase"), haystack)]

    msg_default = format_ambiguity_error(
        edit_index=1,
        target_text="phrase",
        haystack=haystack,
        match_positions=match_positions,
    )

    msg_explicit_false = format_ambiguity_error(
        edit_index=1,
        target_text="phrase",
        haystack=haystack,
        match_positions=match_positions,
        terse=False,
    )

    assert msg_default == msg_explicit_false
    assert "5." in msg_default  # Default cap is 5


def test_terse_stale_id_error_lists_at_most_eight_ids():
    # Create a document with 12 tracked changes
    doc = Document()
    p = doc.add_paragraph("Sample document text.")
    # Add ins nodes with w:id 1 through 12
    p_el = p._element
    ns_w = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
    for i in range(1, 13):
        xml_str = f'<w:ins {ns_w} w:id="{i}" w:author="Author"><w:r><w:t>text{i}</w:t></w:r></w:ins>'
        ins_node = parse_xml(xml_str)
        p_el.append(ins_node)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    # Full engine (default)
    engine_full = RedlineEngine(BytesIO(buf.getvalue()), terse_errors=False)
    act = AcceptChange(type="accept", target_id="Chg:888", comment=None)
    err_full = engine_full._action_not_found_error("Chg:888", "888", act)

    # Terse engine
    buf.seek(0)
    engine_terse = RedlineEngine(BytesIO(buf.getvalue()), terse_errors=True)
    err_terse = engine_terse._action_not_found_error("Chg:888", "888", act)

    assert "Chg:8" in err_terse
    assert "Chg:9" not in err_terse
    assert "(+4 more)" in err_terse
    assert "Chg:12" in err_full


def test_cli_apply_accepts_terse_errors_flag(tmp_path, capsys):
    # Create a docx with ambiguous target text
    doc = Document()
    for _ in range(5):
        doc.add_paragraph("The quick brown fox jumps over the lazy dog.")

    orig_path = tmp_path / "orig.docx"
    doc.save(orig_path)

    # Create changes json with an ambiguous edit
    changes_json = tmp_path / "changes.json"
    changes_json.write_text(
        '[{"type": "modify", "target_text": "quick brown fox", "new_text": "fast brown fox"}]',
        encoding="utf-8",
    )

    out_path = tmp_path / "out.docx"

    class DefaultArgs:
        original = orig_path
        changes = changes_json
        output = out_path
        author = "Test"
        json = False
        live = False
        partial = False
        terse_errors = False
        report = "standard"
        allow_major_deletions = False

    class TerseArgs(DefaultArgs):
        terse_errors = True

    # Default run (terse_errors=False)
    with pytest.raises(SystemExit) as exc_info_default:
        handle_apply(DefaultArgs())
    assert exc_info_default.value.code == 1
    err_default = capsys.readouterr().err

    # Terse run (terse_errors=True)
    with pytest.raises(SystemExit) as exc_info_terse:
        handle_apply(TerseArgs())
    assert exc_info_terse.value.code == 1
    err_terse = capsys.readouterr().err

    assert "RECOMMENDED:" in err_default
    assert "RECOMMENDED:" not in err_terse
    assert len(err_terse) < len(err_default)
    assert approx_tokens(err_terse) < approx_tokens(err_default)
