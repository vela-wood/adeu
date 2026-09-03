"""
Inserting a whole new SECTION in front of an existing heading, expressed the
way an agent naturally writes it::

    target_text = "# SCOPE"
    new_text    = "# NEW SECTION\\n\\nBody of the new section.\\n\\n# SCOPE"

The Node twin welded the replacement's first paragraph onto the heading's own
text ("# NEW SECTIONSCOPE"), promoted the body paragraph to a heading and left
an empty "# " behind (Asteria v Northstar demo, 2026-08-12). These are the
parity pins for that shape, plus the style-resolution rule the corruption
depended on: a template heading declares its heading-ness as <w:outlineLvl>
inside styles.xml under a house name, so a style-NAME test does not see it.
"""

from io import BytesIO

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement

from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine

NEW_SECTION = "# NEW SECTION\n\nBody of the new section.\n\n# SCOPE"


def _declare_outline_style(doc, style_id: str, outline_lvl: int) -> None:
    """Declares a CUSTOM paragraph style whose heading-ness lives in styles.xml."""
    style = OxmlElement("w:style")
    style.set(qn("w:type"), "paragraph")
    style.set(qn("w:customStyle"), "1")
    style.set(qn("w:styleId"), style_id)
    name = OxmlElement("w:name")
    name.set(qn("w:val"), "Legal Num 2 L1")
    style.append(name)
    based_on = OxmlElement("w:basedOn")
    based_on.set(qn("w:val"), "Normal")
    style.append(based_on)
    pPr = OxmlElement("w:pPr")
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(outline_lvl))
    pPr.append(outline)
    style.append(pPr)
    doc.styles.element.append(style)


def _add_styled_paragraph(doc, text: str, style_id: str):
    p = doc.add_paragraph(text)
    pPr = p._p.get_or_add_pPr()
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style_id)
    pPr.insert(0, pStyle)
    return p


def _style_of_paragraph_with_text(engine, text: str):
    for p in engine.doc.element.iter(qn("w:p")):
        runs = "".join(t.text or "" for t in p.iter(qn("w:t")))
        if runs == text:
            pPr = p.find(qn("w:pPr"))
            if pPr is None:
                return None
            pStyle = pPr.find(qn("w:pStyle"))
            return pStyle.get(qn("w:val")) if pStyle is not None else None
    return None


def _stream(doc) -> BytesIO:
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


class TestHeadingSectionInsert:
    def test_does_not_weld_the_replacement_onto_the_heading(self):
        doc = Document()
        _add_styled_paragraph(doc, "SCOPE", "Heading1")
        doc.add_paragraph("The protections conferred by this agreement are broad.")

        engine = RedlineEngine(_stream(doc), author="Reviewer")
        stats = engine.process_batch([ModifyText(type="modify", target_text="# SCOPE", new_text=NEW_SECTION)])
        assert stats["edits_skipped"] == 0, stats["skipped_details"]

        engine.accept_all_revisions(remove_comments=True)
        final = extract_text_from_stream(engine.save_to_stream(), clean_view=True)
        assert final == (
            "# NEW SECTION\n\nBody of the new section.\n\n# SCOPE\n\n"
            "The protections conferred by this agreement are broad."
        )

    def test_inserted_body_paragraph_is_not_promoted_to_a_heading(self):
        doc = Document()
        _add_styled_paragraph(doc, "SCOPE", "Heading1")
        doc.add_paragraph("Tail.")

        engine = RedlineEngine(_stream(doc), author="Reviewer")
        engine.process_batch([ModifyText(type="modify", target_text="# SCOPE", new_text=NEW_SECTION)])
        engine.accept_all_revisions(remove_comments=True)

        final = extract_text_from_stream(engine.save_to_stream(), clean_view=True)
        assert "# Body of the new section." not in final
        assert _style_of_paragraph_with_text(engine, "Body of the new section.") != "Heading1"

    def test_custom_outline_style_heading_survives_and_its_body_stays_plain(self):
        doc = Document()
        _declare_outline_style(doc, "LegalNum2L1", 0)
        _add_styled_paragraph(doc, "SCOPE", "LegalNum2L1")
        doc.add_paragraph("Tail.")

        engine = RedlineEngine(_stream(doc), author="Reviewer")
        stats = engine.process_batch([ModifyText(type="modify", target_text="# SCOPE", new_text=NEW_SECTION)])
        assert stats["edits_skipped"] == 0, stats["skipped_details"]

        engine.accept_all_revisions(remove_comments=True)
        final = extract_text_from_stream(engine.save_to_stream(), clean_view=True)
        assert final == "# NEW SECTION\n\nBody of the new section.\n\n# SCOPE\n\nTail."
        assert _style_of_paragraph_with_text(engine, "Body of the new section.") != "LegalNum2L1"

    def test_existing_heading_paragraph_keeps_its_template_style(self):
        doc = Document()
        _declare_outline_style(doc, "LegalNum2L1", 0)
        _add_styled_paragraph(doc, "SCOPE", "LegalNum2L1")
        doc.add_paragraph("Tail.")

        engine = RedlineEngine(_stream(doc), author="Reviewer")
        engine.process_batch([ModifyText(type="modify", target_text="# SCOPE", new_text=NEW_SECTION)])
        engine.accept_all_revisions(remove_comments=True)

        # Re-homing the heading's text into a freshly minted paragraph silently
        # drops the template style (and with it the legal auto-numbering).
        assert _style_of_paragraph_with_text(engine, "SCOPE") == "LegalNum2L1"

    def test_appending_a_body_paragraph_after_a_custom_outline_heading(self):
        """The scrub itself: the new paragraph must not inherit heading-ness."""
        doc = Document()
        _declare_outline_style(doc, "LegalNum2L1", 0)
        _add_styled_paragraph(doc, "SCOPE", "LegalNum2L1")

        engine = RedlineEngine(_stream(doc), author="Reviewer")
        stats = engine.process_batch(
            [
                ModifyText(
                    type="modify",
                    target_text="SCOPE",
                    new_text="SCOPE\n\nBody of the new section.",
                )
            ]
        )
        assert stats["edits_skipped"] == 0, stats["skipped_details"]

        engine.accept_all_revisions(remove_comments=True)
        final = extract_text_from_stream(engine.save_to_stream(), clean_view=True)
        assert final == "# SCOPE\n\nBody of the new section."

    def test_insertion_at_paragraph_start_relocates_host_text(self):
        doc = Document()
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("00.")

        engine = RedlineEngine(_stream(doc), author="Reviewer")
        stats = engine.process_batch([ModifyText(type="modify", target_text="00.", new_text="0.\n\n0 00.")])
        assert stats["edits_skipped"] == 0, stats["skipped_details"]

        engine.accept_all_revisions(remove_comments=True)
        final = extract_text_from_stream(engine.save_to_stream(), clean_view=True)
        assert final == "First paragraph.\n\n0.\n\n0 00."
