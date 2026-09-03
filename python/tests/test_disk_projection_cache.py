import json
import os
import sys
from pathlib import Path
from unittest.mock import patch

import pytest

from adeu.cli import main
from adeu.disk_cache import disk_projection_cache


def get_fixture_path(name: str) -> Path:
    p = Path(__file__).resolve()
    for parent in p.parents:
        if (parent / "shared" / "fixtures").is_dir():
            return parent / "shared" / "fixtures" / name
    raise FileNotFoundError(f"Could not find fixtures directory for {name}")


def make_heading_doc(path: Path) -> Path:
    """
    Writes a DOCX with real L1/L2 headings, so outline output and the budget
    guard heading map are non-empty (golden.docx has no headings at all).
    """
    from docx import Document

    doc = Document()
    doc.add_heading("Chapter One", level=1)
    doc.add_paragraph("Body text under chapter one. " * 20)
    doc.add_heading("Section One Point One", level=2)
    doc.add_paragraph("Body text under section one point one. " * 20)
    doc.add_heading("Chapter Two", level=1)
    doc.add_paragraph("Body text under chapter two. " * 20)
    doc.save(str(path))
    return path


def read_cache_views(cache_dir: Path) -> dict:
    """Returns the `views` object of the single on-disk cache entry."""
    cache_files = list(cache_dir.glob("*.json"))
    assert len(cache_files) == 1, cache_files
    return json.loads(cache_files[0].read_text(encoding="utf-8"))["views"]


@pytest.fixture(autouse=True)
def _isolate_cache(tmp_path, monkeypatch):
    """Isolate disk projection cache to a temporary directory per test."""
    cache_dir = tmp_path / "adeu_cache"
    monkeypatch.setenv("ADEU_CACHE_DIR", str(cache_dir))
    monkeypatch.delenv("ADEU_NO_CACHE", raising=False)
    monkeypatch.delenv("ADEU_DISABLE_DISK_CACHE", raising=False)
    disk_projection_cache.cache_dir = cache_dir
    disk_projection_cache.clear_stats()
    yield
    disk_projection_cache.cache_dir = None
    disk_projection_cache.clear_stats()


def _run_cli(args: list[str], capsys) -> str:
    test_args = ["adeu"] + args
    with patch.object(sys, "argv", test_args):
        try:
            main()
        except SystemExit as e:
            assert e.code == 0, f"CLI exited with code {e.code}"
    captured = capsys.readouterr()
    return captured.out


def _run_cli_warm(args: list[str], capsys) -> str:
    """Runs the CLI with DOCX loading fatal: a warm read must not re-parse."""
    with patch("adeu.cli._load_docx_or_exit", side_effect=AssertionError("warm read re-parsed the DOCX")):
        return _run_cli(args, capsys)


def test_second_read_is_byte_identical_and_hits_the_cache(tmp_path, capsys):
    doc_path = make_heading_doc(tmp_path / "headings.docx")

    # First read (cold cache miss)
    out1 = _run_cli(["extract", str(doc_path)], capsys)
    assert disk_projection_cache.stats["misses"] >= 1

    # Second read (warm cache hit): no DOCX parse allowed.
    out2 = _run_cli_warm(["extract", str(doc_path)], capsys)
    assert out1 == out2


def test_plain_cold_read_stays_lazy(tmp_path, capsys):
    """
    A plain read must not walk the outline or open the redline engine: those
    cost 25-69% on uncached reads and only `--mode outline` / `--mode changes`
    render them.
    """
    doc_path = make_heading_doc(tmp_path / "lazy.docx")

    with patch("adeu.cli._extract_outline_nodes", side_effect=AssertionError("outline walked on a plain read")):
        with patch(
            "adeu.cli._open_redline_engine_or_exit", side_effect=AssertionError("engine opened on a plain read")
        ):
            out = _run_cli(["extract", str(doc_path)], capsys)

    assert out
    views = read_cache_views(disk_projection_cache.cache_dir)
    assert set(views["raw"]) == {"base_text"}


def test_outline_mode_is_byte_identical_from_cache(tmp_path, capsys):
    doc_path = make_heading_doc(tmp_path / "outline.docx")

    # First read (outline mode, cold miss)
    out1 = _run_cli(["extract", str(doc_path), "--mode", "outline"], capsys)
    assert "Chapter One" in out1

    # Second read (outline mode, warm hit): served from cache, no re-parse.
    out2 = _run_cli_warm(["extract", str(doc_path), "--mode", "outline"], capsys)
    assert out1 == out2


def test_mtime_change_invalidates(tmp_path, capsys):
    fixture_path = get_fixture_path("golden.docx")
    doc_copy = tmp_path / "test_doc.docx"
    doc_copy.write_bytes(fixture_path.read_bytes())

    # Read 1
    out1 = _run_cli(["extract", str(doc_copy)], capsys)
    hits_1 = disk_projection_cache.stats["hits"]

    # Read 2 (cache hit)
    out2 = _run_cli(["extract", str(doc_copy)], capsys)
    assert out1 == out2
    assert disk_projection_cache.stats["hits"] == hits_1 + 1

    # Modify mtime
    st = os.stat(doc_copy)
    new_mtime = st.st_mtime + 10.0
    os.utime(doc_copy, (new_mtime, new_mtime))

    # Read 3 (cache invalidated due to mtime change -> miss)
    misses_before = disk_projection_cache.stats["misses"]
    out3 = _run_cli(["extract", str(doc_copy)], capsys)
    assert out1 == out3
    assert disk_projection_cache.stats["misses"] > misses_before


def test_disable_switch_and_unwritable_dir_are_non_fatal(tmp_path, monkeypatch, capsys):
    fixture_path = get_fixture_path("golden.docx")

    # 1. Test ADEU_NO_CACHE=1
    monkeypatch.setenv("ADEU_NO_CACHE", "1")
    hits_before = disk_projection_cache.stats["hits"]
    out1 = _run_cli(["extract", str(fixture_path)], capsys)
    assert out1
    assert disk_projection_cache.stats["hits"] == hits_before  # No hit when disabled

    monkeypatch.delenv("ADEU_NO_CACHE")

    # 2. Test unwritable cache dir
    unwritable_dir = tmp_path / "unwritable_cache"
    unwritable_dir.write_text(
        "i_am_a_file_not_a_directory"
    )  # Path exists as a file, so mkdir will fail with NotADirectoryError/OSError
    monkeypatch.setenv("ADEU_CACHE_DIR", str(unwritable_dir))
    disk_projection_cache.cache_dir = unwritable_dir

    # Extract should succeed without crashing
    out2 = _run_cli(["extract", str(fixture_path)], capsys)
    assert out2
    assert out1 == out2


def test_corrupt_cache_entry_is_ignored(tmp_path, capsys):
    fixture_path = get_fixture_path("golden.docx")
    doc_copy = tmp_path / "corrupt_test.docx"
    doc_copy.write_bytes(fixture_path.read_bytes())

    # Read 1 to populate cache
    out1 = _run_cli(["extract", str(doc_copy)], capsys)

    # Locate cache file and corrupt it
    cache_files = list(disk_projection_cache.cache_dir.glob("*.json"))
    assert len(cache_files) == 1
    cache_file = cache_files[0]
    cache_file.write_text("CORRUPTED_NOT_JSON {{{", encoding="utf-8")

    # Read 2: should ignore corrupt entry, perform fresh extract, and succeed
    out2 = _run_cli(["extract", str(doc_copy)], capsys)
    assert out1 == out2

    # Verify cache file was regenerated and contains valid JSON again
    repaired_text = cache_file.read_text(encoding="utf-8")
    parsed = json.loads(repaired_text)
    assert "key" in parsed


def test_view_merging_across_mode_switches(tmp_path, capsys):
    doc_path = make_heading_doc(tmp_path / "modes.docx")

    # 1. Plain read (cold): caches base_text only.
    out_full1 = _run_cli(["extract", str(doc_path)], capsys)
    assert set(read_cache_views(disk_projection_cache.cache_dir)["raw"]) == {"base_text"}

    # 2. Appendix mode (cold for the appendix projection): merges alongside base_text.
    assert _run_cli(["extract", str(doc_path), "--mode", "appendix"], capsys)

    # 3. Outline mode (cold for the heading map): merges alongside the other two.
    out_outline1 = _run_cli(["extract", str(doc_path), "--mode", "outline"], capsys)

    views = read_cache_views(disk_projection_cache.cache_dir)
    assert {"base_text", "text_with_appendix", "outline_nodes"} <= set(views["raw"])
    assert views["raw"]["outline_nodes"], "outline nodes should be non-empty for a document with headings"

    # 4. Both earlier modes still serve from the merged entry without re-parsing.
    assert _run_cli_warm(["extract", str(doc_path)], capsys) == out_full1
    assert _run_cli_warm(["extract", str(doc_path), "--mode", "outline"], capsys) == out_outline1


def test_budget_guard_output_identity_cold_vs_warm(tmp_path, capsys, monkeypatch):
    doc_path = make_heading_doc(tmp_path / "guard.docx")
    monkeypatch.setattr("adeu.payloads.response_budget_limit", lambda: 10)

    # Cold read (budget exceeded)
    with patch.object(sys, "argv", ["adeu", "extract", str(doc_path), "--page", "all"]):
        with pytest.raises(SystemExit) as exc1:
            main()
        assert exc1.value.code == 1
    err1 = capsys.readouterr().err
    assert "Chapter One" in err1, "the refusal must carry the L1 heading map"

    # Warm read: the guard's heading map comes from the cache, not a re-parse.
    with patch("adeu.cli._load_docx_or_exit", side_effect=AssertionError("warm read re-parsed the DOCX")):
        with patch.object(sys, "argv", ["adeu", "extract", str(doc_path), "--page", "all"]):
            with pytest.raises(SystemExit) as exc2:
                main()
            assert exc2.value.code == 1
    err2 = capsys.readouterr().err

    assert err1 == err2


def test_mode_changes_cache_hit_does_not_reopen_docx(capsys):
    fixture_path = get_fixture_path("golden.docx")

    # Cold read
    out1 = _run_cli(["extract", str(fixture_path), "--mode", "changes"], capsys)

    # Warm read: neither the DOCX nor the redline engine may be re-opened.
    with patch(
        "adeu.cli._open_redline_engine_or_exit", side_effect=AssertionError("Should not re-open DOCX on cache hit")
    ):
        out2 = _run_cli_warm(["extract", str(fixture_path), "--mode", "changes"], capsys)

    assert out1 == out2
