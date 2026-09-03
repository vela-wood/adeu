import json
import re
from io import BytesIO
from pathlib import Path

from adeu.cli import handle_extract
from adeu.ingest import _extract_text_from_doc
from adeu.mcp_components._response_builders import (
    build_changes_response,
    build_outline_response,
    build_paginated_response,
    build_search_response,
)
from adeu.models import ModifyText
from adeu.redline.comments import CommentsManager
from adeu.redline.engine import RedlineEngine
from tests.fixtures_synth import build_long_docx
from tests.utils import approx_tokens


class _Args:
    """Mock argparse Namespace for CLI tests."""

    def __init__(self, **kwargs):
        self.input = kwargs.get("input", None)
        self.live = kwargs.get("live", False)
        self.output = kwargs.get("output", None)
        self.force = kwargs.get("force", False)
        self.clean_view = kwargs.get("clean_view", False)
        self.mode = kwargs.get("mode", "full")
        self.changes_author = kwargs.get("changes_author", None)
        self.changes_offset = kwargs.get("changes_offset", 0)
        self.page = kwargs.get("page", None)
        self.search_query = kwargs.get("search_query", None)
        self.search_regex = kwargs.get("search_regex", False)
        self.search_case_insensitive = kwargs.get("search_case_insensitive", False)
        self.max_matches = kwargs.get("max_matches", 20)
        self.match_offset = kwargs.get("match_offset", 0)
        self.full_paragraph = kwargs.get("full_paragraph", False)
        self.outline_max_level = kwargs.get("outline_max_level", 2)
        self.outline_verbose = kwargs.get("outline_verbose", False)
        self.json = kwargs.get("json", False)
        self.no_chrome = kwargs.get("no_chrome", False)


def test_no_chrome_drops_file_path_header_and_prose(tmp_path: Path):
    doc_path = tmp_path / "test.docx"
    build_long_docx(doc_path, pages=3)

    from docx import Document

    doc = Document(str(doc_path))
    text = _extract_text_from_doc(doc)

    res_default = build_paginated_response(text, page=1, file_path=str(doc_path), no_chrome=False)
    res_no_chrome = build_paginated_response(text, page=1, file_path=str(doc_path), no_chrome=True)

    assert "**File Path:**" in str(res_default.content)
    assert "(synthetic page — a length-based chunk" in str(res_default.content)

    assert "**File Path:**" not in str(res_no_chrome.content)
    assert "(synthetic page — a length-based chunk" not in str(res_no_chrome.content)
    assert "Continues on page" not in str(res_no_chrome.content)


def test_no_chrome_page_content_is_byte_identical_apart_from_chrome(tmp_path: Path):
    doc_path = tmp_path / "test_single.docx"
    build_long_docx(doc_path, pages=1)

    from docx import Document

    doc = Document(str(doc_path))
    text = _extract_text_from_doc(doc)

    res_default = build_paginated_response(text, page=1, file_path=str(doc_path), no_chrome=False)
    res_no_chrome = build_paginated_response(text, page=1, file_path=str(doc_path), no_chrome=True)

    # In single page doc with no_chrome=True, content is pure body text.
    body_text = str(res_no_chrome.content)
    # The default content includes '> **File Path:** `path`\n\n' prefix and body_text
    default_content = str(res_default.content)
    header_prefix = f"> **File Path:** `{doc_path}`\n\n"
    assert default_content.startswith(header_prefix)
    extracted_body = default_content[len(header_prefix) :]

    assert extracted_body == body_text


def test_no_chrome_keeps_bare_page_marker_on_multipage(tmp_path: Path):
    doc_path = tmp_path / "test_multi.docx"
    build_long_docx(doc_path, pages=3)

    from docx import Document

    doc = Document(str(doc_path))
    text = _extract_text_from_doc(doc)

    res_no_chrome_p1 = build_paginated_response(text, page=1, file_path=str(doc_path), no_chrome=True)
    res_no_chrome_p2 = build_paginated_response(text, page=2, file_path=str(doc_path), no_chrome=True)

    content_p1 = str(res_no_chrome_p1.content)
    content_p2 = str(res_no_chrome_p2.content)

    assert content_p1.startswith("[p1/3]\n\n")
    assert content_p2.startswith("[p2/3]\n\n")
    assert "synthetic page" not in content_p1
    assert "**File Path:**" not in content_p1


def test_no_chrome_saves_tokens(tmp_path: Path):
    doc_path = tmp_path / "test_tokens.docx"
    build_long_docx(doc_path, pages=3)

    from docx import Document

    doc = Document(str(doc_path))
    text = _extract_text_from_doc(doc)

    res_default = build_paginated_response(text, page=1, file_path=str(doc_path), no_chrome=False)
    res_no_chrome = build_paginated_response(text, page=1, file_path=str(doc_path), no_chrome=True)

    default_tokens = approx_tokens(str(res_default.content))
    no_chrome_tokens = approx_tokens(str(res_no_chrome.content))

    assert default_tokens - no_chrome_tokens >= 20


def test_no_chrome_composes_with_json(tmp_path: Path, capsys):
    doc_path = tmp_path / "test_json.docx"
    build_long_docx(doc_path, pages=3)

    args = _Args(input=doc_path, page="1", json=True, no_chrome=True)
    handle_extract(args)

    captured = capsys.readouterr()
    stdout = captured.out.strip()

    data = json.loads(stdout)
    assert "markdown" in data
    md = data["markdown"]

    assert "**File Path:**" not in md
    assert "(synthetic page" not in md
    assert md.startswith("[p1/3]\n\n")


def test_no_chrome_search_zero_matches(tmp_path: Path):
    doc_path = tmp_path / "test_search_zero.docx"
    build_long_docx(doc_path, pages=1)

    from docx import Document

    doc = Document(str(doc_path))
    text = _extract_text_from_doc(doc)
    assert isinstance(text, str)

    res = build_search_response(
        text,
        search_query="nonexistent_xyz_query",
        search_regex=False,
        search_case_sensitive=False,
        page=None,
        file_path=str(doc_path),
        no_chrome=True,
    )

    content = str(res.content)
    assert "**File Path:**" not in content
    assert "**Search Results**" not in content
    assert "Verify your search spelling" not in content
    assert "No matches found" in content


def test_no_chrome_search_offset_past_total(tmp_path: Path):
    doc_path = tmp_path / "test_search_offset.docx"
    build_long_docx(doc_path, pages=1)

    from docx import Document

    doc = Document(str(doc_path))
    text = _extract_text_from_doc(doc)
    assert isinstance(text, str)

    res = build_search_response(
        text,
        search_query="Section",
        search_regex=False,
        search_case_sensitive=False,
        page=None,
        file_path=str(doc_path),
        match_offset=100,
        no_chrome=True,
    )

    content = str(res.content)
    assert "**File Path:**" not in content
    assert "**Search Results**" not in content
    assert "No matches in this window" in content


def test_no_chrome_search_keeps_regex_downgrade_note_with_hits(tmp_path: Path):
    """`no_chrome` strips chrome, not the notice that the query lost its regex semantics."""
    doc_path = tmp_path / "test_search_bad_regex.docx"
    from docx import Document

    doc = Document()
    doc.add_paragraph("The deposit is [USD 2500 and the rent is [USD 1200 per month.")
    doc.save(str(doc_path))

    text = _extract_text_from_doc(Document(str(doc_path)))
    assert isinstance(text, str)

    res = build_search_response(
        text,
        search_query="[USD",
        search_regex=True,
        search_case_sensitive=False,
        page=None,
        file_path=str(doc_path),
        no_chrome=True,
    )

    content = str(res.content)
    assert "**File Path:**" not in content
    assert "**Search Results**" not in content
    assert "not a valid regular expression" in content
    assert "### Match 1" in content


def test_no_chrome_deep_outline_level(tmp_path: Path):
    doc_path = tmp_path / "test_outline_deep.docx"
    from docx import Document

    doc = Document()
    doc.add_heading("Deep Heading", level=3)
    doc.save(str(doc_path))

    doc_obj = Document(str(doc_path))
    text = _extract_text_from_doc(doc_obj)
    assert isinstance(text, str)

    res = build_outline_response(
        doc_obj,
        projected_text=text,
        file_path=str(doc_path),
        outline_max_level=1,
        no_chrome=True,
    )

    content = str(res.content)
    assert "**File Path:**" not in content
    assert "Run `adeu extract" not in content
    assert "Call read_docx" not in content
    assert str(doc_path) not in content
    assert "No headings at level <=" in content


def _get_doc_text(docx_path: Path) -> str:
    from docx import Document

    res = _extract_text_from_doc(Document(str(docx_path)), clean_view=False)
    if isinstance(res, tuple):
        return str(res[0])
    return str(res)


def _build_clean_docx(path: Path) -> Path:
    """A document with body text but no tracked changes and no comments."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("The quick brown fox jumps over the lazy dog.")
    doc.save(str(path))
    return path


def _build_tracked_docx(path: Path) -> Path:
    """A document carrying tracked changes authored by `Jane Doe`."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("The quick brown fox jumps over the lazy dog.")
    base_path = path.parent / "tracked_base.docx"
    doc.save(str(base_path))

    with open(base_path, "rb") as f:
        engine = RedlineEngine(BytesIO(f.read()), author="Jane Doe")
    engine.process_batch([ModifyText(target_text="fox", new_text="cat")])

    with open(path, "wb") as f:
        f.write(engine.save_to_stream().getvalue())
    return path


def test_no_chrome_changes_clean_document_states_zero_counts(tmp_path: Path):
    doc_path = _build_clean_docx(tmp_path / "clean.docx")
    text = _get_doc_text(doc_path)

    res = build_changes_response(text, str(doc_path), is_cli=True, no_chrome=True)

    content = str(res.content)
    assert content == "0 change(s), 0 comment(s)"
    assert (res.structured_content or {})["markdown"] == "0 change(s), 0 comment(s)"
    assert "**File Path:**" not in content
    assert "**Changes ledger**" not in content


def test_no_chrome_changes_cli_clean_document_prints_zero_counts(tmp_path: Path, capsys):
    doc_path = _build_clean_docx(tmp_path / "clean_cli.docx")

    args = _Args(input=doc_path, mode="changes", no_chrome=True)
    handle_extract(args)

    assert capsys.readouterr().out.strip() == "0 change(s), 0 comment(s)"


def test_no_chrome_changes_unmatched_author_filter_states_zero_counts(tmp_path: Path):
    from docx import Document

    doc_path = _build_tracked_docx(tmp_path / "tracked.docx")
    text = _get_doc_text(doc_path)
    comments_data = CommentsManager(Document(str(doc_path))).extract_comments_data()

    matched = build_changes_response(
        text,
        str(doc_path),
        comments_data=comments_data,
        author_filter="Jane Doe",
        no_chrome=True,
    )
    assert "Chg:" in str(matched.content)

    res = build_changes_response(
        text,
        str(doc_path),
        comments_data=comments_data,
        author_filter="No Such Author",
        no_chrome=True,
    )
    assert str(res.content) == "0 change(s), 0 comment(s)"


def test_no_chrome_changes_offset_past_last_entry_keeps_counts(tmp_path: Path):
    doc_path = _build_tracked_docx(tmp_path / "tracked_offset.docx")
    text = _get_doc_text(doc_path)

    res = build_changes_response(text, str(doc_path), offset=500, no_chrome=True)

    content = str(res.content)
    assert re.fullmatch(r"[1-9]\d* change\(s\), \d+ comment\(s\)", content), content
