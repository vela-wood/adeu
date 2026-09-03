"""
Guard: a nested table must not leak its rows/cells into the enclosing table.

A `w:tbl` inside a `w:tc` is common (RFP response grids, signature blocks,
FedRAMP control tables). Any traversal that enumerates rows/cells with a
DESCENDANT search instead of a direct-child search corrupts these three ways:

  1. the outer row absorbs every inner cell
       "AfterInner | InnerA1 | InnerB1 | InnerA2 | InnerB2 | OuterB2"
  2. the inner rows are re-emitted as rows of the *outer* table
  3. each inner cell's text is therefore projected three times

Duplicated text is worse than missing text: `find_text`/`modify` match on the
projection, so a duplicated target reads as ambiguous or patches the wrong run,
and the outer row's column count stops matching its grid.

Python has always been correct here (`python-docx` resolves `Table.rows` via
`CT_Tbl.tr_lst` = `./w:tr` and `_Row.cells` via `CT_Row.tc_lst` = `./w:tc`,
both direct-child XPaths), while Node regressed on exactly this shape and had
to be repaired. Nothing pinned the Python behaviour, though, and the
sdt-transparent row/cell helpers in adeu.utils.docx now sit directly in this
path — so this file stops a future change there from reintroducing the leak in
the engine that currently gets it right.

Visibility only — no edit/apply semantics are exercised here.
"""

import io
import re

import pytest
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn

from adeu.ingest import extract_text_from_stream
from adeu.redline.mapper import DocumentMapper
from tests.sdt_fixtures import load_shared_fixture_xml

NS = (
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
    'xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"'
)


def _p(text: str) -> str:
    return f'<w:p><w:r><w:t xml:space="preserve">{text}</w:t></w:r></w:p>'


def _tc(inner: str) -> str:
    return f'<w:tc><w:tcPr><w:tcW w:w="3000" w:type="dxa"/></w:tcPr>{inner}</w:tc>'


OUTER_XML = load_shared_fixture_xml("nested_table_leakage.xml")

# Same shape, but the outer row and the inner row are behind content controls.
# This is where a naive "just use direct children" fix breaks: the traversal
# has to be sdt-transparent AND stop at the nested w:tbl.
OUTER_SDT_XML = f"""<w:tbl {NS}>
  <w:tblPr><w:tblW w:w="0" w:type="auto"/></w:tblPr>
  <w:tblGrid><w:gridCol w:w="3000"/><w:gridCol w:w="3000"/></w:tblGrid>
  <w:tr>{_tc(_p("OuterA1"))}{_tc(_p("OuterB1"))}</w:tr>
  <w:sdt>
    <w:sdtPr><w:alias w:val="NestRow"/><w:id w:val="201"/></w:sdtPr>
    <w:sdtContent>
      <w:tr>
        <w:tc><w:tcPr><w:tcW w:w="3000" w:type="dxa"/></w:tcPr>
          <w:tbl>
            <w:tblPr><w:tblW w:w="0" w:type="auto"/></w:tblPr>
            <w:tblGrid><w:gridCol w:w="1000"/><w:gridCol w:w="1000"/></w:tblGrid>
            <w:sdt>
              <w:sdtPr><w:id w:val="202"/></w:sdtPr>
              <w:sdtContent>
                <w:tr>{_tc(_p("InnerA1"))}{_tc(_p("InnerB1"))}</w:tr>
              </w:sdtContent>
            </w:sdt>
          </w:tbl>
          {_p("AfterInner")}
        </w:tc>
        {_tc(_p("OuterB2"))}
      </w:tr>
    </w:sdtContent>
  </w:sdt>
</w:tbl>"""

INNER_TOKENS = ("InnerA1", "InnerB1", "InnerA2", "InnerB2")


def _build(table_xml: str) -> bytes:
    doc = Document()
    doc.add_paragraph("Intro.")
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    element = parse_xml(table_xml)
    if sect_pr is not None:
        sect_pr.addprevious(element)
    else:
        body.append(element)
    doc.add_paragraph("Outro.")
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


def _project(data: bytes, clean_view: bool = True) -> str:
    return extract_text_from_stream(io.BytesIO(data), clean_view=clean_view, include_appendix=False)


_CC_TOKEN_RE = re.compile(r"\{#/?cc:\d+[^}]*\}")


def _lines(text: str) -> list[str]:
    """Non-blank projected lines, with content-control anchors stripped.

    CC-1b projects `{#cc:N}` pairs around the very controls this suite wraps
    its fixtures in. This suite's subject is whether the sdt-transparent walk
    STOPS at a nested `w:tbl` — a structural question that must read the same
    however much chrome CC-1 adds. The anchors themselves are asserted in the
    CC-1 suites.
    """
    return [_CC_TOKEN_RE.sub("", line) for line in text.splitlines() if line.strip()]


@pytest.fixture(scope="module")
def nested_bytes() -> bytes:
    return _build(OUTER_XML)


@pytest.fixture(scope="module")
def nested_sdt_bytes() -> bytes:
    return _build(OUTER_SDT_XML)


@pytest.mark.parametrize("clean_view", [False, True])
def test_outer_row_keeps_its_own_column_count(nested_bytes: bytes, clean_view: bool) -> None:
    text = _project(nested_bytes, clean_view)
    outer_row = next((line for line in _lines(text) if line.startswith("AfterInner")), None)
    assert outer_row is not None, f"outer row missing:\n{text}"
    for token in INNER_TOKENS:
        assert token not in outer_row, f"inner cell {token!r} leaked into the outer row: {outer_row}"
    assert len(outer_row.split(" | ")) == 2, f"outer row column count drifted: {outer_row}"


@pytest.mark.parametrize("token", [*INNER_TOKENS, "AfterInner"])
def test_each_nested_cell_is_projected_exactly_once(nested_bytes: bytes, token: str) -> None:
    text = _project(nested_bytes)
    count = text.count(token)
    assert count == 1, f"{token!r} projected {count}x (expected 1):\n{text}"


def test_nested_table_is_projected_in_document_order(nested_bytes: bytes) -> None:
    assert _lines(_project(nested_bytes)) == [
        "Intro.",
        "OuterA1 | OuterB1",
        "--- | ---",
        "InnerA1 | InnerB1",
        "--- | ---",
        "InnerA2 | InnerB2",
        "AfterInner | OuterB2",
        "Outro.",
    ]


def test_one_divider_per_table_not_per_leaked_row(nested_bytes: bytes) -> None:
    # Outer and inner tables are both 2 columns, so both dividers read
    # "--- | ---". The count is what proves the inner rows did not also
    # become rows of the outer table.
    assert _project(nested_bytes, clean_view=False).count("--- | ---") == 2


def test_nested_table_behind_content_controls(nested_sdt_bytes: bytes) -> None:
    """The sdt-transparent walk must still stop at the nested w:tbl."""
    assert _lines(_project(nested_sdt_bytes)) == [
        "Intro.",
        "OuterA1 | OuterB1",
        "--- | ---",
        # The single-row inner table gets its own divider after its first row.
        "InnerA1 | InnerB1",
        "--- | ---",
        "AfterInner | OuterB2",
        "Outro.",
    ]


@pytest.mark.parametrize("clean_view", [False, True])
@pytest.mark.parametrize("xml", [OUTER_XML, OUTER_SDT_XML], ids=["plain", "sdt"])
def test_ingest_and_mapper_stay_synchronized(xml: str, clean_view: bool) -> None:
    data = _build(xml)
    mapper = DocumentMapper(Document(io.BytesIO(data)), clean_view=clean_view)
    assert mapper.full_text == _project(data, clean_view), "DocumentMapper drifted from ingest"


@pytest.mark.parametrize("target", ["InnerB2", "AfterInner"])
def test_nested_cell_text_stays_addressable(nested_bytes: bytes, target: str) -> None:
    mapper = DocumentMapper(Document(io.BytesIO(nested_bytes)))
    start = mapper.full_text.find(target)
    assert start >= 0, f"{target!r} absent from the mapper projection"
    assert mapper.full_text.find(target, start + 1) == -1, f"{target!r} duplicated"
    end = start + len(target)
    covering = [s for s in mapper.spans if s.run is not None and s.start < end and s.end > start]
    assert covering, f"no run-backed span covers {target!r}"
    assert "".join(s.text for s in covering) == target
