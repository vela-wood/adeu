# FILE: tests/test_corpus_validation.py
"""Corpus validation against real government documents.

These run against real public-sector .docx/.dotx files that are deliberately
NOT committed. `corpus_path()` skips cleanly when a document is absent, so CI
is green without a single download; the optional corpus job fetches first.

* A5.9 — the fetch mechanism itself.
* A5.7 (partial) — the .dotx opens through the standard path.
* A5.8 (partial) — the negative `w:sdt` id survives a no-op round trip.
* A5.1 (partial) — the CC-0 repair holds at production scale, asserted in the
  DISCRIMINATING form: cell-level SDT content that is invisible without the fix.

The deferred assertions are listed in PROGRESS.md against the task that unblocks
each one. They are not stubbed here: a skipped test that can never run is
indistinguishable from a passing one at a glance, and this suite's whole purpose
is to not be vacuously green.
"""

from __future__ import annotations

import io
import json
import re
import subprocess
import sys
import zipfile

import pytest

from adeu.ingest import extract_text_from_stream
from tests.utils import CLI_OUTPUT_ENCODING, CORPUS_MANIFEST, corpus_path

REPO_ROOT = CORPUS_MANIFEST.parents[2]


def _project(data: bytes, clean_view: bool = True) -> str:
    return extract_text_from_stream(io.BytesIO(data), clean_view=clean_view, include_appendix=False)


# ---------------------------------------------------------------------------
# A5.9 — fetch mechanism smoke (no network)
# ---------------------------------------------------------------------------


def _run_fetch_script(*args: str) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, str(REPO_ROOT / "scripts" / "fetch_corpus.py"), *args],
        capture_output=True,
        encoding=CLI_OUTPUT_ENCODING,
        cwd=REPO_ROOT,
    )


def test_a5_9_fetch_corpus_list_reports_every_manifest_key():
    """`fetch_corpus.py --list` exits 0 and reports presence per manifest key.

    The one part of the corpus machinery that must work on a machine with no
    corpus and no network — it is how a developer finds out what to download.
    """
    result = _run_fetch_script("--list")
    assert result.returncode == 0, result.stderr
    keys = json.loads(CORPUS_MANIFEST.read_text(encoding="utf-8"))["documents"]
    for key in keys:
        line = next((ln for ln in result.stdout.splitlines() if ln.startswith(key)), None)
        assert line is not None, f"--list omitted {key!r}:\n{result.stdout}"
        assert re.search(r"\b(present|missing)\b", line), f"no on-disk status for {key!r}: {line}"


def test_a5_9_fetch_corpus_rejects_an_unknown_key():
    """A typo must fail loudly, not fetch nothing and exit 0.

    `--only` naming a key that does not exist is the shape that would otherwise
    let the optional CI job "succeed" having downloaded nothing at all.
    """
    result = _run_fetch_script("--only", "no_such_document")
    assert result.returncode == 2
    assert "unknown manifest key" in result.stderr
    assert "known keys:" in result.stderr, "the error must be self-service"


def test_corpus_path_raises_on_an_unknown_key():
    """The test helper draws the same distinction as the fetcher.

    Absent document → skip; unknown key → raise. A helper that skipped on both
    would turn every typo into a permanently green test.
    """
    with pytest.raises(KeyError, match="unknown corpus key"):
        corpus_path("no_such_document")


# ---------------------------------------------------------------------------
# A5.1 (partial) — CC-0 at production scale, in discriminating form
# ---------------------------------------------------------------------------


def _cell_level_sdt_texts(data: bytes) -> list[str]:
    """Every distinct text that lives inside a cell-level SDT (`sdtContent > w:tc`).

    Derived from the document rather than hardcoded: upstream revises these
    templates in place (spec-corpus §1), and a list of literal strings would rot
    into a skip-shaped failure. Each entry is a single `w:t` node's text, not a
    join of several — runs split at arbitrary points and the projection
    reassembles them with its own whitespace rules, so joined text is not a
    substring of the output even when nothing is wrong.
    """
    from lxml import etree

    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    with zipfile.ZipFile(io.BytesIO(data)) as package:
        root = etree.fromstring(package.read("word/document.xml"))

    texts: set[str] = set()
    for sdt in root.iter(f"{W}sdt"):
        content = sdt.find(f"{W}sdtContent")
        if content is None:
            continue
        if not any(etree.QName(child).localname == "tc" for child in content):
            continue
        for node in content.iter(f"{W}t"):
            value = (node.text or "").strip()
            if len(value) >= 20:
                texts.add(value)
    return sorted(texts)


def test_a5_1_cell_level_sdt_content_is_visible_at_scale():
    """The FedRAMP SSP's 371 cell-level SDTs project their text.

    A0.5 asserts a 400,000-char floor on this document, and PROGRESS.md records
    that the floor does not discriminate: with row/cell descent disabled the
    same file still projects 490,345 chars, so that guard passes with the bug
    present. This one asserts on content reachable ONLY through a cell-level
    `sdtContent > w:tc`, so it fails the moment the CC-0 repair regresses.
    """
    data = corpus_path("fedramp_ssp_rev4").read_bytes()
    text = _project(data)

    assert len(text) > 400_000, f"clean view projected only {len(text):,} chars"

    cell_texts = _cell_level_sdt_texts(data)
    # ~95% of the 2026-08-21 scan's 371 cell-level controls, per spec-corpus §1.
    assert len(cell_texts) >= 20, f"fixture drifted: only {len(cell_texts)} cell-level texts"

    missing = [value for value in cell_texts if value not in text]
    assert not missing, (
        f"{len(missing)} of {len(cell_texts)} cell-level SDT texts are invisible in the "
        f"projection (CC-0 data loss): {missing[:5]}"
    )


def test_a5_1_no_raw_sdt_markup_leaks_into_the_projection():
    """Descending into `w:sdt` must not emit the wrapper itself.

    The failure mode opposite to CC-0: a traversal that "fixes" invisibility by
    stringifying the element would put OOXML in front of the model. Cheap to
    check, and it covers the whole corpus rather than one document.
    """
    text = _project(corpus_path("fedramp_ssp_rev4").read_bytes())

    for token in ("<w:sdt", "sdtContent", "w:sdtPr", "showingPlcHdr"):
        assert token not in text, f"raw OOXML {token!r} leaked into the text projection"


# ---------------------------------------------------------------------------
# A5.7 (partial) — .dotx through the standard path
# ---------------------------------------------------------------------------


def test_a5_7_the_fixture_really_is_a_template():
    """Guards the guard: A5.7 means nothing if the file stops being a .dotx."""
    with zipfile.ZipFile(io.BytesIO(corpus_path("odot_uic_drywell").read_bytes())) as package:
        content_types = package.read("[Content_Types].xml").decode("utf-8")
    assert "template.main+xml" in content_types


def test_a5_7_dotx_template_opens_through_the_standard_path():
    """A .dotx is an OPC package like any other; the engine must not sniff content types.

    The ledger/picture halves of A5.7 need CC-1. This is the half that is
    testable today. It failed until CC-11: `python-docx`'s `Document()` accepts
    exactly one main-part content type and raised `ValueError: ... is not a Word
    file` on `template.main+xml`, which the CLI surfaced as an unhandled
    traceback, while `@adeu/core` read the same file happily. `adeu.utils.opc`
    now registers the template and macro-enabled content types against
    `DocumentPart`, so the part keeps its own content type and a `.dotx` saves
    back as a `.dotx` — see `tests/test_opc_document_types.py` for the
    round-trip guard.
    """
    text = _project(corpus_path("odot_uic_drywell").read_bytes())
    assert text.strip(), "the .dotx projected nothing at all"
    assert "<w:sdt" not in text


# ---------------------------------------------------------------------------
# A5.8 (partial) — negative w:sdt id survives a round trip
# ---------------------------------------------------------------------------

_SDT_ID_RE = re.compile(r'<w:sdt>.*?<w:id w:val="(-?\d+)"', re.DOTALL)


def test_a5_8_negative_sdt_id_round_trips_untouched():
    """`w:sdt/w:id` is signed, and the wild contains negative values.

    AI_CONTEXT §8's ST_LongHexNumber lesson does NOT apply here — `w:id` on an
    sdt is `ST_DecimalNumber`, where a negative value is legal and Word keeps
    it. The risk is Adeu "helpfully" normalising it. Surgical mode must leave
    the bytes alone, so the id is asserted identical after a no-op open→save.

    Note where it lives: the negative id in this document is in `word/footer1.xml`,
    not `word/document.xml`. Scanning only the main part finds nothing and the
    test passes vacuously — which is how this was first written.
    """
    data = corpus_path("hc_diagnostic_nonlab").read_bytes()

    def sdt_ids(package_bytes: bytes) -> dict[str, list[str]]:
        found: dict[str, list[str]] = {}
        with zipfile.ZipFile(io.BytesIO(package_bytes)) as package:
            for name in package.namelist():
                if not name.endswith(".xml"):
                    continue
                ids = _SDT_ID_RE.findall(package.read(name).decode("utf-8", "replace"))
                if ids:
                    found[name] = ids
        return found

    before = sdt_ids(data)
    negative = [(part, value) for part, ids in before.items() for value in ids if value.startswith("-")]
    assert negative, f"fixture no longer carries a negative sdt id (found {before})"

    from adeu.redline.engine import RedlineEngine

    saved = RedlineEngine(io.BytesIO(data)).save_to_stream().getvalue()

    assert sdt_ids(saved) == before, "a no-op save rewrote the sdt ids"


# ---------------------------------------------------------------------------
# CC-1c — checkbox census
#
# Engine-independent corpus facts, so no node twin: these read XML, not
# projections. They exist because spec-projection.md §4 makes assumptions about
# what checkboxes look like in the wild, and the wild is right here.
# ---------------------------------------------------------------------------

_W14_CHECKBOX = re.compile(r"<w14:checkbox\b.*?</w14:checkbox>", re.S)
_W14_CHECKED = re.compile(r'<w14:checked\s+w14:val="([^"]*)"')
_CHECKED_STATE = re.compile(r"<w14:checkedState\s+([^/>]*)/>")
_UNCHECKED_STATE = re.compile(r"<w14:uncheckedState\s+([^/>]*)/>")
_LEGACY_CHECKBOX = re.compile(r"<w:checkBox>.*?</w:checkBox>", re.S)
_SDT_ELEMENT = re.compile(r"<w:sdt>.*?</w:sdt>", re.S)

BALLOT_EMPTY = "\u2610"
BALLOT_X = "\u2612"

# key -> (w14:checkbox count, checked count). Every corpus document; the zeros
# are as load-bearing as the counts, since a document that GAINS checkboxes
# changes what §4 has to cope with.
CHECKBOX_CENSUS = {
    "ca_talent_recruitment": (0, 0),
    "dau_acquisition_plan": (0, 0),
    "fedramp_sar": (0, 0),
    "fedramp_ssp_appx_a_moderate": (3_804, 0),
    "fedramp_ssp_rev4": (3_881, 0),
    "fedramp_ssp_rev5": (3, 0),
    "hc_diagnostic_nonlab": (0, 0),
    "odot_uic_drywell": (19, 0),
    "on_juries_form1": (0, 0),
    "wawd_esi_agreement": (0, 0),
}


def _word_xml(data: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(data)) as package:
        return "\n".join(
            package.read(name).decode("utf-8", "replace")
            for name in package.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        )


@pytest.mark.parametrize("key", sorted(CHECKBOX_CENSUS))
def test_cc1c_checkbox_census_is_pinned(key):
    """How many checkboxes each corpus document has, and how many are ticked."""
    data = corpus_path(key).read_bytes()
    xml = _word_xml(data)

    total = len(_W14_CHECKBOX.findall(xml))
    checked = sum(1 for value in _W14_CHECKED.findall(xml) if value in ("1", "true"))

    assert (total, checked) == CHECKBOX_CENSUS[key], (
        f"{key}: {total} checkboxes / {checked} checked, expected {CHECKBOX_CENSUS[key]}"
    )


def test_cc1c_the_corpus_contains_no_legacy_form_field_checkboxes():
    """`w14:checkbox` is the only checkbox mechanism here — the legacy one is absent.

    Word has two: the modern content control (`w14:checkbox`, Word 2010+) and the
    legacy form field (`w:fldChar` + `w:ffData/w:checkBox`, Word 97 era), which is
    NOT a content control and would not be reached by any of this initiative's
    traversal work. spec-projection.md §4 only describes the modern one.

    Across ~7,700 checkboxes in ten real government documents there is not one
    legacy field, which is what makes that scope choice safe rather than lucky.
    Pinned so that adding a corpus document containing legacy fields fails here
    and forces the scope question to be re-asked, instead of silently projecting
    nothing where a form has checkboxes.
    """
    legacy = {}
    for key in sorted(CHECKBOX_CENSUS):
        xml = _word_xml(corpus_path(key).read_bytes())
        found = len(_LEGACY_CHECKBOX.findall(xml))
        if found:
            legacy[key] = found

    assert legacy == {}, f"legacy w:checkBox form fields found, §4 does not cover these: {legacy}"


def test_cc1c_every_corpus_checkbox_uses_the_same_glyph_pair():
    """`MS Gothic` 2612/2610 throughout — the state glyphs are not per-document.

    `w14:checkedState`/`w14:uncheckedState` are free-form (font + code point), so
    a document may legitimately use Wingdings `F0FE`, a private-use code point
    whose meaning depends entirely on the font. That case would break any
    projection that recognises checkboxes by their character.

    It does not occur here: all four documents that have checkboxes use the same
    pair. This is the empirical half of why §4 projects from `w14:checked` rather
    than from the glyph — the glyph is *usually* recognisable, and "usually" is
    not a contract.
    """
    pairs = set()
    for key in sorted(k for k, (total, _) in CHECKBOX_CENSUS.items() if total):
        xml = _word_xml(corpus_path(key).read_bytes())
        pairs.update(_CHECKED_STATE.findall(xml))
        pairs.update(_UNCHECKED_STATE.findall(xml))

    fonts = {re.search(r'w14:font="([^"]*)"', p).group(1) for p in pairs if 'w14:font="' in p}
    vals = {re.search(r'w14:val="([^"]*)"', p).group(1) for p in pairs if 'w14:val="' in p}

    assert fonts == {"MS Gothic"}, f"unexpected checkbox state fonts: {fonts}"
    assert vals == {"2612", "2610"}, f"unexpected checkbox state code points: {vals}"


def test_cc1c_the_corpus_has_ballot_glyphs_that_are_not_checkboxes():
    """Two bare `☐` characters in `odot_uic_drywell`, outside every content control.

    Segoe UI Symbol runs in ordinary prose — a human typing a box rather than
    inserting a control. They are TEXT, and must survive projection as `☐`.

    This is the trap for CC-1c's implementation half. The obvious way to satisfy
    A1.8 ("the view contains `[x]`/`[ ]` and NO `☒`/`☐` characters") is to
    substitute on the character; do that and these two turn into `[ ]`, inventing
    two checkboxes that do not exist, in a document that has 19 real ones to hide
    among. The substitution has to be driven by the `w14:checkbox` control, and
    A1.8's "no glyphs" clause has to be read as scoped to control content.
    """
    xml = _word_xml(corpus_path("odot_uic_drywell").read_bytes())
    outside = _SDT_ELEMENT.sub("", xml)

    assert xml.count(BALLOT_EMPTY) == 21
    assert outside.count(BALLOT_EMPTY) == 2, "bare ballot glyphs outside any control"
    assert BALLOT_X not in xml, "no ticked glyph anywhere in the corpus"


def test_cc1c_no_corpus_checkbox_is_ticked():
    """Nothing in the corpus exercises the `[x]` half of A1.8.

    ~7,700 checkboxes, every one unchecked: these are blank templates, which is
    what public bodies publish. So the corpus can validate `[ ]` at scale and can
    say nothing at all about `[x]`, and a corpus-driven implementation would be
    half-tested while looking thoroughly exercised.

    The checked path is covered by synthetic fixtures and by the live-Word probes
    in `test_live_word_content_controls.py` instead. Recorded as a test so the
    gap is a stated fact rather than an omission nobody noticed.
    """
    ticked = {
        key: sum(1 for v in _W14_CHECKED.findall(_word_xml(corpus_path(key).read_bytes())) if v in ("1", "true"))
        for key in sorted(k for k, (total, _) in CHECKBOX_CENSUS.items() if total)
    }
    assert set(ticked.values()) == {0}, f"a corpus document now has ticked checkboxes: {ticked}"


# ---------------------------------------------------------------------------
# CC-3b — the A5 examples that needed CC-1, CC-2, CC-4 and CC-5
#
# A5.2-A5.6 asserted against the ledger, the gates and `set_field`. No new
# mechanism: `corpus_path()` and the skip-if-missing discipline above are
# unchanged, and every floor below is the frozen number from A5, which is ~95%
# of the 2026-08-21 scan so upstream revisions do not turn these red.
#
# Floors, not equalities, everywhere except where an exact count is the point.
# These are living government templates that get revised in place (spec-corpus
# §1); pinning 162 rather than ">= 154" would make the suite a tripwire for
# other people's editing rather than for our own regressions.
# ---------------------------------------------------------------------------

STOCK_PLACEHOLDER = "Click or tap here to enter text."


def _ledger_for(key: str):
    """(entries, rendered lines, protection, raw projection) for a corpus doc."""
    from docx import Document

    from adeu.fields import collect_fields, read_document_protection, render_line
    from adeu.ingest import _extract_text_from_doc

    doc = Document(io.BytesIO(corpus_path(key).read_bytes()))
    text = _extract_text_from_doc(doc, clean_view=False, include_appendix=False)
    if isinstance(text, tuple):
        text = text[0]
    entries = collect_fields(doc, text, None)
    width = max((len(f"CC:{e.ordinal}") for e in entries), default=4)
    lines = [render_line(e, width) for e in entries]
    return entries, lines, read_document_protection(doc), text


# --- A5.1 (ledger half) ----------------------------------------------------


def test_a5_1_ledger_class_and_state_floors():
    """The FedRAMP SSP's ledger, by class and by state.

    The scale assertion in A5.0 does not discriminate the CC-0 bug; this does,
    on a different axis. 371 of these controls wrap table cells, and every one
    of them was invisible before CC-0 — so the `table cell` floor fails loudly
    if cell descent ever regresses, where a character count merely sags.

    Left OUT deliberately: A5.1's python/node identical-count parity assertion.
    It is still blocked on two non-sdt divergences (emphasis-run coalescing, and
    header lines node projects that python omits) — CC-10's closing note
    measured that closing the page-break gap alone flipped the sign rather than
    closing the difference. Asserting it here would be red for reasons that have
    nothing to do with content controls.
    """
    entries, lines, _protection, text = _ledger_for("fedramp_ssp_rev4")

    assert len(entries) >= 4_750, f"ledger lists {len(entries)} controls"

    classes: dict[str, int] = {}
    for entry in entries:
        classes[entry.cls_word] = classes.get(entry.cls_word, 0) + 1

    for cls_word, floor in (
        ("checkbox", 3_690),
        ("text", 430),
        ("date", 315),
        ("richtext", 260),
        ("combobox", 25),
        ("dropdown", 19),
        ("picture", 4),
    ):
        assert classes.get(cls_word, 0) >= floor, f"{cls_word}: {classes.get(cls_word, 0)} < {floor} ({classes})"

    # 94, not 89: the five missing controls lived in a header referenced by a
    # section break wrapped in a content control (CC-17).
    assert sum(1 for e in entries if e.bound) >= 94
    assert sum(1 for e in entries if e.empty) >= 680
    assert sum(1 for e in entries if e.container_kind == "table cell") >= 350
    assert sum(1 for line in lines if "TEMPORARY" in line) >= 2


def test_cc17_a_section_break_inside_a_content_control_is_still_a_section():
    """CC-17 — the mechanism, on a synthetic document.

    python-docx enumerates sections with `./w:body/w:p/w:pPr/w:sectPr`, which
    takes only DIRECT children of the body. Wrap the section-terminating
    paragraph in a content control — as Word does for a cover page inserted from
    the document-part gallery — and the section vanishes from `doc.sections`
    along with the header it references.
    """
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    from adeu.utils.docx import iter_sections_including_wrapped

    doc = Document()
    body = doc.element.body

    doc.add_paragraph("wrapped")
    para = doc.paragraphs[-1]._p
    pPr = para.get_or_add_pPr()
    pPr.append(OxmlElement("w:sectPr"))

    # Wrap that paragraph in w:sdt/w:sdtContent, in place.
    sdt = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    sdt.append(content)
    para.addprevious(sdt)
    content.append(para)

    assert body.find(qn("w:sdt")) is not None, "fixture did not wrap the paragraph"

    visible = len(doc.sections)
    found = len(list(iter_sections_including_wrapped(doc)))
    assert found == visible + 1, (
        f"python-docx sees {visible} sections; the walk must see one more (the wrapped one), saw {found}"
    )


def test_cc17_the_wrapped_sections_header_controls_are_reachable():
    """CC-17 — the consequence, on the document that exposed it.

    `fedramp_ssp_rev4` puts two section breaks inside content controls. The
    header one of them references holds five DATA-BOUND controls — exactly what
    `set_field` exists to fill — and they were unreachable: `resolve_field`
    cannot resolve what `collect_fields` never listed. An agent asked to fill
    the SSP's header got "no such field" for a field the user is looking at.
    """
    from adeu.fields import resolve_field

    entries, _lines, _prot, _text = _ledger_for("fedramp_ssp_rev4")

    assert len(entries) == 5_007
    assert sum(1 for e in entries if e.bound) == 94

    # header2 holds five controls: four tagged, one untagged. Each tagged one is
    # a SECOND occurrence of a tag that already existed elsewhere, so the proof
    # is the count, not mere resolvability — before the fix each was one lower.
    for tag, expected in (
        ("cspname", 20),
        ("informationsystemname", 25),
        ("versionnumber", 5),
        ("versiondate", 5),
    ):
        hits = resolve_field(entries, tag, match_mode="all")
        assert len(hits) == expected, (
            f"{tag!r}: {len(hits)} != {expected} — the occurrence in the header "
            f"behind a wrapped section break is missing"
        )


def test_cc17_no_opc_part_is_projected_twice():
    """The other direction of the same bug.

    Walking every header/footer part in the package would also fix the count —
    and would project inherited headers, unreferenced orphans and first-page
    headers Word never renders. The node port was already corrected for exactly
    that, so pin it: every part walked appears exactly once.
    """
    import io

    from docx import Document

    from adeu.utils.docx import iter_document_parts_with_kind

    for key in ("fedramp_ssp_rev4", "fedramp_ssp_rev5", "fedramp_sar"):
        doc = Document(io.BytesIO(corpus_path(key).read_bytes()))
        names = []
        for part, kind in iter_document_parts_with_kind(doc):
            if kind in ("header", "footer"):
                names.append(str(part.part.partname))
        assert len(names) == len(set(names)), f"{key}: part projected twice: {names}"


def test_a5_1_checkboxes_project_as_ascii_not_glyphs():
    """3,881 checkboxes reach the raw view as `[ ]`, and no ballot glyph survives.

    Scoped to the whole projection on purpose. A1.8's "no glyphs" clause is
    scoped to control CONTENT elsewhere (see the odot test above, where two bare
    glyphs live outside any control) — this document has no such bare glyphs, so
    here the stronger form is the true one and worth pinning.
    """
    _entries, _lines, _protection, text = _ledger_for("fedramp_ssp_rev4")

    assert text.count("[x]") + text.count("[ ]") >= 3_690
    assert BALLOT_X not in text and BALLOT_EMPTY not in text


# --- A5.2 — dau_acquisition_plan: locks and anonymity ----------------------


def test_a5_2_locks_and_anonymity():
    """162 controls, 48 content-locked, and almost none of them named.

    The anonymity floor is the interesting half. 162 of 162 controls carry
    neither alias nor tag, which makes this the document that proves `set_field`
    and the gate errors cannot lean on names being present: every reference to
    one of these has to be by CC ordinal, and every gate error about one has to
    stay readable with the alias and tag segments empty.
    """
    entries, lines, _protection, _text = _ledger_for("dau_acquisition_plan")

    assert len(entries) >= 154
    assert sum(1 for line in lines if "LOCKED (contents)" in line) >= 45
    assert sum(1 for e in entries if not e.alias and not e.tag) >= 150

    empty_with_placeholder = [e for e in entries if e.empty and e.placeholder]
    assert len(empty_with_placeholder) >= 38

    custom = [e for e in empty_with_placeholder if e.placeholder.strip() != STOCK_PLACEHOLDER]
    assert custom, "every placeholder is the stock string; the prose case is untested"


def test_a5_2_modify_into_a_locked_control_is_refused():
    """A3.1 holds on a real document, against a control chosen from the ledger.

    The target is read out of the ledger rather than hardcoded — the point is
    that the gate fires on whatever this document's first locked control happens
    to be after upstream's next revision, not on a string that may not survive
    it.

    It must also be UNIQUE in the projection. This document's first locked
    control reads "ACQUISITION PLAN for", which appears twice, and an ambiguous
    target is refused by the matcher before any gate runs — so the naive version
    of this test passed while proving nothing about locks at all.
    """
    from adeu.models import ModifyText
    from adeu.redline.engine import RedlineEngine

    entries, _lines, _protection, text = _ledger_for("dau_acquisition_plan")
    locked = [e for e in entries if e.locked and e.value and e.value.strip() and text.count(e.value.strip()) == 1]
    assert locked, "no locked control with a unique text value to target"
    target = locked[0]

    engine = RedlineEngine(io.BytesIO(corpus_path("dau_acquisition_plan").read_bytes()), author="A5 Corpus")
    errors = engine.validate_edits(
        [ModifyText(type="modify", target_text=target.value.strip(), new_text="REPLACED BY A TEST")]
    )

    assert errors, f"the edit into locked CC:{target.ordinal} was allowed"
    joined = "\n".join(errors)
    assert f"CC:{target.ordinal}" in joined, joined
    assert "content-locked" in joined.lower(), joined
    assert "ignore_control_locks" in joined, "the refusal must name its override"


# --- A5.3 — wawd_esi_agreement: bound court fields -------------------------


def test_a5_3_bound_court_fields_are_empty_and_bound():
    """Three data-bound caption fields, and the shape `set_field` has to handle.

    All three are EMPTY and BOUND, which is exactly the intersection CC-4's G13
    refuses text edits into and CC-5's dual-write exists to serve.
    """
    entries, lines, _protection, _text = _ledger_for("wawd_esi_agreement")

    by_tag = {e.tag: e for e in entries if e.tag}
    for tag in ("Plaintiff", "Defendant", "Case #"):
        assert tag in by_tag, f"tag {tag!r} missing from {sorted(by_tag)}"
        entry = by_tag[tag]
        assert entry.cls_word == "text"
        assert entry.empty, f"{tag} is not EMPTY"
        assert entry.bound, f"{tag} is not BOUND"

    assert all("BOUND" in line for line in lines)


def test_a5_3_placeholders_render_as_bubbles_not_bare_text():
    """`[Plaintiff]` appears only inside a placeholder bubble.

    The discriminating form of A5.3's last clause. A projection that dropped the
    control and emitted its placeholder run as ordinary text would still contain
    the string — and would read as though the caption were already filled in
    with a literal `[Plaintiff]`. So every occurrence is required to sit inside a
    bubble, which is a claim the bare-substring version cannot make.
    """
    _entries, _lines, _protection, text = _ledger_for("wawd_esi_agreement")

    for token in ("[Plaintiff]", "[Defendant]", "[Case #]"):
        assert token in text, f"{token} vanished from the projection entirely"
        bubbled = text.count(f"{{>>placeholder: {token}<<}}")
        assert bubbled == text.count(token), f"{token} also appears as bare body text"


def test_a5_3_set_field_fills_a_bound_court_field():
    """`set_field` by tag fills a bound control on a real document.

    A4.8's dangling-or-resolving rule decides whether the XML store is updated
    too; either outcome is acceptable here and the assertion deliberately does
    not pick one. What must hold is that the fill is not silently refused — this
    is the operation CC-4's G13 error tells callers to use.
    """
    from adeu.models import SetField
    from adeu.redline.engine import RedlineEngine

    engine = RedlineEngine(io.BytesIO(corpus_path("wawd_esi_agreement").read_bytes()), author="A5 Corpus")
    edit = SetField(type="set_field", field="Case #", value="2:26-cv-01234")

    assert engine.validate_edits([edit]) == [], "set_field into a bound control was refused"

    stats = engine.process_batch([edit])
    assert not stats["failed"], stats.get("skipped_details")
    assert stats["edits_applied"] == 1


# --- A5.4 — on_juries_form1: enforced forms protection ---------------------


def test_a5_4_forms_protection_banner_and_gate():
    """A form with protection and no content controls at all.

    The banner has to say both halves — the protection AND the absence — because
    "no content controls" is what stops an agent concluding the ledger failed to
    load. This document is the corpus's only enforced-protection sample, so it
    is the only place the enforced branch of the banner is exercised for real.
    """
    from adeu.fields import render_banner

    entries, _lines, protection, _text = _ledger_for("on_juries_form1")

    assert entries == [], f"document gained content controls: {len(entries)}"
    assert protection.edit == "forms"
    assert protection.enforced is True

    banner = render_banner(entries, protection)
    assert banner is not None, "a protected document must still get a banner"
    assert "fill-in-forms only (enforced)" in banner
    assert "no content controls" in banner


def test_a5_4_body_edit_is_refused_and_the_override_lets_it_through():
    """G5 on a real enforced-forms document, both directions.

    Asserting only the refusal would leave the override untested, and an
    override that does not work turns a correct gate into a dead end — the exact
    failure the CC-4/CC-5 merge produced elsewhere.

    **A5.4 as frozen says `ignore_document_protection=true` alone makes the edit
    apply, and that is wrong** — the same contradiction A3.5 carried, from the
    same cause. Mikko's spec-gates §1a decision put a SECOND, deliberately
    separate gate behind forms protection: Word records writes to a
    forms-protected document as untracked, so clearing the protection gate still
    leaves the untracked-write gate, which `allow_untracked_writes` clears. Two
    overrides, because they authorise two different things — "I know it is
    protected" is not "I accept an untracked write". A5.4 is corrected
    accordingly; §1a governs.
    """
    from adeu.models import ModifyText
    from adeu.redline.engine import RedlineEngine

    data = corpus_path("on_juries_form1").read_bytes()
    _entries, _lines, _protection, text = _ledger_for("on_juries_form1")

    # Unique, plain, and not an image marker: `![image](docx-image:...)` is the
    # longest line in this document and is not literal text, so the obvious
    # "first long line" pick fails on target resolution before reaching G5.
    target = next(
        (
            line
            for line in (ln.strip() for ln in text.splitlines())
            if 25 <= len(line) <= 90 and "docx-image" not in line and "{#" not in line and text.count(line) == 1
        ),
        None,
    )
    assert target, "no unique plain body line to target"
    edit = ModifyText(type="modify", target_text=target, new_text=target + " (edited)")

    refused = RedlineEngine(io.BytesIO(data), author="A5 Corpus").validate_edits([edit])
    assert refused, "an edit into an enforced fill-in-forms document was allowed"
    joined = "\n".join(refused)
    assert "ignore_document_protection" in joined, joined

    # First override alone: still refused, now by the untracked-write gate, and
    # the new error must name the override that clears IT. A gate that refuses
    # without naming its own escape hatch is where an agent gets stuck.
    half = RedlineEngine(io.BytesIO(data), author="A5 Corpus", ignore_document_protection=True).validate_edits([edit])
    assert half, "the untracked-write gate did not fire"
    assert "allow_untracked_writes" in "\n".join(half), half

    allowed = RedlineEngine(
        io.BytesIO(data),
        author="A5 Corpus",
        ignore_document_protection=True,
        allow_untracked_writes=True,
    ).validate_edits([edit])
    assert allowed == [], f"both overrides did not clear the gates: {allowed}"


# --- A5.5 — ca_talent_recruitment: prompt-as-option ------------------------


def test_a5_5_prompt_is_carried_as_a_real_option():
    """The dropdown's first option is a prompt, and it is a real option.

    `Choose a type.` is Word's prompt text, but it is stored as a genuine
    `w:listItem` — so the ledger must list it rather than filtering it out as
    chrome, and G10 must accept it as a valid value. Dropping it would make the
    ledger disagree with what Word offers the user.
    """
    entries, lines, _protection, _text = _ledger_for("ca_talent_recruitment")

    dropdowns = [e for e in entries if e.cls_word == "dropdown"]
    assert len(dropdowns) == 1, f"expected exactly one dropdown, got {len(dropdowns)}"
    options = list(dropdowns[0].options)
    assert options[:3] == ["Choose a type.", "Internal", "External"], options

    line = next(ln for ln in lines if "dropdown" in ln)
    assert "Choose a type. | Internal | External" in line, line


def test_a5_5_set_field_enforces_dropdown_membership():
    """G10 on a real document: a listed value applies, an unlisted one is refused.

    The refusal has to list the options. A dropdown rejection that does not say
    what IS allowed leaves the agent guessing at a closed set it cannot see.

    The batch is refused as a whole rather than reported as a failed edit,
    because G10 now runs during validation (Mikko, 2026-08-22) — see the
    validate-time test below.
    """
    from adeu.models import SetField
    from adeu.redline.engine import BatchValidationError, RedlineEngine

    data = corpus_path("ca_talent_recruitment").read_bytes()
    entries, _lines, _protection, _text = _ledger_for("ca_talent_recruitment")
    ordinal = next(e.ordinal for e in entries if e.cls_word == "dropdown")

    good = RedlineEngine(io.BytesIO(data), author="A5 Corpus")
    stats = good.process_batch([SetField(type="set_field", field=f"CC:{ordinal}", value="Internal")])
    assert not stats["failed"], stats.get("skipped_details")
    assert stats["edits_applied"] == 1

    bad = RedlineEngine(io.BytesIO(data), author="A5 Corpus")
    with pytest.raises(BatchValidationError) as excinfo:
        bad.process_batch([SetField(type="set_field", field=f"CC:{ordinal}", value="External Hire")])

    joined = "\n".join(excinfo.value.errors)
    assert f"CC:{ordinal}" in joined, joined
    assert "Internal" in joined and "External" in joined, f"the refusal did not list the options: {joined}"


def test_a5_5_dropdown_membership_is_caught_during_validation():
    """G10 refuses before anything is written, like every CC-4 gate.

    This test was originally written the other way round — pinning that G10
    fired only at apply time, which is how CC-5 shipped it. Mikko's 2026-08-22
    ruling moved it into `validate_edits` so the gate contract is uniform: a
    caller who validates a batch and then applies it now learns about a bad
    dropdown value at the same point it learns about locks, protection and
    bound controls, instead of one round trip later with the write already
    begun.

    Kept (inverted) rather than deleted, because the asymmetry is the kind of
    thing that reappears when the next value-bearing gate is added.
    """
    from adeu.models import SetField
    from adeu.redline.engine import RedlineEngine

    entries, _lines, _protection, _text = _ledger_for("ca_talent_recruitment")
    ordinal = next(e.ordinal for e in entries if e.cls_word == "dropdown")

    engine = RedlineEngine(io.BytesIO(corpus_path("ca_talent_recruitment").read_bytes()), author="A5 Corpus")
    errors = engine.validate_edits([SetField(type="set_field", field=f"CC:{ordinal}", value="External Hire")])

    assert errors, "G10 did not fire during validation"
    assert "Internal" in "\n".join(errors), errors


# --- A5.6 — token cost -----------------------------------------------------

_CC_OPEN = re.compile(r"\{#cc:\d+[^}]*\}")
_CC_CLOSE = re.compile(r"\{#/cc:\d+\}")
_CC_BUBBLE = re.compile(r"\{>>[^<]*<<\}")


def test_a5_6_projection_chrome_stays_within_budget():
    """spec-projection §8's budget, on the densest document in the corpus.

    Counts EVERYTHING the projection adds — anchors, flags and placeholder
    bubbles in full, placeholder text included — against a 7% bound. Mikko's
    2026-08-22 ruling, and the reasoning is worth keeping next to the assertion:
    the same document measures 3.69% counting anchors alone, ~4.7% adding bubble
    delimiters, and 6.04% counting bubbles whole. Three defensible conventions,
    two of which cleared the old 5% bound and one of which did not, which made
    "≤5%" a statement about the measurement rather than about read cost.

    So the widest reading wins and the bound moved to fit it. A caller reading a
    control-dense document pays for placeholder text whether or not we classify
    it as content.
    """
    _entries, _lines, _protection, text = _ledger_for("fedramp_ssp_rev4")

    chrome = (
        sum(len(t) for t in _CC_OPEN.findall(text))
        + sum(len(t) for t in _CC_CLOSE.findall(text))
        + sum(len(t) for t in _CC_BUBBLE.findall(text))
    )
    assert chrome > 0, "no projection chrome at all"

    ratio = chrome / len(text)
    assert ratio <= 0.07, f"projection chrome {ratio:.2%} exceeds the 7% budget"


def test_a5_6_the_budget_is_measured_against_a_document_with_anchors():
    """Guards the guard: the ratio is trivially 0% if anchors stop being emitted.

    A5.6 is a ceiling, so it passes hardest when the feature is broken. This
    pins the floor that makes the ceiling meaningful.
    """
    _entries, _lines, _protection, text = _ledger_for("fedramp_ssp_rev4")

    assert len(_CC_OPEN.findall(text)) >= 1_000
    assert len(_CC_CLOSE.findall(text)) >= 1_000
