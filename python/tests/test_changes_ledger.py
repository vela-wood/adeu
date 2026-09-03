import re
from io import BytesIO
from pathlib import Path

from docx import Document

from adeu.ingest import _extract_text_from_doc
from adeu.mcp_components._response_builders import build_changes_response
from adeu.mcp_components.tools.document import read_docx
from adeu.models import ModifyText
from adeu.redline.comments import CommentsManager
from adeu.redline.engine import RedlineEngine
from tests.utils import approx_tokens, get_mock_ctx, run_async, run_cli


def make_sample_docx(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("The quick brown fox jumps over the lazy dog.")
    doc.add_paragraph("This is a second paragraph for testing changes.")

    base_path = tmp_path / "base.docx"
    doc.save(base_path)

    with open(base_path, "rb") as f:
        stream = BytesIO(f.read())

    engine = RedlineEngine(stream, author="Jane Doe")
    engine.process_batch(
        [
            ModifyText(target_text="fox", new_text="cat"),
            ModifyText(target_text="dog", new_text="wolf"),
        ]
    )

    c1 = engine.comments_manager.add_comment("Bob Ross", "Is it really quick?")
    engine.comments_manager.add_comment("Jane Doe", "Yes, very quick!", parent_id=c1)

    out_path = tmp_path / "sample.docx"
    with open(out_path, "wb") as f:
        f.write(engine.save_to_stream().getvalue())
    return out_path


def make_table_docx(tmp_path: Path) -> Path:
    doc = Document()
    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).paragraphs[0].text = "Header 1"
    tbl.cell(0, 1).paragraphs[0].text = "Header 2"
    tbl.cell(1, 0).paragraphs[0].text = "Data Cell"
    tbl.cell(1, 1).paragraphs[0].text = "Value"

    base_path = tmp_path / "tbl_base.docx"
    doc.save(base_path)

    with open(base_path, "rb") as f:
        stream = BytesIO(f.read())
    engine = RedlineEngine(stream, author="Adeu AI")
    engine.process_batch(
        [
            ModifyText(target_text="Data Cell", new_text="Updated Cell"),
        ]
    )
    out_path = tmp_path / "tbl_sample.docx"
    with open(out_path, "wb") as f:
        f.write(engine.save_to_stream().getvalue())
    return out_path


def _get_doc_text(docx_path: Path) -> str:
    doc = Document(docx_path)
    res = _extract_text_from_doc(doc, clean_view=False)
    if isinstance(res, tuple):
        return str(res[0])
    return str(res)


def test_ledger_lists_every_change_id_exactly_once(tmp_path: Path):
    docx_path = make_sample_docx(tmp_path)
    with open(docx_path, "rb") as f:
        engine = RedlineEngine(BytesIO(f.read()))

    expected_change_ids = engine._existing_change_ids()
    assert len(expected_change_ids) > 0

    doc = Document(docx_path)
    text = _get_doc_text(docx_path)
    comments_data = CommentsManager(doc).extract_comments_data()

    res = build_changes_response(text, str(docx_path), comments_data=comments_data)
    content = str(res.content)

    chg_ids = re.findall(r"^Chg:(\d+)", content, re.MULTILINE)
    unique_chg_ids = sorted(list(set(chg_ids)), key=lambda x: (int(x) if x.isdigit() else 0, x))

    assert unique_chg_ids == expected_change_ids
    assert len(chg_ids) == len(set(chg_ids))


def test_ledger_marks_pairs(tmp_path: Path):
    docx_path = make_sample_docx(tmp_path)
    doc = Document(docx_path)
    text = _get_doc_text(docx_path)
    comments_data = CommentsManager(doc).extract_comments_data()

    res = build_changes_response(text, str(docx_path), comments_data=comments_data)
    content = str(res.content)

    assert "(pairs Chg:" in content

    lines = content.splitlines()
    pair_lines = [line for line in lines if "(pairs Chg:" in line]
    assert len(pair_lines) >= 2


def test_ledger_renders_reply_chain(tmp_path: Path):
    docx_path = make_sample_docx(tmp_path)
    doc = Document(docx_path)
    text = _get_doc_text(docx_path)
    comments_data = CommentsManager(doc).extract_comments_data()

    res = build_changes_response(text, str(docx_path), comments_data=comments_data)
    content = str(res.content)

    assert "(reply to Com:" in content


def test_ledger_includes_table_cell_change(tmp_path: Path):
    docx_path = make_table_docx(tmp_path)
    with open(docx_path, "rb") as f:
        engine = RedlineEngine(BytesIO(f.read()))
    expected_ids = engine._existing_change_ids()
    assert len(expected_ids) > 0

    doc = Document(docx_path)
    text = _get_doc_text(docx_path)
    comments_data = CommentsManager(doc).extract_comments_data()

    res = build_changes_response(text, str(docx_path), comments_data=comments_data)
    content = str(res.content)

    for cid in expected_ids:
        assert f"Chg:{cid}" in content


def test_ledger_header_has_totals_authors_and_distribution(tmp_path: Path):
    docx_path = make_sample_docx(tmp_path)
    doc = Document(docx_path)
    text = _get_doc_text(docx_path)
    comments_data = CommentsManager(doc).extract_comments_data()

    res = build_changes_response(text, str(docx_path), comments_data=comments_data)
    content = str(res.content)

    assert "> **Changes ledger** — " in content
    assert "change(s)" in content
    assert "comment(s)" in content
    assert "> Distribution — " in content
    assert "> Authors — " in content
    assert "Jane Doe" in content
    assert "Bob Ross" in content


def test_ledger_token_budget(tmp_path: Path):
    docx_path = make_sample_docx(tmp_path)
    doc = Document(docx_path)
    text = _get_doc_text(docx_path)
    comments_data = CommentsManager(doc).extract_comments_data()

    res = build_changes_response(text, str(docx_path), comments_data=comments_data)
    content = str(res.content)

    tokens = approx_tokens(content)
    assert tokens <= 40 * 18


def test_changes_author_filter(tmp_path: Path):
    docx_path = make_sample_docx(tmp_path)
    doc = Document(docx_path)
    text = _get_doc_text(docx_path)
    comments_data = CommentsManager(doc).extract_comments_data()

    res = build_changes_response(text, str(docx_path), comments_data=comments_data, author_filter="Jane Doe")
    content = str(res.content)

    assert "Jane Doe" in content
    assert "Bob Ross" not in content


def test_changes_page_filter(tmp_path: Path):
    docx_path = make_sample_docx(tmp_path)
    doc = Document(docx_path)
    text = _get_doc_text(docx_path)
    comments_data = CommentsManager(doc).extract_comments_data()

    res = build_changes_response(text, str(docx_path), comments_data=comments_data, page=1)
    content = str(res.content)

    assert "p1" in content


def test_clean_view_with_mode_changes_is_a_usage_error(tmp_path: Path):
    docx_path = make_sample_docx(tmp_path)
    proc = run_cli("extract", str(docx_path), "--mode", "changes", "--clean-view")

    assert proc.returncode == 2
    assert "clean-view" in proc.stderr.lower() or "clean-view" in proc.stdout.lower()


def test_ledger_paginates_above_300_entries(tmp_path: Path):
    lines = []
    for i in range(1, 321):
        lines.append(f"Paragraph {i} {{++ inserted text {i} ++}}{{>>[Chg:{i} insert] Author{i}<<}}")
    text = "\n\n".join(lines)

    res = build_changes_response(text, "synthetic.docx", offset=0, is_cli=True)
    content = str(res.content)

    chg_ids = re.findall(r"\bChg:(\d+)\b", content)
    assert len(chg_ids) == 300
    assert "--changes-offset 300" in content or "offset=300" in content


def test_mcp_mode_changes(tmp_path: Path):
    docx_path = make_sample_docx(tmp_path)
    proc = run_cli("extract", str(docx_path), "--mode", "changes")
    assert proc.returncode == 0
    cli_output = proc.stdout.strip()

    mcp_res = run_async(read_docx(reasoning="test", file_path=str(docx_path), mode="changes", ctx=get_mock_ctx()))
    if isinstance(mcp_res.content, list):
        mcp_content = mcp_res.content[0].text.strip()
    else:
        mcp_content = str(mcp_res.content).strip()

    assert cli_output == mcp_content


def test_read_docx_mode_changes_without_page_returns_all_changes(tmp_path: Path):
    docx_path = make_sample_docx(tmp_path)
    res = run_async(read_docx(reasoning="test", file_path=str(docx_path), mode="changes", ctx=get_mock_ctx()))
    if isinstance(res.content, list):
        content = res.content[0].text
    else:
        content = str(res.content)

    assert "Chg:1" in content
    assert "Chg:2" in content
    assert "Com:1" in content
    assert "> **Changes ledger**" in content


def test_cli_mode_changes_page_range(tmp_path: Path):
    fill = ("Paragraph " + "x" * 100 + "\n\n") * 200
    text = (
        "p1 text {++ ins1 ++}{>>[Chg:1 insert] Jane<<}\n\n"
        + fill
        + "p2 text {++ ins2 ++}{>>[Chg:2 insert] Jane<<}\n\n"
        + fill
        + "p3 text {++ ins3 ++}{>>[Chg:3 insert] Jane<<}"
    )
    res = build_changes_response(text, "multi_page.docx", page="2-3", is_cli=True)
    content = str(res.content)
    assert "Chg:2" in content
    assert "Chg:3" in content
    assert "Chg:1" not in content

    docx_path = make_sample_docx(tmp_path)
    proc = run_cli("extract", str(docx_path), "--mode", "changes", "--page", "1-2")
    assert proc.returncode == 0
    assert "Chg:1" in proc.stdout


def test_large_edit_snippet_not_empty():
    large_del = "A" * 500
    text = f"Some prefix text {{--{large_del}--}}{{>>[Chg:1 delete] Author<<}} suffix text"
    res = build_changes_response(text, "large_edit.docx")
    content = str(res.content)
    chg_lines = [line for line in content.splitlines() if line.startswith("Chg:1")]
    assert len(chg_lines) == 1
    assert '""' not in chg_lines[0]
    assert "AAAAA" in chg_lines[0]


def test_negative_offset_clamped_to_zero():
    lines = [f"Paragraph {i} {{++ text {i} ++}}{{>>[Chg:{i} insert] Jane<<}}" for i in range(1, 321)]
    text = "\n\n".join(lines)
    res = build_changes_response(text, "test.docx", offset=-5, is_cli=True)
    content = str(res.content)
    chg_ids = re.findall(r"\bChg:(\d+)\b", content)
    assert len(chg_ids) == 300
    assert "Chg:1" in content
    assert "Chg:300" in content
    assert "Chg:301" not in content


def test_missing_change_ids_synthesized_from_existing_change_ids():
    text = "Paragraph with {++ inserted text ++}{>>[Chg:1 insert] Author<<}"
    res = build_changes_response(text, "missing_chg.docx", existing_change_ids={"1", "2"})
    content = str(res.content)

    assert "Chg:1" in content
    assert "Chg:2" in content
    chg_2_line = [line for line in content.splitlines() if line.startswith("Chg:2")][0]
    assert "del" in chg_2_line
    assert "Unknown" in chg_2_line
    assert '""' in chg_2_line


def test_multi_change_multi_partner_bubble_ledger():
    text = (
        "Projected text {--AAA--}{++CCC++}{--BBB--}{++DDD++}"
        "{>>[Chg:1 delete] [Chg:2 insert] [Chg:3 delete] [Chg:4 insert] Jane Doe (pairs with Chg:2, Chg:3, Chg:4)<<}"
    )
    res = build_changes_response(text, "multi_partner.docx")
    content = str(res.content)

    lines = {line.split()[0]: line for line in content.splitlines() if line.startswith("Chg:")}

    assert "Chg:1" in lines
    assert "Chg:3" in lines
    assert '"AAA"' in lines["Chg:1"]
    assert '"BBB"' in lines["Chg:3"]
    assert '"CCC"' in lines["Chg:2"]
    assert '"DDD"' in lines["Chg:4"]

    assert "Jane Doe (pairs" not in content
    assert "Jane Doe" in lines["Chg:1"]

    assert "> Authors — Jane Doe" in content

    assert "(pairs Chg:2, Chg:3, Chg:4)" in lines["Chg:1"]
    assert "(pairs Chg:1)" in lines["Chg:2"]
    assert "(pairs Chg:1)" in lines["Chg:3"]
    assert "(pairs Chg:1)" in lines["Chg:4"]


def test_continuation_hint_mcp_format_and_filters():
    lines = [f"Paragraph {i} {{++ text {i} ++}}{{>>[Chg:{i} insert] Jane<<}}" for i in range(1, 321)]
    text = "\n\n".join(lines)

    res = build_changes_response(text, "sample.docx", offset=0, author_filter="Jane", page=1, is_cli=False)
    content = str(res.content)

    assert "changes_offset=300" in content
    assert ", offset=" not in content
    assert 'file_path="sample.docx"' in content
    assert 'mode="changes"' in content
    assert 'changes_author="Jane"' in content
    assert "page=1" in content


def test_continuation_hint_cli_format_and_filters():
    lines = [f"Paragraph {i} {{++ text {i} ++}}{{>>[Chg:{i} insert] Jane<<}}" for i in range(1, 321)]
    text = "\n\n".join(lines)

    res = build_changes_response(text, "sample.docx", offset=0, author_filter="Jane", page=1, is_cli=True)
    content = str(res.content)

    assert "adeu extract sample.docx" in content
    assert "--mode changes" in content
    assert '--changes-author "Jane"' in content
    assert "--page 1" in content
    assert "--changes-offset 300" in content


def test_continuation_hint_mcp_empty_filepath():
    lines = [f"Paragraph {i} {{++ text {i} ++}}{{>>[Chg:{i} insert] Jane<<}}" for i in range(1, 321)]
    text = "\n\n".join(lines)

    res = build_changes_response(text, "", offset=0, is_cli=False)
    content = str(res.content)

    assert 'read_docx(mode="changes", changes_offset=300)' in content
    assert "file_path=" not in content


def test_read_docx_changes_lists_header_footer_change_ids(tmp_path: Path):
    """The changes ledger and the engine's actionable-id list must agree.

    Until issue #114 the engine's revision reads were body-only, so header/
    footer change ids were filtered OUT of the ledger — advertising them
    would have invited accept/reject calls that failed. With revision state
    read across every story part those ids are actionable, so the ledger
    lists them; the invariant under test is unchanged: the ledger's Chg ids
    equal _existing_change_ids() exactly.
    """

    doc = Document()
    doc.add_paragraph("Body text with a change.")
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "Header paragraph"

    base_path = tmp_path / "base.docx"
    doc.save(base_path)

    with open(base_path, "rb") as f:
        stream = BytesIO(f.read())

    engine = RedlineEngine(stream, author="Body Author")
    engine.process_batch([ModifyText(target_text="change", new_text="tracked change")])

    out_doc = engine.doc
    hdr_p = out_doc.sections[0].header.paragraphs[0]._p
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    ins_xml = parse_xml(
        f'<w:ins {nsdecls("w")} w:id="999" w:author="Header Author" '
        'w:date="2026-08-07T00:00:00Z"><w:r><w:t>Header insert</w:t></w:r></w:ins>'
    )
    hdr_p.append(ins_xml)

    docx_path = tmp_path / "sample_header_chg.docx"
    out_doc.save(docx_path)

    with open(docx_path, "rb") as f:
        check_engine = RedlineEngine(BytesIO(f.read()))

    expected_change_ids = check_engine._existing_change_ids()
    assert "999" in expected_change_ids  # header revisions are actionable (issue #114)

    tool_res = run_async(
        read_docx(
            reasoning="test",
            file_path=str(docx_path),
            mode="changes",
            ctx=get_mock_ctx(),
        )
    )
    content = tool_res.content[0].text
    chg_ids = re.findall(r"^Chg:(\d+)", content, re.MULTILINE)
    unique_chg_ids = sorted(list(set(chg_ids)), key=lambda x: (int(x) if x.isdigit() else 0, x))

    assert unique_chg_ids == expected_change_ids


def test_snippet_clamped_to_48_chars():
    long_snip = "A" * 60
    text = f"Some prefix {{++{long_snip}++}}{{>>[Chg:1 insert] Author<<}}"
    res = build_changes_response(text, "test.docx")
    content = str(res.content)
    chg_line = [line for line in content.splitlines() if line.startswith("Chg:1")][0]
    expected_snip = "A" * 45 + "..."
    assert f'"{expected_snip}"' in chg_line


def test_pair_partner_ids_filtered_against_existing_change_ids():
    text = "Projected text {--AAA--}{++CCC++}{>>[Chg:1 delete] [Chg:2 insert] Jane Doe (pairs with Chg:2)<<}"
    res = build_changes_response(text, "pair_test.docx", existing_change_ids={"1"})
    content = str(res.content)
    chg_lines = [line for line in content.splitlines() if line.startswith("Chg:1")]
    assert len(chg_lines) == 1
    assert "(pairs" not in chg_lines[0]

    res_both = build_changes_response(text, "pair_test.docx", existing_change_ids={"1", "2"})
    content_both = str(res_both.content)
    chg_lines_both = [line for line in content_both.splitlines() if line.startswith("Chg:1")]
    assert "(pairs Chg:2)" in chg_lines_both[0]


def test_comment_body_false_positive_chg_id_ignored():
    text = (
        "Some text {++ inserted ++}{>>[Com:1] Bob Ross @ 2026-08-07T17:26:34Z: Please see [Chg:4 insert] for details<<}"
    )
    res = build_changes_response(text, "fp_test.docx", comments_data=None)
    content = str(res.content)
    assert "Com:1" in content
    assert not any(line.startswith("Chg:4") for line in content.splitlines())


def test_chg_tag_does_not_adopt_com_author():
    text = "Some text {-- deleted --}{>>[Chg:1 delete]\n[Com:1] Bob Ross @ 2026-08-07T17:26:34Z: Hello world<<}"
    res = build_changes_response(text, "adopt_test.docx", comments_data=None)
    content = str(res.content)
    chg_lines = [line for line in content.splitlines() if line.startswith("Chg:1")]
    assert len(chg_lines) == 1
    assert "Unknown" in chg_lines[0]
    assert "Bob Ross" not in chg_lines[0]


def test_fallback_comment_parser_iso_timestamp():
    text = "Some text {== format ==}{>>[Com:1] Bob Ross @ 2026-08-07T17:26:34Z: Is it really 17:26:34Z?<<}"
    res = build_changes_response(text, "iso_test.docx", comments_data=None)
    content = str(res.content)
    com_lines = [line for line in content.splitlines() if line.startswith("Com:1")]
    assert len(com_lines) == 1
    assert "Bob Ross" in com_lines[0]
    assert '"Is it really 17:26:34Z?"' in com_lines[0]


def test_ledger_parses_chg_and_com_in_same_bubble_chg_first():
    text = "Some text {-- deleted --}{>>[Chg:1 delete] Jane Doe [Com:1] Bob Ross @ 2026-08-07T17:26:34Z: Hello<<}"
    res = build_changes_response(text, "chg_com_test.docx", comments_data=None)
    content = str(res.content)
    assert any(line.startswith("Chg:1") for line in content.splitlines())
    assert any(line.startswith("Com:1") for line in content.splitlines())


def test_ledger_parses_chg_and_com_in_same_bubble_com_first():
    text = "Some text {-- deleted --}{>>[Com:1] Bob @ 2026-08-07T17:26:34Z: note\n[Chg:1 delete] Jane<<}"
    res = build_changes_response(text, "com_chg_test.docx", comments_data=None)
    content = str(res.content)
    assert any(line.startswith("Com:1") for line in content.splitlines())
    assert any(line.startswith("Chg:1") for line in content.splitlines())


def test_ledger_parses_email_author():
    text = "Some text {== format ==}{>>[Com:1] bob@example.com @ 2026-08-07T17:26:34Z: Hi<<}"
    res = build_changes_response(text, "email_author_test.docx", comments_data=None)
    content = str(res.content)
    com_lines = [line for line in content.splitlines() if line.startswith("Com:1")]
    assert len(com_lines) == 1
    assert "bob@example.com" in com_lines[0]


def test_ledger_parses_author_with_colon():
    text = "Some text {== format ==}{>>[Com:1] Dr: Smith @ 2026-08-07T17:26:34Z: hi<<}"
    res = build_changes_response(text, "colon_author_test.docx", comments_data=None)
    content = str(res.content)
    com_lines = [line for line in content.splitlines() if line.startswith("Com:1")]
    assert len(com_lines) == 1
    assert "Dr: Smith" in com_lines[0]


def test_chg_header_after_com_body_with_lowercase_author():
    text = "Some text {-- deleted --}{>>[Com:1] Bob @ 2026-08-07T17:26:34Z: note\n[Chg:1 delete] jsmith<<}"
    res = build_changes_response(text, "lower_author_test.docx", comments_data=None)
    content = str(res.content)
    com_lines = [line for line in content.splitlines() if line.startswith("Com:1")]
    chg_lines = [line for line in content.splitlines() if line.startswith("Chg:1")]
    assert len(com_lines) == 1
    assert "Bob" in com_lines[0]
    assert '"note"' in com_lines[0]
    assert len(chg_lines) == 1
    assert "del" in chg_lines[0]
    assert "jsmith" in chg_lines[0]
    assert '"deleted"' in chg_lines[0]


def test_chg_header_after_com_body_with_lowercase_name_particle():
    text = "Some text {-- deleted --}{>>[Com:1] Bob @ 2026-08-07T17:26:34Z: note\n[Chg:1 delete] de Vries<<}"
    res = build_changes_response(text, "particle_author_test.docx", comments_data=None)
    content = str(res.content)
    com_lines = [line for line in content.splitlines() if line.startswith("Com:1")]
    chg_lines = [line for line in content.splitlines() if line.startswith("Chg:1")]
    assert len(com_lines) == 1
    assert '"note"' in com_lines[0]
    assert len(chg_lines) == 1
    assert "del" in chg_lines[0]
    assert "de Vries" in chg_lines[0]
    assert '"deleted"' in chg_lines[0]


def test_mid_line_chg_mention_stays_in_comment_body():
    text = (
        "Some text {++ inserted ++}{>>[Com:1] Bob @ 2026-08-07T17:26:34Z: "
        "Please see [Chg:4 insert] Section 5 for details<<}"
    )
    res = build_changes_response(text, "mid_line_mention_test.docx", comments_data=None)
    content = str(res.content)
    com_lines = [line for line in content.splitlines() if line.startswith("Com:1")]
    assert len(com_lines) == 1
    assert "Bob" in com_lines[0]
    assert '"Please see [Chg:4 insert] Section 5 for details"' in com_lines[0]
    assert not any(line.startswith("Chg:") for line in content.splitlines())
