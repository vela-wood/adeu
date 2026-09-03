"""Markdown block markers restyle a paragraph only when they govern the whole
paragraph — and when they do, the marker is consumed rather than leaked.

Two defects are pinned here, both found while closing a dual-engine parity gap.

1. Python leaked the marker. `_maybe_paragraph_replace` admitted HEADINGS ONLY,
   so a whole-paragraph replace like modify("Alpha", "- Beta") fell through to
   the inline path and wrote a literal "- Beta" with no style, while the
   TypeScript engine consumed the "- " and applied ListParagraph + numPr. Same
   edit, same input, two different documents.

2. TypeScript corrupted mid-paragraph edits. Its restyle trigger was a bare
   `target_style !== new_style` with no paragraph-boundary test, so
   modify("Gamma", "- Delta") against "Alpha Gamma" bulleted the entire host
   paragraph AND emitted an EMPTY <w:del> — "Gamma" survived, projecting
   "* Alpha DeltaGamma". Python was already correct here via its bounds test.

The unifying rule, now enforced in both engines: a block marker means "block"
only when the edit spans the whole block. It is the same principle that governs
the line-break gate in track_insert (see test_repro_surgical_word_diff.py).
"""

import io

from docx import Document

from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine


def _apply(paragraph_text: str, target_text: str, new_text: str):
    doc = Document()
    doc.add_paragraph(paragraph_text)
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="Test AI")
    engine.process_batch([ModifyText(type="modify", target_text=target_text, new_text=new_text)])

    out = engine.save_to_stream()
    raw = extract_text_from_stream(out, clean_view=False)
    out.seek(0)
    clean = extract_text_from_stream(out, clean_view=True)
    out.seek(0)
    result_doc = Document(out)
    styles = [p.style.name for p in result_doc.paragraphs]
    has_numpr = any(p._element.xpath(".//w:numPr") for p in result_doc.paragraphs)
    return raw, clean, styles, has_numpr


def test_whole_paragraph_bullet_marker_restyles_and_is_consumed():
    """modify("Alpha" -> "- Beta") over the whole paragraph: bullet style
    applied, "- " consumed as a marker rather than written as text."""
    raw, clean, styles, has_numpr = _apply("Alpha", "Alpha", "- Beta")

    assert "List Paragraph" in styles, f"Expected a bullet paragraph, got styles {styles}"
    assert has_numpr, "Bullet paragraph is missing w:numPr"
    assert clean == "* Beta", f"Unexpected accepted text: {clean!r}"
    # The marker must not survive as literal text.
    assert "{++- Beta++}" not in raw, f"Bullet marker leaked as literal text: {raw!r}"
    assert "{++Beta++}" in raw, f"Expected the marker to be stripped: {raw!r}"


def test_whole_paragraph_asterisk_marker_matches_hyphen_marker():
    """ "* " and "- " are the same bullet marker and must behave identically."""
    hyphen = _apply("Alpha", "Alpha", "- Beta")
    asterisk = _apply("Alpha", "Alpha", "* Beta")
    assert hyphen == asterisk, f"Bullet spellings diverge:\n  '- ': {hyphen}\n  '* ': {asterisk}"


def test_whole_paragraph_heading_marker_still_restyles():
    """Headings were already handled; guard against regressing them while
    widening the gate to bullets."""
    raw, clean, styles, _ = _apply("Alpha", "Alpha", "# Beta")

    assert "Heading 1" in styles, f"Expected a heading, got styles {styles}"
    assert clean == "# Beta", f"Unexpected accepted text: {clean!r}"
    assert "{++# Beta++}" not in raw, f"Heading marker leaked as literal text: {raw!r}"


def test_mid_paragraph_marker_does_not_restyle_and_keeps_deletion_intact():
    """The corruption regression. A fragment replacement must not restyle its
    host paragraph, and the replaced word must actually be deleted."""
    raw, clean, styles, has_numpr = _apply("Alpha Gamma", "Gamma", "- Delta")

    assert clean == "Alpha - Delta", f"Fragment edit corrupted the text: {clean!r}"
    assert "List Paragraph" not in styles, f"Fragment edit restyled its host paragraph: {styles}"
    assert not has_numpr, "Fragment edit applied spurious numbering"
    # The deletion must be real — the TS defect produced an empty <w:del>
    # and left the original word in the document.
    assert "{--Gamma--}" in raw, f"Target was not tracked-deleted: {raw!r}"
    assert "{++- Delta++}" in raw, f"Marker should stay literal in a fragment: {raw!r}"


def test_mid_paragraph_heading_marker_also_stays_inline():
    raw, clean, styles, _ = _apply("Alpha Gamma", "Gamma", "# Delta")

    assert clean == "Alpha # Delta", f"Fragment edit corrupted the text: {clean!r}"
    assert not any(s.startswith("Heading") for s in styles), f"Fragment edit restyled its host paragraph: {styles}"
    assert "{--Gamma--}" in raw, f"Target was not tracked-deleted: {raw!r}"


def test_unmarked_replacement_is_unaffected():
    """Control: no marker, no restyle, ordinary tracked modification."""
    raw, clean, styles, has_numpr = _apply("Alpha Gamma", "Gamma", "Delta")

    assert clean == "Alpha Delta"
    assert styles == ["Normal"], f"Unexpected restyle: {styles}"
    assert not has_numpr
    assert "{--Gamma--}{++Delta++}" in raw
