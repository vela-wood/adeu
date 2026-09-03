import inspect
import io
import logging
import re
import sys
from pathlib import Path

import pytest
import structlog
from docx import Document

from adeu.utils.console import dynamic_stderr

# This fork deliberately excludes the Python MCP server and live-Word adapter
# from its distribution. Keep upstream's server tests in the tree for easier
# future merges, but do not collect any module that imports that surface at
# import time (fastmcp is not installed here).
_EXCLUDED_SURFACE_IMPORT_RE = re.compile(
    r"^\s*(?:from|import)\s+(?:fastmcp|mcp\b|adeu\.server|adeu\.mcp_components\.(?:tools|resources))"
    r"|^from adeu\.cli import .*\b(?:handle_init|_get_claude_config_path)\b",
    re.MULTILINE,
)


def _imports_excluded_surface(path: Path) -> bool:
    try:
        return bool(_EXCLUDED_SURFACE_IMPORT_RE.search(path.read_text(encoding="utf-8")))
    except OSError:
        return False


collect_ignore = [p.name for p in Path(__file__).parent.glob("test_*.py") if _imports_excluded_surface(p)]

# Unconfigured structlog prints DEBUG lines to STDOUT, so any test that calls
# engine/ingest helpers in its setup pollutes captured stdout (a `--json` CLI
# test then fails json.loads on "2026-…" log lines). The CLI and server both
# bind structlog to stderr at entry; tests only got that binding if some
# earlier test happened to trigger it — an order dependence pytest-xdist's
# fresh workers expose. Configure it here, once per worker, exactly the way
# the CLI does (WARNING level, dynamic stderr proxy so capsys replacement and
# teardown are honored — never pin the sys.stderr object itself).
structlog.configure(
    wrapper_class=structlog.make_filtering_bound_logger(logging.WARNING),
    logger_factory=structlog.PrintLoggerFactory(file=dynamic_stderr),  # type: ignore[arg-type]
)

try:
    from hypothesis import HealthCheck
    from hypothesis import settings as _hyp_settings

    # Property-test profiles (tests/test_property_invariants.py). Registered
    # here so `--hypothesis-profile=hunt` resolves at pytest configure time.
    _hyp_settings.register_profile(
        "default", deadline=None, max_examples=25, suppress_health_check=[HealthCheck.too_slow]
    )
    _hyp_settings.register_profile(
        "hunt", deadline=None, max_examples=300, suppress_health_check=[HealthCheck.too_slow]
    )
    _hyp_settings.load_profile("default")
except ImportError:
    pass


@pytest.hookimpl(tryfirst=True)
def pytest_collection_modifyitems(config, items):
    """Serialize live-Word tests under pytest-xdist.

    tryfirst matters: in each xdist worker, WorkerInteractor's own
    pytest_collection_modifyitems consumes the xdist_group marker (it rewrites
    the nodeid to "…@group" for the loadgroup scheduler). Our marker must be
    attached before that hook runs, or the grouping silently does nothing.

    Tests that drive the real Word COM instance (the `active_word_app` and
    `word_app` fixtures, plus everything in test_live_word*.py) all bind to the
    single active Word application — two xdist workers doing that concurrently
    corrupt each other's document state. `xdist_group` + `--dist loadgroup`
    (set in pyproject addopts) pins them to ONE worker, where they run
    sequentially; every other test still distributes freely.
    """
    com_fixtures = {"active_word_app", "word_app"}
    for item in items:
        try:
            source = inspect.getsource(item.obj)
        except (OSError, TypeError):
            source = ""
        unsupported_python_server_surface = any(
            marker in source
            for marker in (
                "fastmcp",
                "adeu.server",
                "mcp_components.tools",
                "handle_init",
                "_get_claude_config_path",
                'run_cli(["init"',
                "--live",
            )
        )
        if unsupported_python_server_surface and item.name != "test_extract_never_imports_fastmcp":
            item.add_marker(pytest.mark.skip(reason="Python MCP/live-Word surface is excluded from the slim fork"))

        is_live_word = bool(com_fixtures & set(getattr(item, "fixturenames", ())))
        if not is_live_word:
            basename = item.path.name if getattr(item, "path", None) else ""
            is_live_word = basename.startswith("test_live_word")
        if is_live_word:
            item.add_marker(pytest.mark.xdist_group("live_word"))


@pytest.fixture(scope="session", autouse=True)
def _isolate_windows_appdata(tmp_path_factory):
    """On Windows, `adeu init` resolves the Claude Desktop config via %APPDATA%.
    A test that runs init without patching _get_claude_config_path would rewrite
    the developer's real claude_desktop_config.json (this happened 2026-07-20:
    two QA-repro tests injected fake uvx entries into a live config). Pointing
    APPDATA at a throwaway directory for the whole session makes that class of
    accident impossible; tests that assert on the config still patch the path
    getter explicitly."""
    if sys.platform != "win32":
        yield
        return
    mp = pytest.MonkeyPatch()
    mp.setenv("APPDATA", str(tmp_path_factory.mktemp("appdata")))
    yield
    mp.undo()


@pytest.fixture
def simple_docx_stream():
    """Returns a BytesIO stream containing a simple DOCX."""
    doc = Document()
    doc.add_heading("Contract Agreement", 0)
    doc.add_paragraph("This is a simple contract.")
    doc.add_paragraph("The party of the first part shall be known as the Seller.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


# Only define COM fixtures on Windows
if sys.platform == "win32":
    import pythoncom
    import win32com.client

    def _await_active_document(doc, attempts: int = 20, delay: float = 0.05) -> None:
        """Block until the tools would resolve `doc`, or fail saying what they see.

        The check deliberately goes through `GetActiveObject`, not through the
        `Dispatch` handle the fixture already holds. Those are not guaranteed to
        be the same object: `Dispatch` reuses a running instance *or starts one*,
        while `GetActiveObject` returns whatever is registered in the COM Running
        Object Table. With two `WINWORD.EXE` processes alive they resolve to
        different applications, and a fixture that verified its own handle would
        report everything fine while the tool under test read another Word's
        document entirely. Asserting on the production lookup path is the whole
        point (CC-13).

        The retry loop exists because `Document.Activate()` is not synchronous —
        Word acknowledges it before the window manager has finished, so an
        immediate read can still return the outgoing document.
        """
        import time

        seen = None
        for _ in range(attempts):
            try:
                app = win32com.client.GetActiveObject("Word.Application")
                seen = app.ActiveDocument.FullName
                if seen == doc.FullName:
                    return
                doc.Activate()
            except Exception as exc:  # pragma: no cover - diagnostic path
                seen = f"<{type(exc).__name__}: {exc}>"
            time.sleep(delay)

        try:
            app = win32com.client.GetActiveObject("Word.Application")
            open_docs = [d.FullName for d in app.Documents]
        except Exception:  # pragma: no cover - diagnostic path
            open_docs = ["<unavailable>"]

        pytest.fail(
            "The live-Word tools would not resolve this fixture's document: "
            f"GetActiveObject().ActiveDocument is {seen!r}, expected {doc.FullName!r}. "
            f"Open documents: {open_docs!r}. Word processes may have multiplied, or "
            "another live-Word test leaked a document into the shared instance (CC-13)."
        )

    @pytest.fixture
    def active_word_app():
        """
        Creates an ephemeral, visible MS Word instance with a fresh document.
        Ensures it is torn down properly after the test.

        **Three things here are load-bearing, not tidiness** (all CC-13). The
        tools under test bind through `GetActiveObject` and read
        `app.ActiveDocument`, so every assertion in these suites is really an
        assertion about whichever document Word happens to consider active.

        1. `doc.Activate()`. `Documents.Add()` usually makes the new document
           active and `app.Activate()` raises the *application* — neither
           guarantees the *document*.
        2. `_await_active_document`, which confirms the claim through
           `GetActiveObject` (the production lookup) rather than through this
           fixture's own `Dispatch` handle, and waits, because activation is
           asynchronous.
        3. Closing every document that appeared during the test. The tools open
           documents by path and never close them, so they pile up in the shared
           instance and each one is a candidate for `ActiveDocument` in later
           tests.

        Without these the failure is spectacularly confusing from the outside:
        the assertion reports text belonging to a different test file entirely
        ("assert '{++Title++}' in 'Initial {==manuscript==}...'"), the set of
        failures changes on every run, and each suite passes in isolation.
        """
        pythoncom.CoInitialize()

        app = None
        try:
            # Dispatch starts a new background instance if one doesn't exist.
            # GetActiveObject will then be able to hook into it in the tool.
            app = win32com.client.Dispatch("Word.Application")
            app.Visible = True  # Needs to be visible/active for GetActiveObject sometimes

            doc = app.Documents.Add()

            # Bring to front so GetActiveObject definitely binds to this instance
            app.Activate()
            # ...and make OUR document the one GetActiveObject will resolve to.
            doc.Activate()

            # Seed initial content
            doc.Range(0, 0).Text = "Hello world! This is a live testing document.\n"

            _await_active_document(doc)

            yield app, doc

        except Exception as e:
            pytest.skip(f"Could not initialize Word COM for testing: {e}")

        finally:
            if app:
                # Only this fixture's own document. Closing every document that
                # appeared during the test was tried and REVERTED: the live-Word
                # tools hand back Ranges into documents they opened, and reaping
                # those turned the ambient-activation bug into
                # "(-2147417848) The object invoked has disconnected from its
                # clients" plus "Object has been deleted" — a worse failure,
                # because it looks like a COM fault rather than a test-isolation
                # one. Document accumulation is real and still unfixed; see CC-13,
                # which now has that dead end recorded so nobody re-walks it.
                try:
                    doc.Close(0)  # 0 = wdDoNotSaveChanges
                except Exception:
                    pass
                # We intentionally omit app.Quit() and pythoncom.CoUninitialize()
                # to avoid Windows Access Violations (0x800706be) when Pytest holds COM locals.

    @pytest.fixture(scope="session")
    def word_app():
        """
        A Word application for INSPECTING packages the engine wrote to disk
        (tests/word_com.py). Distinct from `active_word_app`, which drives an
        open document through the live-Word tools: this one opens saved files
        read-only and is the oracle for ids Word reinterprets on load —
        w14:paraId threading, w16cid:durableId anchoring
        (BUG_paraId_signed_int32_thread_collapse.md).

        Session-scoped because starting Word costs seconds and the tests only
        read. Like `active_word_app` it deliberately omits app.Quit() and
        pythoncom.CoUninitialize(): tearing COM down while pytest still holds
        proxies raises 0x800706be. Not quitting is also what makes it safe to
        attach to a developer's already-running Word.

        `Visible` is deliberately NOT touched. Dispatch attaches to the running
        Word instance, so this is the SAME application object `active_word_app`
        drives and the developer has open; forcing it hidden changes what
        `GetActiveObject` binds to in the live-Word tools. Reading document
        properties does not need visibility. `DisplayAlerts` does have to go:
        without it, opening a package Word considers corrupted — which is one
        of the outcomes under test — blocks on a modal dialog forever.
        """
        pythoncom.CoInitialize()
        try:
            app = win32com.client.Dispatch("Word.Application")
            app.DisplayAlerts = 0
        except Exception as e:
            pytest.skip(f"Could not initialize Word COM for testing: {e}")
            return
        yield app

else:  # pragma: no cover - non-Windows CI

    @pytest.fixture(scope="session")
    def word_app():
        pytest.skip("Live Word COM tests require Windows")
