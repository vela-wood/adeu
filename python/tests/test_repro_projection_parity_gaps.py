"""
Parity twin of node/packages/core/src/repro_projection_parity_gaps.test.ts.

Two divergences kept the engines' projections from being byte-identical on real
documents. Both were found by running the FedRAMP SSP Moderate rev4 template
(shared/corpus) through both engines and diffing — neither is reachable from
the synthetic fixtures either suite used before.

1. EMPHASIS COALESCING. Adjacent runs with identical formatting must project as
   ONE marker span. Python already did this correctly (`build_paragraph_text`
   ignores trailing whitespace before testing for the closing marker); Node
   tested the literal tail, so hoisted boundary whitespace defeated it and it
   emitted ``**A** **B**`` where Python emitted ``**A B**``.

2. HEADER/FOOTER ENUMERATION. Python walks `w:sectPr` references, honouring
   Link-to-Previous, ``w:titlePg`` and ``w:evenAndOddHeaders``, so it projects
   what Word renders. Node listed every header/footer part in the package.

Python was the correct side of both, so these tests are characterization tests:
they pin the behaviour Node was brought into line with, so a future change to
the Python side cannot silently re-open the gap.
"""

import io

import pytest
from docx import Document  # noqa: F401 — Document() with no args builds an empty doc; load_document opens bytes
from docx.oxml import parse_xml
from docx.oxml.ns import qn

from adeu.ingest import extract_text_from_stream
from adeu.redline.mapper import DocumentMapper
from adeu.utils.opc import load_document

NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'

BOLD = "<w:rPr><w:b/></w:rPr>"
ITALIC = "<w:rPr><w:i/></w:rPr>"


def _paragraph_of(runs: list[tuple[str, str]]) -> str:
    body = "".join(f'<w:r>{rpr}<w:t xml:space="preserve">{t}</w:t></w:r>' for t, rpr in runs)
    return f"<w:p {NS}>{body}</w:p>"


def _docx(*body_xml: str) -> bytes:
    doc = Document()
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    for xml in body_xml:
        el = parse_xml(xml)
        if sect_pr is not None:
            sect_pr.addprevious(el)
        else:
            body.append(el)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _project(data: bytes) -> str:
    return extract_text_from_stream(io.BytesIO(data), clean_view=True, include_appendix=False)


# ---------------------------------------------------------------------------
# 1. Emphasis coalescing
# ---------------------------------------------------------------------------
def test_bold_runs_separated_by_a_hoisted_space_merge():
    """The corpus shape: '**Name of Organization** **CSP Name...**' was wrong."""
    text = _project(_docx(_paragraph_of([("Name of Organization", BOLD), (" CSP Name System Connects To", BOLD)])))
    assert text.strip() == "**Name of Organization CSP Name System Connects To**"


def test_three_italic_runs_merge_into_one_span():
    """python: '_Version #.#,  Date_'; node emitted '_Version_ _#.#,_  _Date_'."""
    text = _project(_docx(_paragraph_of([("Version", ITALIC), (" #.#,", ITALIC), ("  Date", ITALIC)])))
    assert text.strip() == "_Version #.#,  Date_"


def test_runs_with_no_whitespace_between_them_merge():
    # A fully-bold paragraph also trips heading detection, hence the "## ".
    text = _project(_docx(_paragraph_of([("A", BOLD), ("B", BOLD)])))
    assert "**AB**" in text
    assert "**A****B**" not in text


def test_differing_formatting_is_not_merged():
    text = _project(_docx(_paragraph_of([("bold", BOLD), (" and ", ""), ("italic", ITALIC)])))
    assert text.strip() == "**bold** and _italic_"


def test_markers_stay_balanced_when_a_whitespace_only_same_style_run_follows():
    """Eliding the closer without a matching opener lost marker balance."""
    text = _project(_docx(_paragraph_of([("March 2012", BOLD), ("  ", BOLD)])))
    assert text.count("**") % 2 == 0, f"unbalanced emphasis markers: {text!r}"


def test_ingest_and_mapper_agree_on_merged_emphasis():
    data = _docx(_paragraph_of([("Name of Organization", BOLD), (" CSP Name", BOLD)]))
    assert DocumentMapper(load_document(io.BytesIO(data)), clean_view=True).full_text == _project(data)


# ---------------------------------------------------------------------------
# 2. Header / footer enumeration
# ---------------------------------------------------------------------------
def _docx_with_header_footer(header: str | None, footer: str | None) -> bytes:
    doc = Document()
    sec = doc.sections[0]
    if header is not None:
        sec.header.paragraphs[0].text = header
    if footer is not None:
        sec.footer.paragraphs[0].text = footer
    doc.add_paragraph("Body text.")
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def test_referenced_header_and_footer_project():
    text = _project(_docx_with_header_footer("HEAD", "FOOT"))
    assert "HEAD" in text
    assert "Body text." in text
    assert "FOOT" in text


def test_first_page_header_is_ignored_without_title_pg():
    """python-docx creates the part, but Word renders it only with w:titlePg."""
    doc = Document()
    sec = doc.sections[0]
    sec.header.paragraphs[0].text = "DEFAULT HEAD"
    sec.different_first_page_header_footer = False
    sec.first_page_header.paragraphs[0].text = "FIRST HEAD"
    doc.add_paragraph("Body text.")
    out = io.BytesIO()
    doc.save(out)

    text = _project(out.getvalue())
    assert "DEFAULT HEAD" in text
    assert "FIRST HEAD" not in text


def test_first_page_header_projects_once_the_section_opts_in():
    doc = Document()
    sec = doc.sections[0]
    sec.header.paragraphs[0].text = "DEFAULT HEAD"
    sec.different_first_page_header_footer = True
    sec.first_page_header.paragraphs[0].text = "FIRST HEAD"
    doc.add_paragraph("Body text.")
    out = io.BytesIO()
    doc.save(out)

    text = _project(out.getvalue())
    assert "DEFAULT HEAD" in text
    assert "FIRST HEAD" in text


@pytest.mark.parametrize("opt_in", [False, True])
def test_even_page_header_follows_the_document_toggle(opt_in):
    doc = Document()
    sec = doc.sections[0]
    sec.header.paragraphs[0].text = "DEFAULT HEAD"
    sec.even_page_header.paragraphs[0].text = "EVEN HEAD"
    doc.settings.odd_and_even_pages_header_footer = opt_in
    doc.add_paragraph("Body text.")
    out = io.BytesIO()
    doc.save(out)

    text = _project(out.getvalue())
    assert "DEFAULT HEAD" in text
    assert ("EVEN HEAD" in text) is opt_in


# ---------------------------------------------------------------------------
# 3. Run-level elements that used to fall through silently
# ---------------------------------------------------------------------------
@pytest.mark.parametrize(
    "name,run_xml,expected",
    [
        # A real hyphen glyph: dropping it merged the words either side.
        ("noBreakHyphen", "<w:r><w:t>e</w:t><w:noBreakHyphen/><w:t>mail</w:t></w:r>", "e-mail"),
        # Absolute-position tab: separates content, like w:tab.
        (
            "ptab",
            '<w:r><w:t>A</w:t><w:ptab w:relativeTo="margin" w:alignment="left" w:leader="none"/><w:t>B</w:t></w:r>',
            "A B",
        ),
        # Optional break hint. Word shows it only when the line actually
        # breaks, so projecting nothing is CORRECT — pinned so nobody
        # "fixes" it into a visible character.
        ("softHyphen", "<w:r><w:t>co</w:t><w:softHyphen/><w:t>operate</w:t></w:r>", "cooperate"),
    ],
)
def test_run_level_elements_project_their_glyph(name, run_xml, expected):
    assert _project(_docx(f"<w:p {NS}>{run_xml}</w:p>")).strip() == expected


def test_symbol_runs_are_still_dropped_deliberately():
    """
    w:sym is NOT projected. Symbol fonts map glyphs into the Unicode
    private-use area (Wingdings F0FE is a checked box), so the code point
    alone does not identify the character and guessing corrupts text. CC-1
    owns checkbox projection and needs a font-aware decision; this pins the
    status quo so the loss is a recorded choice rather than an oversight.
    """
    text = _project(_docx(f'<w:p {NS}><w:r><w:sym w:font="Wingdings" w:char="F0FE"/></w:r></w:p>'))
    assert text.strip() == ""


# ---------------------------------------------------------------------------
# 4. The regression guard for all of the above: python and node must project
#    real documents identically, character for character.
#
#    This is the assertion A5.1 specifies ("Engines: python + node — identical
#    counts"). Character counts are recorded here as a python-side tripwire:
#    the true cross-engine comparison needs both runtimes, so it lives in the
#    parity harness, but any python-side drift moves these numbers and fails
#    here first.
#
#    Every value was verified byte-identical against the node engine on
#    2026-08-21 across 4 fixtures x 2 views and 4 corpus documents x 2 views
#    (16/16), with zero DocumentMapper drift.
# ---------------------------------------------------------------------------
CORPUS_PROJECTION_SIZES = {
    # key: (raw_view_chars, clean_view_chars)
    # CC-1c moved both views by +7,762 = 3,881 checkboxes x 2, the width a
    # `w14:checkbox` gains going from a one-character ballot glyph to the
    # three-character `[ ]` token. No emphasis markers are involved here,
    # unlike odot_uic_drywell below.
    # CC-17 then moved both views by +365: two of this document's nine `sectPr`
    # sit at `body/sdt/sdtContent/p/pPr`, so the sections were invisible and the
    # headers they reference were never walked. header2's running-header content
    # is the +365; header1 is empty, and the two footers likewise. Nothing is
    # projected twice — the growth is one part that was missing, not a duplicate.
    "fedramp_ssp_rev4": (621_040, 521_089),
    "dau_acquisition_plan": (19_611, 17_254),
    "wawd_esi_agreement": (15_978, 15_891),
    "on_juries_form1": (5_505, 3_199),
    "ca_talent_recruitment": (5_613, 5_109),
    # A .dotx. Absent from this table until CC-11, because python could not open
    # one at all and there was nothing to pin against. Both views are identical
    # because the template carries no tracked changes; the count rose from 7,221
    # when CC-1b started projecting content-control anchors.
    #
    # CC-1c then moved it DOWN, 7,449 -> 7,435, which looks wrong for a change
    # that widens glyphs into tokens and is not. Attributed exactly: of this
    # document's 21 ballot glyphs, 19 sit in controls and 13 of those arrived
    # wrapped in emphasis markers, projecting as `**<glyph>**`. The mark is
    # chrome, so it now carries no markers: -52 for the 13 x 4 dropped marker
    # characters, +38 for 19 x 2 of token width, net -14. The two surviving
    # glyphs are bare prose outside any control and are deliberately untouched.
    "odot_uic_drywell": (7_435, 7_435),
}


@pytest.mark.parametrize("key", sorted(CORPUS_PROJECTION_SIZES))
@pytest.mark.parametrize("clean_view", [False, True])
def test_corpus_projection_size_is_pinned_to_the_node_engine(key, clean_view):
    from tests.utils import corpus_path

    path = corpus_path(key)  # skips cleanly when the document is absent
    data = path.read_bytes()

    text = extract_text_from_stream(io.BytesIO(data), clean_view=clean_view, include_appendix=False)
    expected = CORPUS_PROJECTION_SIZES[key][1 if clean_view else 0]
    assert len(text) == expected, (
        f"{key} {'clean' if clean_view else 'raw'} view projects {len(text):,} chars, "
        f"expected {expected:,}. If this change is intentional, re-run the parity "
        f"harness against the node engine and update BOTH engines' pinned values — "
        f"a python-only change here re-opens the divergence."
    )

    # No markup may reach the character stream (CC-10).
    assert "<w:" not in text


@pytest.mark.parametrize("key", sorted(CORPUS_PROJECTION_SIZES))
def test_corpus_ingest_and_mapper_agree(key):
    """The Virtual Text contract, on real documents rather than fixtures."""
    from tests.utils import corpus_path

    data = corpus_path(key).read_bytes()
    for clean_view in (False, True):
        projected = extract_text_from_stream(io.BytesIO(data), clean_view=clean_view, include_appendix=False)
        mapped = DocumentMapper(load_document(io.BytesIO(data)), clean_view=clean_view).full_text
        assert mapped == projected, f"{key}: mapper drifted from ingest (clean={clean_view})"


# ---------------------------------------------------------------------------
# 5 — shapes the corpus cannot reach
#
# The corpus is published documents: no tracked changes, so a whole class of
# clean-view behaviour never gets exercised by the size pins above. Both cases
# below were already correct in python and both were WRONG in node (node's
# ingest and mapper were consistently wrong together, so they agreed with each
# other and only cross-engine comparison caught them). Pinned here so the
# python side stays the oracle.
# ---------------------------------------------------------------------------

_DEL_ATTRS = 'w:id="900" w:author="A" w:date="2026-01-01T00:00:00Z"'


def _doc_with_deleted_paragraph_mark() -> bytes:
    """Alpha / a paragraph whose mark AND content are tracked deletions / Beta."""
    d = Document()
    d.add_paragraph("Alpha")
    p = d.add_paragraph()
    p._element.append(parse_xml(f"<w:del {NS} {_DEL_ATTRS}><w:r><w:delText>gone</w:delText></w:r></w:del>"))
    p._element.get_or_add_pPr().append(parse_xml(f"<w:rPr {NS}><w:del {_DEL_ATTRS}/></w:rPr>"))
    d.add_paragraph("Beta")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_clean_view_drops_paragraph_whose_mark_is_deleted():
    """Accepting a paragraph-mark deletion merges the paragraph away.

    When nothing visible survives inside it the accepted view must render no
    container at all — not an empty one. An empty container costs a whole
    "\\n\\n" block separator, so the bug shows up as a doubled blank line.
    """
    data = _doc_with_deleted_paragraph_mark()

    clean = extract_text_from_stream(io.BytesIO(data), clean_view=True, include_appendix=False)
    assert clean == "Alpha\n\nBeta", f"clean view kept an empty container for a deleted paragraph mark: {clean!r}"

    # The raw view still shows the deletion — this is a clean-view-only skip.
    raw = extract_text_from_stream(io.BytesIO(data), clean_view=False, include_appendix=False)
    assert "{--gone--}" in raw

    for clean_view in (False, True):
        mapped = DocumentMapper(Document(io.BytesIO(data)), clean_view=clean_view).full_text
        projected = extract_text_from_stream(io.BytesIO(data), clean_view=clean_view, include_appendix=False)
        assert mapped == projected, f"mapper drifted from ingest (clean={clean_view})"


def test_empty_styled_run_contributes_no_style_markers():
    """A styled run with no projected text emits nothing — not even markers.

    A bold run whose only child is a drawing or a footnote reference would
    otherwise leave a dangling "****" pair that the reader never emits, since
    apply_formatting_to_segments("") is "".
    """
    d = Document()
    p = d.add_paragraph()
    p._element.append(parse_xml(f'<w:r {NS}><w:rPr><w:b/></w:rPr><w:footnoteReference w:id="2"/></w:r>'))
    run = p.add_run("Visible")
    run.bold = True
    buf = io.BytesIO()
    d.save(buf)
    data = buf.getvalue()

    text = extract_text_from_stream(io.BytesIO(data), clean_view=False, include_appendix=False)
    assert "****" not in text, f"empty styled run emitted a dangling marker pair: {text!r}"

    mapped = DocumentMapper(Document(io.BytesIO(data)), clean_view=False).full_text
    assert mapped == text, f"mapper drifted from ingest: {mapped!r} != {text!r}"
