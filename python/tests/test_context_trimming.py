import io

from docx import Document

from adeu.diff import trim_common_context
from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine


def test_trim_logic_basic():
    """Prefix and Suffix exist."""
    t = "Context A Context"
    n = "Context B Context"
    p, s = trim_common_context(t, n)
    assert p == 8  # "Context "
    assert s == 8  # " Context"
    # Remainder: "A", "B"


def test_trim_logic_prefix_only():
    t = "Hello World"
    n = "Hello User"
    p, s = trim_common_context(t, n)
    assert p == 6  # "Hello "
    assert s == 0


def test_trim_logic_suffix_only():
    t = "Old Item"
    n = "New Item"
    p, s = trim_common_context(t, n)
    assert p == 0
    assert s == 5  # " Item"


def test_trim_logic_morph_to_insert():
    t = "Prefix"
    n = "Prefix Added"
    p, s = trim_common_context(t, n)
    assert p == 6
    assert s == 0


def test_end_to_end_context_cleanup():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Start ")
    p.add_run("Middle")
    p.add_run(" End")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = ModifyText(target_text="Start Middle End", new_text="Start Center End")

    engine = RedlineEngine(stream)
    engine.apply_edits([edit])

    result_stream = engine.save_to_stream()
    doc = Document(result_stream)
    xml = doc.element.xml

    assert "w:del" not in xml.split("Start ")[0][-20:]
    assert "<w:delText>Middle</w:delText>" in xml
    assert "<w:t>Center</w:t>" in xml
    assert "<w:t>Center End</w:t>" not in xml


def test_auto_strip_insertion_duplication():
    doc = Document()
    doc.add_paragraph("Liability Cap.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    edit = ModifyText(target_text="Liability Cap.", new_text="Liability Cap. SLA Clause.")

    engine = RedlineEngine(stream)
    engine.apply_edits([edit])

    result_stream = engine.save_to_stream()
    doc = Document(result_stream)
    xml = doc.element.xml

    assert xml.count("Liability Cap.") == 1
    assert "SLA Clause." in xml


def test_trailing_whitespace_only_removal_is_a_real_deletion():
    """
    `target="A.\\n\\n"`, `new="A."` is the paragraph-MERGE shape that
    `make_edits_self_contained` emits (widening a bare "\\n\\n" deletion with
    its left neighbour word). The apply path's rstrip "smart fallback" for
    trailing-space omissions used to swallow it: `new` starts with
    `target.rstrip()`, so the edit resolved to an INSERTION of "" — a no-op
    that the batch still reported as applied, silently leaving the paragraph
    break in place. It must resolve to a DELETION of the paragraph mark and
    the paragraphs must merge on accept.
    """
    doc = Document()
    doc.add_paragraph("A.")
    doc.add_paragraph("B.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="Tester")
    stats = engine.process_batch([ModifyText(target_text="A.\n\n", new_text="A.")])
    assert stats["edits_skipped"] == 0

    engine.accept_all_revisions(remove_comments=True)
    merged = Document(engine.save_to_stream())
    assert [p.text for p in merged.paragraphs] == ["A.B."]


def test_trailing_whitespace_omission_fallback_still_appends():
    """The counterpart the fallback exists for: the caller dropped the target's
    trailing space but genuinely EXTENDS it. That stays a pure insertion."""
    doc = Document()
    doc.add_paragraph("Alpha beta gamma.")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="Tester")
    stats = engine.process_batch([ModifyText(target_text="Alpha beta ", new_text="Alpha beta delta ")])
    assert stats["edits_skipped"] == 0

    engine.accept_all_revisions(remove_comments=True)
    merged = Document(engine.save_to_stream())
    assert [p.text for p in merged.paragraphs] == ["Alpha beta delta gamma."]


def test_trim_logic_full_suffix_overlap_crash_repro():
    """
    Regression test for IndexError when suffix consumes the entire target.
    Target: "Agreement"
    New: "New Agreement"
    Suffix match is "Agreement" (len 9). Accessing target[-(9+1)] is out of bounds.
    """
    target = "Agreement"
    new_val = "New Agreement"

    # This should not raise IndexError
    p, s = trim_common_context(target, new_val)

    # Expectation:
    # Prefix: 0
    # Suffix: 9 ("Agreement")
    assert p == 0
    assert s == 9
