import io
import json
import sys
from unittest.mock import patch

import pytest
from docx import Document

from adeu.cli import main
from adeu.mcp_components.tools.document import process_document_batch
from adeu.models import AcceptChange, ModifyText, RejectChange
from adeu.redline.engine import BatchValidationError, RedlineEngine
from tests.utils import approx_tokens, extract_content, get_mock_ctx, run_async


def _create_sample_docx(tmp_path, filename="sample.docx"):
    path = tmp_path / filename
    doc = Document()
    doc.add_paragraph("The quick brown fox jumps over the lazy dog.")
    doc.save(path)
    return path


def test_cli_partial_lands_valid_edits_and_leads_with_failures(tmp_path, capsys):
    doc_path = _create_sample_docx(tmp_path)
    out_path = tmp_path / "out.docx"
    changes_path = tmp_path / "changes.json"

    changes_path.write_text(
        json.dumps(
            [
                {"type": "modify", "target_text": "quick brown", "new_text": "fast blue"},
                {"type": "modify", "target_text": "nonexistent text", "new_text": "replacement"},
            ]
        ),
        encoding="utf-8",
    )

    test_args = ["adeu", "apply", str(doc_path), str(changes_path), "-o", str(out_path), "--partial"]
    with patch.object(sys, "argv", test_args):
        try:
            main()
        except SystemExit as e:
            assert e.code == 0 or e.code is None

    captured = capsys.readouterr()
    err_output = captured.err
    assert "PARTIAL: applied 1 of 2" in err_output

    assert out_path.exists()
    out_doc = Document(out_path)
    from adeu.ingest import _extract_text_from_doc

    text = _extract_text_from_doc(out_doc, clean_view=True)
    assert "fast blue" in text


def test_cli_partial_all_edits_fail_exits_1(tmp_path, capsys):
    doc_path = _create_sample_docx(tmp_path)
    out_path = tmp_path / "out.docx"
    changes_path = tmp_path / "changes.json"

    changes_path.write_text(
        json.dumps(
            [
                {"type": "modify", "target_text": "nonexistent text 1", "new_text": "replacement 1"},
                {"type": "modify", "target_text": "nonexistent text 2", "new_text": "replacement 2"},
            ]
        ),
        encoding="utf-8",
    )

    test_args = ["adeu", "apply", str(doc_path), str(changes_path), "-o", str(out_path), "--partial"]
    with patch.object(sys, "argv", test_args):
        with pytest.raises(SystemExit) as exc_info:
            main()
        assert exc_info.value.code == 1

    captured = capsys.readouterr()
    assert "Batch failed — no output was written" in captured.err
    assert not out_path.exists()


def test_cli_default_is_still_atomic(tmp_path, capsys):
    doc_path = _create_sample_docx(tmp_path)
    initial_bytes = doc_path.read_bytes()
    out_path = tmp_path / "out.docx"
    default_redlined = tmp_path / f"{doc_path.stem}_redlined.docx"
    changes_path = tmp_path / "changes.json"

    changes_path.write_text(
        json.dumps(
            [
                {"type": "modify", "target_text": "quick brown", "new_text": "fast blue"},
                {"type": "modify", "target_text": "nonexistent text", "new_text": "replacement"},
            ]
        ),
        encoding="utf-8",
    )

    test_args = ["adeu", "apply", str(doc_path), str(changes_path), "-o", str(out_path)]
    with patch.object(sys, "argv", test_args):
        with pytest.raises(SystemExit) as exc_info:
            main()
        assert exc_info.value.code == 1

    assert not out_path.exists()
    assert not default_redlined.exists()
    assert doc_path.read_bytes() == initial_bytes


def test_cli_mutually_exclusive_salvage_flags(tmp_path, capsys):
    doc_path = _create_sample_docx(tmp_path)
    changes_path = tmp_path / "changes.json"
    changes_path.write_text("[]", encoding="utf-8")

    test_args = ["adeu", "apply", str(doc_path), str(changes_path), "--partial", "--atomic"]
    with patch.object(sys, "argv", test_args):
        with pytest.raises(SystemExit) as exc_info:
            main()
        assert exc_info.value.code == 2

    # argparse must be the one refusing: exit code 2 alone is not evidence,
    # since other input errors also exit 2. The rejection has to name the
    # conflicting flag.
    captured = capsys.readouterr()
    assert "not allowed with argument --partial" in captured.err


def test_partial_json_failed_indices_are_machine_readable(tmp_path):
    doc_path = _create_sample_docx(tmp_path)
    stream = io.BytesIO(doc_path.read_bytes())
    engine = RedlineEngine(stream)

    changes = [
        ModifyText(type="modify", target_text="quick brown", new_text="fast blue", comment=None),
        ModifyText(type="modify", target_text="nonexistent text", new_text="replacement", comment=None),
    ]

    stats = engine.process_batch(changes, partial=True)
    assert stats["status"] == "partial"
    assert "failed" in stats
    assert len(stats["failed"]) == 1
    assert stats["failed"][0]["index"] == 1


def test_partial_rejected_for_text_file_input(tmp_path, capsys):
    doc_path = _create_sample_docx(tmp_path)
    out_path = tmp_path / "out.docx"
    text_path = tmp_path / "modified.txt"
    text_path.write_text("The fast blue fox jumps over the lazy dog.", encoding="utf-8")

    test_args = ["adeu", "apply", str(doc_path), str(text_path), "-o", str(out_path), "--partial"]
    with patch.object(sys, "argv", test_args):
        with pytest.raises(SystemExit) as exc_info:
            main()
        assert exc_info.value.code == 2

    captured = capsys.readouterr()
    assert "--partial is only supported with JSON batch input" in captured.err


def test_partial_rejected_for_live_word_input(tmp_path, capsys):
    # Live Word applies edit-by-edit through COM and has no salvage path, so
    # --partial must be refused up front rather than silently ignored.
    changes_path = tmp_path / "changes.json"
    changes_path.write_text(
        json.dumps([{"type": "modify", "target_text": "quick brown", "new_text": "fast blue"}]),
        encoding="utf-8",
    )

    test_args = ["adeu", "apply", "--live", str(changes_path), "--partial"]
    with patch.object(sys, "argv", test_args):
        with pytest.raises(SystemExit) as exc_info:
            main()
        assert exc_info.value.code == 2

    captured = capsys.readouterr()
    assert "--partial is not supported" in captured.err


def test_pairing_contradiction_still_rejects_whole_batch_under_partial(tmp_path):
    doc_path = _create_sample_docx(tmp_path)
    stream = io.BytesIO(doc_path.read_bytes())
    engine = RedlineEngine(stream)

    changes = [
        AcceptChange(type="accept", target_id="Chg:1"),
        RejectChange(type="reject", target_id="Chg:1"),
    ]

    with pytest.raises(BatchValidationError):
        engine.process_batch(changes, partial=True)


def test_mcp_defaults_to_partial_and_reports_schema_rejects_in_one_list(tmp_path):
    doc_path = _create_sample_docx(tmp_path)
    ctx = get_mock_ctx()

    changes = [
        {"type": "modify", "target_text": "quick brown", "new_text": "fast blue"},
        {"target_text": "missing type field"},
        {"type": "modify", "target_text": "nonexistent text", "new_text": "replacement"},
    ]

    res = run_async(
        process_document_batch(
            reasoning="test",
            original_docx_path=str(doc_path),
            author_name="Tester",
            changes=changes,
            ctx=ctx,
        )
    )
    res_text = extract_content(res)
    assert res_text.startswith("PARTIAL: applied 1 of 3")
    assert "Unable to extract tag using discriminator 'type'" in res_text
    assert "nonexistent text" in res_text or "Target text not found" in res_text

    # A partial success WROTE a file, so it must not carry the batch recovery
    # protocol's "Nothing was written" — the two statements contradict. It
    # reports like any other success: the saved path plus one report section.
    expected_out = doc_path.parent / f"{doc_path.stem}_processed.docx"
    assert "Nothing was written" not in res_text
    assert str(expected_out) in res_text
    assert expected_out.exists()
    assert res_text.count("Detailed Edit Reports:") == 1


def test_partial_failure_payload_within_token_budget(tmp_path):
    doc_path = tmp_path / "sample_budget.docx"
    doc = Document()
    words = [f"token_{i:02d}" for i in range(19)]
    doc.add_paragraph(" ".join(words))
    doc.save(doc_path)

    stream = io.BytesIO(doc_path.read_bytes())
    engine = RedlineEngine(stream)

    changes = []
    for i in range(19):
        changes.append(
            ModifyText(
                type="modify",
                target_text=f"token_{i:02d}",
                new_text=f"replaced_{i:02d}",
                comment=None,
            )
        )
    changes.append(ModifyText(type="modify", target_text="nonexistent text", new_text="replacement", comment=None))

    stats = engine.process_batch(changes, partial=True)
    header = f"PARTIAL: applied {stats['edits_applied']} of {len(changes)} changes. {len(stats['failed'])} failed:"
    assert approx_tokens(header) <= 60

    failure_block = json.dumps(stats["failed"])
    assert approx_tokens(failure_block) <= 500
