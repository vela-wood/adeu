# FILE: tests/test_repro_issue_114_cross_part_revisions.py
#
# Regression tests for GitHub issue #114 (mirrors the Node suite in
# node/packages/core/src/repro_issue_114_cross_part_revisions.test.ts, per
# the dual-engine parity mandate).
#
# RedlineEngine WROTE tracked changes across the whole package while every
# path that READ revision state was rooted at the main part only. Revision
# ids are numbered PER PART, so: a body/header id collision silently
# resolved the body's revision and reported success; an id existing only in
# a header was advertised by the projection yet untargetable; revisions the
# engine itself authored in headers could never be individually resolved;
# and the body-only id scan minted duplicate ids inside header parts.
#
# Fixed by reading revision state across every story part: the id scan
# spans every wordprocessingml part, targeted accept/reject resolve a bare
# id wherever it uniquely lives, a bare id matching several parts is
# REFUSED, and the optional `part` field on accept/reject picks the part
# explicitly.

import io
import re
from zipfile import ZipFile

import pytest
from docx import Document
from docx.oxml import parse_xml

from adeu.ingest import extract_text_from_stream
from adeu.models import AcceptChange, ModifyText, RejectChange
from adeu.redline.engine import BatchValidationError, RedlineEngine

W_ATTR = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'

BODY_MEMBER = "word/document.xml"


def ins_xml(rid: str, author: str, text: str) -> str:
    return (
        f'<w:ins {W_ATTR} w:id="{rid}" w:author="{author}" w:date="2026-01-01T00:00:00Z">'
        f"<w:r><w:t>{text}</w:t></w:r></w:ins>"
    )


def build_doc_bytes(header_ins: str = "", body_ins: str = "") -> bytes:
    """Body paragraph + a header carrying "HEADER MARKER", with optional raw
    tracked insertions injected into each part (per-part id numbering is
    Word's own behavior, so colliding ids across parts are ordinary)."""
    doc = Document()
    sec = doc.sections[0]
    sec.header.paragraphs[0].text = "HEADER MARKER"
    doc.add_paragraph("Body paragraph one.")
    if header_ins:
        sec.header.paragraphs[0]._element.append(parse_xml(header_ins))
    if body_ins:
        doc.paragraphs[-1]._element.append(parse_xml(body_ins))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def header_member(docx_bytes: bytes) -> str:
    with ZipFile(io.BytesIO(docx_bytes)) as z:
        for name in z.namelist():
            if re.fullmatch(r"word/header\d+\.xml", name) and "HEADER MARKER" in z.read(name).decode(
                "utf-8", "replace"
            ):
                return name
    raise AssertionError("fixture has no header part carrying HEADER MARKER")


def part_xml(docx_bytes: bytes, member: str) -> str:
    with ZipFile(io.BytesIO(docx_bytes)) as z:
        return z.read(member).decode("utf-8")


def rev_ids_in(xml: str) -> list:
    return re.findall(r'<w:(?:ins|del)\b[^>]*w:id="(\d+)"', xml)


def text_of(xml: str) -> str:
    return re.sub(r"<[^>]+>", "", xml)


class TestBulkPathsAndProjection:
    """Correct before AND after the fix — pinned so the fix cannot regress them."""

    def test_accept_all_revisions_spans_headers(self):
        src = build_doc_bytes(
            header_ins=ins_xml("900", "Bob", "HeaderInserted"), body_ins=ins_xml("5", "Alice", "BodyInserted")
        )
        hdr = header_member(src)
        engine = RedlineEngine(io.BytesIO(src), author="Reviewer")
        engine.accept_all_revisions()
        out = engine.save_to_stream().getvalue()
        assert rev_ids_in(part_xml(out, BODY_MEMBER)) == []
        assert rev_ids_in(part_xml(out, hdr)) == []
        assert "HeaderInserted" in part_xml(out, hdr)
        assert "BodyInserted" in part_xml(out, BODY_MEMBER)

    def test_reject_all_revisions_spans_headers(self):
        src = build_doc_bytes(
            header_ins=ins_xml("900", "Bob", "HeaderInserted"), body_ins=ins_xml("5", "Alice", "BodyInserted")
        )
        hdr = header_member(src)
        engine = RedlineEngine(io.BytesIO(src), author="Reviewer")
        engine.reject_all_revisions()
        out = engine.save_to_stream().getvalue()
        assert rev_ids_in(part_xml(out, BODY_MEMBER)) == []
        assert rev_ids_in(part_xml(out, hdr)) == []
        assert "HeaderInserted" not in part_xml(out, hdr)
        assert "BodyInserted" not in part_xml(out, BODY_MEMBER)

    def test_unambiguous_body_accept_applies(self):
        src = build_doc_bytes(body_ins=ins_xml("5", "Alice", "BodyInserted"))
        engine = RedlineEngine(io.BytesIO(src), author="Reviewer")
        result = engine.process_batch([AcceptChange(target_id="5")])
        assert result["actions_applied"] == 1
        out = engine.save_to_stream().getvalue()
        assert rev_ids_in(part_xml(out, BODY_MEMBER)) == []
        assert "BodyInserted" in part_xml(out, BODY_MEMBER)

    def test_same_part_different_author_guard_still_fires(self):
        src = build_doc_bytes(body_ins=ins_xml("7", "Alice", "first"))
        doc = Document(io.BytesIO(src))
        doc.paragraphs[-1]._element.append(parse_xml(ins_xml("7", "Bob", "second")))
        buf = io.BytesIO()
        doc.save(buf)
        engine = RedlineEngine(io.BytesIO(buf.getvalue()), author="Reviewer")
        with pytest.raises(BatchValidationError) as exc:
            engine.process_batch([AcceptChange(target_id="7")])
        assert "different authors" in str(exc.value)

    def test_projection_advertises_header_revisions(self):
        src = build_doc_bytes(header_ins=ins_xml("900", "Bob", "HeaderInserted"))
        projection = extract_text_from_stream(io.BytesIO(src))
        assert "Chg:900" in projection
        assert "HeaderInserted" in projection


class TestTargetedResolutionAcrossParts:
    """The fixed behavior for issue #114's four findings plus the `part` selector."""

    def test_f1_bare_id_collision_is_refused_naming_both_parts(self):
        src = build_doc_bytes(
            header_ins=ins_xml("0", "Bob", "HeaderInserted"), body_ins=ins_xml("0", "Alice", "BodyInserted")
        )
        hdr = header_member(src)
        engine = RedlineEngine(io.BytesIO(src), author="Reviewer")
        with pytest.raises(BatchValidationError) as exc:
            engine.process_batch([AcceptChange(target_id="0")])
        message = str(exc.value)
        assert "ambiguous" in message
        assert "word/document.xml" in message
        assert hdr in message
        assert '"part"' in message
        # Refused means NOTHING was resolved — no silent body-wins.
        out = engine.save_to_stream().getvalue()
        assert rev_ids_in(part_xml(out, BODY_MEMBER)) == ["0"]
        assert rev_ids_in(part_xml(out, hdr)) == ["0"]

    def test_f1_part_selector_resolves_exactly_the_named_side(self):
        src = build_doc_bytes(
            header_ins=ins_xml("0", "Bob", "HeaderInserted"), body_ins=ins_xml("0", "Alice", "BodyInserted")
        )
        hdr = header_member(src)
        engine = RedlineEngine(io.BytesIO(src), author="Reviewer")
        result = engine.process_batch([AcceptChange(target_id="0", part=hdr)])
        assert result["actions_applied"] == 1
        out = engine.save_to_stream().getvalue()
        assert rev_ids_in(part_xml(out, hdr)) == []
        assert "HeaderInserted" in part_xml(out, hdr)  # accepted, text kept
        assert rev_ids_in(part_xml(out, BODY_MEMBER)) == ["0"]  # untouched

        # With the header's resolved, the bare id is unique again.
        result2 = engine.process_batch([AcceptChange(target_id="0")])
        assert result2["actions_applied"] == 1
        out2 = engine.save_to_stream().getvalue()
        assert rev_ids_in(part_xml(out2, BODY_MEMBER)) == []
        assert "BodyInserted" in part_xml(out2, BODY_MEMBER)

    def test_f2_header_only_id_resolves_bare(self):
        src = build_doc_bytes(header_ins=ins_xml("900", "Bob", "HeaderInserted"))
        hdr = header_member(src)
        engine = RedlineEngine(io.BytesIO(src), author="Reviewer")
        result = engine.process_batch([AcceptChange(target_id="900")])
        assert result["actions_applied"] == 1
        out = engine.save_to_stream().getvalue()
        assert rev_ids_in(part_xml(out, hdr)) == []
        assert "HeaderInserted" in part_xml(out, hdr)

    def test_f2_not_found_hint_lists_header_ids(self):
        src = build_doc_bytes(header_ins=ins_xml("900", "Bob", "HeaderInserted"))
        engine = RedlineEngine(io.BytesIO(src), author="Reviewer")
        with pytest.raises(BatchValidationError) as exc:
            engine.process_batch([AcceptChange(target_id="555")])
        message = str(exc.value)
        assert "Chg:900" in message
        assert "This document has no tracked changes." not in message

    def test_f3_engine_authored_header_revisions_are_targetable(self):
        src = build_doc_bytes()
        hdr = header_member(src)
        engine = RedlineEngine(io.BytesIO(src), author="Reviewer")
        result = engine.process_batch([ModifyText(target_text="HEADER MARKER", new_text="Amended Header")])
        assert result["edits_applied"] == 1
        mid = engine.save_to_stream().getvalue()
        minted = rev_ids_in(part_xml(mid, hdr))
        assert minted

        # A fresh engine (the normal act-later flow) resolves it; the
        # del+ins pair resolves as one unit.
        engine2 = RedlineEngine(io.BytesIO(mid), author="Reviewer")
        result2 = engine2.process_batch([AcceptChange(target_id=minted[0])])
        assert result2["actions_applied"] == 1
        out = engine2.save_to_stream().getvalue()
        assert rev_ids_in(part_xml(out, hdr)) == []
        assert "Amended Header" in text_of(part_xml(out, hdr))
        assert "HEADER MARKER" not in text_of(part_xml(out, hdr))

    def test_f4_id_scan_spans_parts_so_no_duplicate_minting(self):
        src = build_doc_bytes(
            header_ins=ins_xml("2", "Bob", "HeaderInserted"), body_ins=ins_xml("1", "Alice", "BodyInserted")
        )
        hdr = header_member(src)
        engine = RedlineEngine(io.BytesIO(src), author="Reviewer")
        assert engine.current_id == 2  # scanned from the whole package
        result = engine.process_batch([ModifyText(target_text="HEADER MARKER", new_text="Amended Header")])
        assert result["edits_applied"] == 1
        out = engine.save_to_stream().getvalue()
        header_ids = rev_ids_in(part_xml(out, hdr))
        authors_of_2 = re.findall(r'<w:(?:ins|del)\b[^>]*w:id="2"[^>]*w:author="([^"]*)"', part_xml(out, hdr))
        assert authors_of_2 == ["Bob"]  # only the pre-existing revision
        assert max(int(i) for i in header_ids) > 2  # fresh ids above package max
        assert len(set(header_ids)) == len(header_ids)  # no duplicates

    def test_f5_consecutive_sessions_mint_distinct_ids_and_resolve(self):
        # No foreign/injected revisions anywhere — pure product usage.
        src = build_doc_bytes()
        hdr = header_member(src)

        engine1 = RedlineEngine(io.BytesIO(src), author="Session One")
        engine1.process_batch([ModifyText(target_text="HEADER MARKER", new_text="Amended Header")])
        after1 = engine1.save_to_stream().getvalue()
        header_ids = rev_ids_in(part_xml(after1, hdr))
        assert header_ids

        engine2 = RedlineEngine(io.BytesIO(after1), author="Session Two")
        engine2.process_batch([ModifyText(target_text="Body paragraph one.", new_text="Body paragraph two.")])
        after2 = engine2.save_to_stream().getvalue()
        body_ids = rev_ids_in(part_xml(after2, BODY_MEMBER))
        assert body_ids
        assert not set(body_ids) & set(header_ids)  # no cross-part collision forms

        engine3 = RedlineEngine(io.BytesIO(after2), author="Session Three")
        result = engine3.process_batch([AcceptChange(target_id=header_ids[0]), AcceptChange(target_id=body_ids[0])])
        assert result["actions_applied"] == 2
        out = engine3.save_to_stream().getvalue()
        assert rev_ids_in(part_xml(out, hdr)) == []
        assert rev_ids_in(part_xml(out, BODY_MEMBER)) == []
        assert "Amended Header" in text_of(part_xml(out, hdr))
        assert "Body paragraph two." in text_of(part_xml(out, BODY_MEMBER))

    def test_part_selectors_normalize_and_do_not_collide_across_parts(self):
        src = build_doc_bytes(
            header_ins=ins_xml("1", "Bob", "HeaderInserted"), body_ins=ins_xml("1", "Alice", "BodyInserted")
        )
        hdr = header_member(src)
        engine = RedlineEngine(io.BytesIO(src), author="Reviewer")
        # accept header's Chg:1 but reject the body's Chg:1 — with per-part
        # tracking this is NOT a duplicate or a contradiction.
        result = engine.process_batch(
            [
                AcceptChange(target_id="1", part="/" + hdr),  # leading slash normalizes
                RejectChange(target_id="1", part="word/document.xml"),
            ]
        )
        assert result["actions_applied"] == 2
        out = engine.save_to_stream().getvalue()
        assert "HeaderInserted" in part_xml(out, hdr)  # accepted
        assert "BodyInserted" not in part_xml(out, BODY_MEMBER)  # rejected
        assert rev_ids_in(part_xml(out, hdr)) == []
        assert rev_ids_in(part_xml(out, BODY_MEMBER)) == []

    def test_part_selector_errors_are_actionable(self):
        src = build_doc_bytes(header_ins=ins_xml("900", "Bob", "HeaderInserted"))
        hdr = header_member(src)

        engine = RedlineEngine(io.BytesIO(src), author="Reviewer")
        with pytest.raises(BatchValidationError) as exc:
            engine.process_batch([AcceptChange(target_id="900", part="word/nope.xml")])
        assert "not a package part" in str(exc.value)
        assert hdr in str(exc.value)

        engine = RedlineEngine(io.BytesIO(src), author="Reviewer")
        with pytest.raises(BatchValidationError) as exc:
            engine.process_batch([AcceptChange(target_id="900", part="word/document.xml")])
        message = str(exc.value)
        assert "word/document.xml" in message
        assert hdr in message  # says where the id actually lives
