# FILE: tests/test_repro_comment_threading_anchoring_typography.py
"""
Repro tests for BUG_comment_threading_anchoring_and_typography.md
(reported 2026-08-11, Adeu 2.1.0 / 56a97cf).

Four defects found driving Adeu from a generic agent over MCP against the
WAWD Model Stipulated Protective Order:

  B1  `ReplyComment` silently creates a NEW TOP-LEVEL comment instead of
      threading. `_find_thread_root_para_id` returns None (the parent comment
      carries no w14:paraId — the ordinary shape for a legacy Word comment or
      any third-party generator), `w15:paraIdParent` is simply never written,
      and the reply degrades to a thread root while apply_review_actions
      reports `(1, 0, 0)`. The agent acts on a success it did not get, retries,
      and makes the document worse.

  B2  Accepting a change destroys the HUMAN's comment with no disclosure:
      `removed_comments` is hard-coded to 0 unless the caller asked for
      wholesale comment removal, so a comment deleted because its anchor was
      consumed is reported as zero removals. The MCP `accept_all_changes`
      surface additionally inverted the library default
      (`remove_comments=True` against an API that defaults to False), giving
      an agent "delete every comment" for free with "accept all changes".

  B3  `w16cid:durableId` is minted across the full 32-bit range. Word parses
      ST_LongHexNumber as a SIGNED 32-bit integer, so a high-bit-set id is
      negative, Word fails to bind the comment and collapses the anchor to a
      zero-length point. Roughly half of all Adeu-authored comments open in
      Word anchored to nothing. Word-verified in the report via
      `Comment.Scope`; Word's own ids are always high-bit clear.

  B4  Edits rewrite curly quotes/apostrophes to straight ones in text the
      caller never targeted: the MATCHER is smart-quote-insensitive
      (`DocumentMapper._replace_smart_quotes`) while the WRITER word-diffs the
      document's real slice against the caller's literal `new_text`, so every
      typographic difference becomes a real tracked change on a provision
      nobody touched.

Every test here is written test-first: it fails on the pre-fix engine.
"""

import asyncio
import io
import re
import zipfile

import pytest
from docx import Document
from docx.opc.part import XmlPart
from docx.oxml import parse_xml
from docx.oxml.ns import qn

from adeu.ingest import extract_text_from_stream
from adeu.mcp_components.tools.document import accept_all_changes
from adeu.models import AcceptChange, ModifyText, ReplyComment
from adeu.redline.comments import CommentsManager
from adeu.redline.engine import BatchValidationError, RedlineEngine

W14 = "http://schemas.microsoft.com/office/word/2010/wordml"
W15 = "http://schemas.microsoft.com/office/word/2012/wordml"
W16CID = "http://schemas.microsoft.com/office/word/2016/wordml/cid"

CT_COMMENTS = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
CT_EXTENDED = "application/vnd.openxmlformats-officedocument.wordprocessingml.commentsExtended+xml"
CT_IDS = "application/vnd.openxmlformats-officedocument.wordprocessingml.commentsIds+xml"


class MockContext:
    """Mock FastMCP Context; absorbs the async logging calls."""

    async def info(self, msg, **kwargs):
        pass

    async def debug(self, msg, **kwargs):
        pass

    async def warning(self, msg, **kwargs):
        pass

    async def error(self, msg, **kwargs):
        pass


def doc_stream(*paragraphs: str) -> io.BytesIO:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def part_bytes(stream: io.BytesIO, content_type: str) -> bytes:
    """The raw bytes of the (single) package part with `content_type`."""
    stream.seek(0)
    doc = Document(io.BytesIO(stream.getvalue()))
    for part in doc.part.package.parts:
        if part.content_type == content_type:
            return part.blob
    raise AssertionError(f"no part with content type {content_type}")


def annotated_stream(body: str, target: str, comment: str, author: str) -> io.BytesIO:
    """A document carrying one comment authored by `author`, anchored to `target`."""
    engine = RedlineEngine(doc_stream(body), author=author)
    engine.apply_edits([ModifyText(target_text=target, new_text=target, comment=comment)])
    return engine.save_to_stream()


def _rewrite_comments_part(stream: io.BytesIO, mutate) -> io.BytesIO:
    """Applies `mutate(comments_root)` to word/comments.xml and re-saves."""
    stream.seek(0)
    doc = Document(io.BytesIO(stream.getvalue()))
    comments_part = next(p for p in doc.part.package.parts if p.content_type == CT_COMMENTS)

    element = parse_xml(comments_part.blob)
    mutate(element)

    replacement = XmlPart(comments_part.partname, comments_part.content_type, element, comments_part.package)
    package = comments_part.package
    package.parts[package.parts.index(comments_part)] = replacement
    for rel in doc.part.rels.values():
        if not rel.is_external and rel.target_part is comments_part:
            rel._target = replacement

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def strip_para_ids_from_comments(stream: io.BytesIO) -> io.BytesIO:
    """
    Removes `w14:paraId` from every comment paragraph, producing the ordinary
    LEGACY comment shape: a comment body with no modern paragraph identity.
    Word itself writes such comments (pre-2013 documents), as does every
    generator that does not implement the modern-comments extensions. The
    commentsExtended / commentsIds parts stay in place, so the document is
    still on Word's modern-comments path.
    """

    def mutate(root):
        stripped = 0
        for comment in root.findall(qn("w:comment")):
            for para in comment.findall(qn("w:p")):
                if para.get(qn("w14:paraId")) is not None:
                    del para.attrib[qn("w14:paraId")]
                    stripped += 1
        assert stripped, "fixture precondition: the comment had no w14:paraId to strip"

    return _rewrite_comments_part(stream, mutate)


def empty_comment_bodies(stream: io.BytesIO) -> io.BytesIO:
    """
    Strips every block-level child from each `<w:comment>`, leaving a comment
    with NO paragraph at all. `EG_BlockLevelElts` is `minOccurs="0"` so this is
    schema-legal, and it is the one shape where a paragraph identity genuinely
    cannot be minted — i.e. where threading is truly impossible. What must NOT
    happen is a silent top-level comment.
    """

    def mutate(root):
        emptied = 0
        for comment in root.findall(qn("w:comment")):
            for child in list(comment):
                comment.remove(child)
                emptied += 1
        assert emptied, "fixture precondition: the comment had no body to strip"

    return _rewrite_comments_part(stream, mutate)


# ---------------------------------------------------------------------------
# B3 — durableId high bit silently unanchors comments
# ---------------------------------------------------------------------------


class TestB3DurableIdSignedRange:
    """
    `w16cid:durableId` is ST_LongHexNumber but Word reads it as a SIGNED 32-bit
    integer. A high-bit-set value is negative; Word then fails to bind the
    comment and collapses its anchor to a zero-length point — no error, no
    repair prompt, nothing wrong-looking in the XML. Word's own durable ids are
    never negative, so "high bit clear" is the invariant to hold.
    """

    # 256 samples: a generator drawing from the full 32-bit range fails this
    # with probability 1 - 2^-256, i.e. deterministically in practice.
    SAMPLES = 256

    def test_generated_durable_ids_are_never_negative_int32(self):
        manager = CommentsManager(Document())
        negatives = [
            value
            for value in (manager._generate_durable_id() for _ in range(self.SAMPLES))
            if int(value, 16) > 0x7FFFFFFF
        ]
        assert not negatives, (
            f"{len(negatives)}/{self.SAMPLES} generated durableIds have the high bit set "
            f"(e.g. {negatives[:4]}). Word reads w16cid:durableId as a signed 32-bit "
            "integer: a negative id silently unanchors the comment."
        )

    def test_generated_durable_ids_are_eight_hex_digits(self):
        manager = CommentsManager(Document())
        for _ in range(64):
            value = manager._generate_durable_id()
            assert re.fullmatch(r"[0-9A-F]{8}", value), (
                f"durableId {value!r} is not an 8-digit upper-case ST_LongHexNumber"
            )

    def test_written_comments_carry_word_readable_durable_ids(self):
        """End-to-end: every durableId in the SAVED package must be positive."""
        engine = RedlineEngine(doc_stream("Alpha. Beta. Gamma. Delta. Epsilon."), author="Adeu AI")
        for word in ("Alpha", "Beta", "Gamma", "Delta", "Epsilon"):
            engine.apply_edits([ModifyText(target_text=word, new_text=word, comment=f"Note on {word}.")])

        ids_xml = part_bytes(engine.save_to_stream(), CT_IDS).decode("utf-8")
        durable_ids = re.findall(r'w16cid:durableId="([0-9A-Fa-f]+)"', ids_xml)
        assert len(durable_ids) == 5, f"expected one durableId per comment, got {durable_ids}"

        negatives = [d for d in durable_ids if int(d, 16) > 0x7FFFFFFF]
        assert not negatives, (
            f"saved commentsIds.xml carries negative (signed int32) durableIds {negatives}; "
            "Word drops the anchor for each of them"
        )

    def test_para_id_and_rsid_are_masked_exactly_like_durable_id(self):
        """
        RETRACTION. This test used to assert the opposite — that paraId and
        rsid "keep the full 32-bit range", because only durableId was believed
        to carry the signed-int32 constraint. That belief was wrong, it was
        recorded here as a pin, and it is why the same bug shipped a second
        time three weeks later: `w14:paraId` above 0x7FFFFFFF is discarded by
        Word exactly like a durableId, dangling every `w15:paraIdParent` that
        pointed at it and dropping the reply out of its thread
        (BUG_paraId_signed_int32_thread_collapse.md, B5, Word-verified).

        There is no attribute for which the high half is safe. The full
        coverage lives in tests/test_repro_para_id_signed_int32.py; what stays
        here is the correction, next to the reasoning it corrects.
        """
        manager = CommentsManager(Document())
        values = [manager._generate_para_id() for _ in range(512)]
        values += [manager._generate_rsid() for _ in range(512)]
        values += [manager._generate_durable_id() for _ in range(512)]
        assert all(0 < int(v, 16) <= 0x7FFFFFFF for v in values), (
            "paraId, rsid and durableId are all ST_LongHexNumber and Word reads all three as "
            "signed 32-bit integers: they share one generator and one range"
        )


# ---------------------------------------------------------------------------
# B1 — ReplyComment silently produces a new top-level comment
# ---------------------------------------------------------------------------


class TestB1ReplyThreading:
    def test_reply_threads_when_parent_has_no_para_id(self):
        """
        The reported shape: the parent comment carries no `w14:paraId`, so
        `_find_thread_root_para_id` resolves nothing, `w15:paraIdParent` is
        never written and the "reply" becomes a second top-level thread —
        reported as applied. Threading must actually happen.
        """
        source = strip_para_ids_from_comments(
            annotated_stream(
                "The parties shall confer in good faith before moving to compel.",
                "confer in good faith",
                "This should reference the protective order.",
                "Sarah Chen",
            )
        )

        engine = RedlineEngine(source, author="Agent")
        parent_id = next(iter(engine.comments_manager.extract_comments_data()))
        applied, skipped, _already = engine.apply_review_actions(
            [ReplyComment(target_id=f"Com:{parent_id}", text="Addressed in the revised clause.")]
        )
        assert (applied, skipped) == (1, 0), "the reply was not applied at all"

        final = engine.save_to_stream()
        data = engine.comments_manager.extract_comments_data()
        replies = {cid: c for cid, c in data.items() if cid != parent_id}
        assert len(replies) == 1, f"expected exactly one reply, got {data}"
        reply_id, reply = next(iter(replies.items()))

        assert reply["parent_id"] == parent_id, (
            "the reply was created as a separate TOP-LEVEL comment instead of being "
            f"threaded under Com:{parent_id} — this is the defect an agent cannot detect "
            f"(comments read back as {data})"
        )

        ext_xml = part_bytes(final, CT_EXTENDED).decode("utf-8")
        assert "w15:paraIdParent" in ext_xml, (
            f"commentsExtended.xml carries no w15:paraIdParent, so Word renders two unrelated threads:\n{ext_xml}"
        )

    def test_reply_never_silently_degrades_to_a_thread_root(self):
        """
        The load-bearing invariant ("if only one thing gets fixed: make B1
        loud"). Whatever the reason threading cannot be resolved, a `reply`
        must NEVER silently succeed as a new top-level comment: either it
        threads, or the action fails loudly.
        """
        source = empty_comment_bodies(
            annotated_stream(
                "Discovery Material shall be produced within thirty days.",
                "within thirty days",
                "Confirm this matches the scheduling order.",
                "Sarah Chen",
            )
        )

        engine = RedlineEngine(source, author="Agent")
        parent_id = next(iter(engine.comments_manager.extract_comments_data()))

        applied, skipped, _already = engine.apply_review_actions(
            [ReplyComment(target_id=f"Com:{parent_id}", text="Addressed.")]
        )

        data = engine.comments_manager.extract_comments_data()
        orphans = [cid for cid, c in data.items() if cid != parent_id and c.get("parent_id") is None]
        assert not orphans, (
            f"reply on Com:{parent_id} silently became top-level comment(s) {orphans} and was "
            f"reported as {applied} applied / {skipped} skipped. A reply that quietly becomes a "
            "new thread is worse than a failed call: the caller acts on a success it did not get."
        )
        assert (applied, skipped) == (0, 1), (
            "unresolvable threading must be reported as a skipped action, not silent success"
        )
        assert any("thread" in d.lower() or "repl" in d.lower() for d in engine.skipped_details), (
            "the failure must be explained to the caller, got: " + repr(engine.skipped_details)
        )

    def test_batch_rejects_a_reply_that_cannot_be_threaded(self):
        """A batch must never SAVE a document in which a reply silently rooted."""
        source = empty_comment_bodies(
            annotated_stream(
                "The receiving party shall bear the cost of production.",
                "bear the cost",
                "Whose cost is this really?",
                "Sarah Chen",
            )
        )
        engine = RedlineEngine(source, author="Agent")
        parent_id = next(iter(engine.comments_manager.extract_comments_data()))

        with pytest.raises(BatchValidationError) as excinfo:
            engine.process_batch([ReplyComment(target_id=f"Com:{parent_id}", text="Addressed.")])
        assert "thread" in str(excinfo.value).lower(), str(excinfo.value)

        # And the document must be untouched: no stray comment was written.
        assert len(engine.comments_manager.extract_comments_data()) == 1

    def test_reply_to_a_reply_still_flattens_to_the_thread_root(self):
        """
        Modern Word flattens a thread: every reply points at the ROOT's paraId.
        Fixing B1 must not turn replies-to-replies into nested chains.
        """
        engine = RedlineEngine(doc_stream("Root anchor text here."), author="Alice")
        engine.apply_edits([ModifyText(target_text="anchor", new_text="anchor", comment="Root topic")])
        root_id = next(iter(engine.comments_manager.extract_comments_data()))

        engine.author = "Bob"
        engine.apply_review_actions([ReplyComment(target_id=f"Com:{root_id}", text="First reply")])
        first_reply = next(cid for cid in engine.comments_manager.extract_comments_data() if cid != root_id)

        engine.author = "Carol"
        engine.apply_review_actions([ReplyComment(target_id=f"Com:{first_reply}", text="Second reply")])

        data = engine.comments_manager.extract_comments_data()
        assert len(data) == 3, data
        for cid, comment in data.items():
            if cid == root_id:
                assert comment["parent_id"] is None
            else:
                assert comment["parent_id"] == root_id, (
                    f"Com:{cid} must thread onto the ROOT (Com:{root_id}), got {comment['parent_id']}"
                )

    def test_reply_backfills_a_parent_missing_from_the_auxiliary_parts(self):
        """
        The parent HAS a `w14:paraId` but no `w15:commentEx` / `w16cid:commentId`
        entry — a shape hand-built and third-party documents produce. Word
        consults all three parts, so a `paraIdParent` pointing at an
        unregistered paragraph drops the reply out of its thread just as surely
        as a missing attribute would.
        """
        source = annotated_stream(
            "Discovery shall proceed under the model order.",
            "the model order",
            "Which model order?",
            "Sarah Chen",
        )
        source.seek(0)
        doc = Document(io.BytesIO(source.getvalue()))
        for part in doc.part.package.parts:
            if part.content_type in (CT_EXTENDED, CT_IDS):
                element = parse_xml(part.blob)
                for child in list(element):
                    element.remove(child)
                replacement = XmlPart(part.partname, part.content_type, element, part.package)
                pkg = part.package
                pkg.parts[pkg.parts.index(part)] = replacement
                for rel in doc.part.rels.values():
                    if not rel.is_external and rel.target_part is part:
                        rel._target = replacement
        stripped = io.BytesIO()
        doc.save(stripped)
        stripped.seek(0)

        engine = RedlineEngine(stripped, author="Agent")
        parent_id = next(iter(engine.comments_manager.extract_comments_data()))
        applied, skipped, _ = engine.apply_review_actions(
            [ReplyComment(target_id=f"Com:{parent_id}", text="The WAWD model order.")]
        )
        assert (applied, skipped) == (1, 0)

        final = engine.save_to_stream()
        data = engine.comments_manager.extract_comments_data()
        reply_id = next(cid for cid in data if cid != parent_id)
        assert data[reply_id]["parent_id"] == parent_id, data

        comments_xml = part_bytes(final, CT_COMMENTS).decode("utf-8")
        ext_xml = part_bytes(final, CT_EXTENDED).decode("utf-8")
        ids_xml = part_bytes(final, CT_IDS).decode("utf-8")
        para_ids = set(re.findall(r'w14:paraId="([0-9A-Fa-f]{8})"', comments_xml))
        assert len(para_ids) == 2, para_ids
        assert para_ids <= set(re.findall(r'w15:paraId="([0-9A-Fa-f]{8})"', ext_xml)), (
            "the parent was never backfilled into commentsExtended:\n" + ext_xml
        )
        assert para_ids <= set(re.findall(r'w16cid:paraId="([0-9A-Fa-f]{8})"', ids_xml)), (
            "the parent was never backfilled into commentsIds:\n" + ids_xml
        )

    def test_threaded_reply_registers_in_every_modern_comment_part(self):
        """
        A repaired parent must be fully registered, not half: Word consults
        commentsExtended (threading) and commentsIds (durable identity)
        together, and a paraId present in one but not the other is exactly the
        state that drops a comment out of the modern-comments path.
        """
        source = strip_para_ids_from_comments(
            annotated_stream(
                "Attorney's Eyes Only material stays with outside counsel.",
                "Attorney's Eyes Only",
                "Is this tier defined?",
                "Sarah Chen",
            )
        )
        engine = RedlineEngine(source, author="Agent")
        parent_id = next(iter(engine.comments_manager.extract_comments_data()))
        engine.apply_review_actions([ReplyComment(target_id=f"Com:{parent_id}", text="Defined in §2.")])
        final = engine.save_to_stream()

        comments_xml = part_bytes(final, CT_COMMENTS).decode("utf-8")
        ext_xml = part_bytes(final, CT_EXTENDED).decode("utf-8")
        ids_xml = part_bytes(final, CT_IDS).decode("utf-8")

        para_ids = set(re.findall(r'w14:paraId="([0-9A-Fa-f]{8})"', comments_xml))
        assert len(para_ids) == 2, f"both the parent and the reply need a paraId in comments.xml, found {para_ids}"
        ex_para_ids = set(re.findall(r'w15:paraId="([0-9A-Fa-f]{8})"', ext_xml))
        id_para_ids = set(re.findall(r'w16cid:paraId="([0-9A-Fa-f]{8})"', ids_xml))
        assert para_ids <= ex_para_ids, f"commentsExtended is missing {para_ids - ex_para_ids}"
        assert para_ids <= id_para_ids, f"commentsIds is missing {para_ids - id_para_ids}"


# ---------------------------------------------------------------------------
# B2 — accepting changes destroys the human's comment
# ---------------------------------------------------------------------------


class TestB2CommentDestructionDisclosure:
    """
    B2's substance is that comment destruction was SILENT, ANONYMOUS and
    UNAVOIDABLE — not that it was wrong. The agent-facing surfaces deliberately
    keep `remove_comments=True` as their default: `accept_all_changes` /
    `adeu accept-all` exist to produce a distributable clean document, and
    shipping a counterparty a file that still carries internal review notes is
    the more expensive failure (QA_ISSUES_DISCOVERED #10, "Confidentiality
    risk"). What these tests pin is that the caller can now SEE the default,
    OVERRIDE it, and is told exactly whose comments were destroyed.
    """

    def test_library_api_does_not_remove_comments_by_default(self):
        """
        The library contract is the opposite of the tool's, and stays that way:
        `accept_all_revisions()` accepts REVISIONS. An SDK caller who did not
        ask for comment removal keeps the comments whose anchors survive.
        """
        engine = RedlineEngine(
            annotated_stream(
                "The parties shall meet and confer. A second clause stands alone.",
                "A second clause",
                "Standalone reviewer note.",
                "Sarah Chen",
            ),
            author="Agent",
        )
        engine.apply_edits([ModifyText(target_text="meet and confer", new_text="confer in good faith")])

        counts = engine.accept_all_revisions()
        assert counts["removed_comments"] == 0, counts
        assert "Standalone reviewer note." in extract_text_from_stream(engine.save_to_stream())

    def test_accept_all_reports_comments_it_actually_deletes(self):
        """
        A comment whose anchor is consumed by an accepted DELETION is removed
        even with remove_comments=False (Word does the same) — so the count
        must say so. Reporting 0 while a human's comment is gone is precisely
        the silent data loss that let the agent rationalise the run as a
        success.
        """
        engine = RedlineEngine(
            annotated_stream(
                "Producing party may designate material Attorney's Eyes Only at its discretion.",
                "Attorney's Eyes Only",
                "Strike this tier — it is not in the model order.",
                "Sarah Chen",
            ),
            author="Agent",
        )
        engine.apply_edits([ModifyText(target_text="Attorney's Eyes Only", new_text="")])

        before = set(engine._existing_comment_ids())
        assert before, "fixture precondition: the document must carry the reviewer's comment"

        counts = engine.accept_all_revisions()

        after = set(engine._existing_comment_ids())
        removed = before - after
        assert removed, "fixture precondition: accepting the deletion should consume the comment anchor"
        assert counts["removed_comments"] == len(removed), (
            f"accept_all_revisions deleted {sorted(removed)} but reported "
            f"removed_comments={counts['removed_comments']}: a human's comment vanished with "
            "the books saying nothing happened"
        )

    def test_accept_action_note_names_the_comment_author(self):
        """
        "Never silently delete a comment authored by someone other than the
        caller." The disclosure already exists but is anonymous — an agent
        reading it cannot tell it just destroyed the reviewer's input, which
        is exactly how the run's data loss got rationalised as success.
        """
        engine = RedlineEngine(
            annotated_stream(
                "Producing party may designate material Attorney's Eyes Only at its discretion.",
                "Attorney's Eyes Only",
                "Strike this tier.",
                "Sarah Chen",
            ),
            author="Agent",
        )
        engine.apply_edits([ModifyText(target_text="Attorney's Eyes Only", new_text="")])

        stream = engine.save_to_stream()
        reviewer_engine = RedlineEngine(stream, author="Agent")
        raw = extract_text_from_stream(io.BytesIO(stream.getvalue()))
        change_ids = re.findall(r"\[Chg:(\d+) delete\]", raw)
        assert change_ids, f"no tracked deletion to accept:\n{raw}"

        reviewer_engine.apply_review_actions([AcceptChange(target_id=f"Chg:{change_ids[0]}")])
        details = "\n".join(reviewer_engine.skipped_details)

        assert "Com:" in details, f"comment removal was not disclosed at all:\n{details}"
        assert "Sarah Chen" in details, (
            "the disclosure must name the comment's AUTHOR so the caller can see it destroyed "
            f"someone else's review content:\n{details}"
        )

    @staticmethod
    def _annotated_file(tmp_path):
        source = tmp_path / "protective_order.docx"
        engine = RedlineEngine(
            annotated_stream(
                "The parties shall meet and confer. A second clause stands alone.",
                "A second clause",
                "Standalone reviewer note.",
                "Sarah Chen",
            ),
            author="Agent",
        )
        engine.apply_edits([ModifyText(target_text="meet and confer", new_text="confer in good faith")])
        source.write_bytes(engine.save_to_stream().getvalue())
        return source

    def test_mcp_accept_all_changes_removes_comments_by_default(self, tmp_path):
        """
        The tool produces a DISTRIBUTABLE clean document, so comment removal
        stays its default (QA_ISSUES_DISCOVERED #10 logged the opposite as a
        confidentiality risk). What changed is that the destruction is no longer
        silent: the response names every deleted comment and its author.
        """
        source = self._annotated_file(tmp_path)
        out = tmp_path / "accepted.docx"
        result = asyncio.run(
            accept_all_changes(
                reasoning="test",
                docx_path=str(source),
                ctx=MockContext(),
                output_path=str(out),
            )
        )

        final_raw = extract_text_from_stream(io.BytesIO(out.read_bytes()))
        assert "Standalone reviewer note." not in final_raw
        assert "Comments removed: 1" in result, result
        assert "Sarah Chen" in result, "the response must name whose review content it destroyed:\n" + result

        with zipfile.ZipFile(out, "r") as zf:
            leaked = [n for n in zf.namelist() if "comments" in n.lower()]
        assert not leaked, f"comment parts leaked into the output package: {leaked}"

    def test_mcp_accept_all_changes_can_keep_comments_on_request(self, tmp_path):
        """
        The inversion B2 objected to was that it was UNAVOIDABLE. A caller whose
        review conversation is still live must be able to accept the tracked
        changes without destroying the annotations.
        """
        source = self._annotated_file(tmp_path)
        out = tmp_path / "accepted_annotated.docx"
        result = asyncio.run(
            accept_all_changes(
                reasoning="test",
                docx_path=str(source),
                ctx=MockContext(),
                output_path=str(out),
                remove_comments=False,
            )
        )

        final_raw = extract_text_from_stream(io.BytesIO(out.read_bytes()))
        assert "Standalone reviewer note." in final_raw, (
            f"remove_comments=False still destroyed the reviewer's comment (tool said: {result})"
        )
        assert "{++" not in final_raw and "{--" not in final_raw, (
            "the tracked changes should still have been accepted:\n" + final_raw
        )

    def test_published_tool_documents_the_flag_and_its_default(self):
        """
        Asserted against the PUBLISHED tool (mcp.list_tools), not the source:
        MCP clients drop optional-property descriptions in transit and rewrite
        required[] (AI_CONTEXT §7a), so the tool DESCRIPTION is the only channel
        guaranteed to reach the model. A destructive default the model cannot
        see is exactly the silent inversion B2 reported — the fix is
        disclosure, so pin the disclosure.
        """
        from adeu.server import mcp

        tools = asyncio.run(mcp.list_tools())
        tool = next(t for t in tools if t.name == "accept_all_changes")
        description = tool.description or ""

        assert "remove_comments" in description, "the description must name the parameter:\n" + description
        assert re.search(r"default\w*\s+true", description, re.IGNORECASE), (
            "the description must state that comment removal is the DEFAULT:\n" + description
        )
        assert re.search(r"remove_comments=false", description, re.IGNORECASE), (
            "the description must tell the caller how to opt out:\n" + description
        )
        # §7a: real clients truncate descriptions at ~2048 chars INCLUDING the
        # appended build tag.
        assert len(description) < 2048, f"description is {len(description)} chars"

        # One JSON type per property (§7a strips property-level unions to {}),
        # and the published default must match the documented one.
        prop = tool.parameters["properties"]["remove_comments"]
        assert prop.get("type") == "boolean", prop
        assert prop.get("default") is True, prop
        assert "remove_comments" not in tool.parameters.get("required", [])

    def test_cli_help_states_the_flag_and_its_default(self, capsys):
        """The CLI's human reader gets the same disclosure via --help."""
        import sys
        from unittest.mock import patch

        from adeu.cli import main

        with patch.object(sys, "argv", ["adeu", "accept-all", "--help"]):
            with pytest.raises(SystemExit):
                main()
        help_text = capsys.readouterr().out

        assert "--remove-comments" in help_text, help_text
        assert "--no-remove-comments" in help_text, "the opt-out must be discoverable from --help:\n" + help_text
        assert re.search(r"default", help_text, re.IGNORECASE), help_text


# ---------------------------------------------------------------------------
# B4 — edits rewrite curly quotes to straight quotes in untargeted text
# ---------------------------------------------------------------------------

CURLY_BODY = (
    "All Discovery Material designated as \u201cConfidential\u201d under the "
    "parties\u2019 Master Agreement shall be produced within thirty days."
)


class TestB4TypographyPreservation:
    def test_untargeted_curly_quotes_are_not_rewritten(self):
        """
        The caller changes "thirty" to "sixty". Its target/new strings carry
        STRAIGHT quotes because that is how models normalise typography; the
        matcher forgives that, then the writer word-diffs the document's real
        slice against the caller's literal text and emits a tracked change for
        every quote. Exactly ONE tracked change belongs in this document.
        """
        engine = RedlineEngine(doc_stream(CURLY_BODY), author="Agent")
        engine.apply_edits(
            [
                ModifyText(
                    target_text=(
                        'designated as "Confidential" under the parties\' Master Agreement '
                        "shall be produced within thirty days"
                    ),
                    new_text=(
                        'designated as "Confidential" under the parties\' Master Agreement '
                        "shall be produced within sixty days"
                    ),
                )
            ]
        )

        raw = extract_text_from_stream(engine.save_to_stream())
        deletions = re.findall(r"\{--(.*?)--\}", raw, flags=re.S)
        insertions = re.findall(r"\{\+\+(.*?)\+\+\}", raw, flags=re.S)

        assert deletions == ["thirty"], (
            "the only tracked deletion may be the word the caller actually changed; "
            f"got {deletions} — the rest are pure punctuation rewrites:\n{raw}"
        )
        assert insertions == ["sixty"], f"unexpected insertions {insertions}:\n{raw}"

    def test_document_keeps_its_own_typography(self):
        """
        In a court model the surrounding typography is the court's, not ours.
        The saved document must still carry its curly characters.
        """
        engine = RedlineEngine(doc_stream(CURLY_BODY), author="Agent")
        engine.apply_edits(
            [
                ModifyText(
                    target_text="parties' Master Agreement shall be produced within thirty days",
                    new_text="parties' Master Agreement shall be produced within sixty days",
                )
            ]
        )
        clean = extract_text_from_stream(engine.save_to_stream(), clean_view=True)
        assert "\u2019" in clean, f"the document's curly apostrophe was straightened:\n{clean!r}"
        assert "\u201c" in clean and "\u201d" in clean, (
            f"the document's curly double quotes were straightened:\n{clean!r}"
        )
        assert "'" not in clean.replace("\u2019", ""), f"a straight apostrophe leaked in:\n{clean!r}"

    def test_typography_only_edit_is_a_no_op(self):
        """
        "If target and new differ only by normalised punctuation, the correct
        number of tracked changes is zero." An LLM re-emitting a clause with
        straightened quotes must not redline the clause.
        """
        engine = RedlineEngine(doc_stream(CURLY_BODY), author="Agent")
        engine.apply_edits(
            [
                ModifyText(
                    target_text='designated as "Confidential" under the parties\' Master Agreement',
                    new_text='designated as "Confidential" under the parties\' Master Agreement',
                )
            ]
        )
        raw = extract_text_from_stream(engine.save_to_stream())
        assert "{--" not in raw and "{++" not in raw, f"a punctuation-only round-trip produced tracked changes:\n{raw}"

    def test_deliberate_typography_change_still_applies(self):
        """
        The repair must key on the MATCH being typography-forgiving, not on
        typography being unchangeable: a caller who quotes the document's real
        characters and asks for different ones still gets the change.
        """
        engine = RedlineEngine(doc_stream(CURLY_BODY), author="Agent")
        engine.apply_edits(
            [
                ModifyText(
                    target_text="\u201cConfidential\u201d",
                    new_text='"Confidential"',
                )
            ]
        )
        raw = extract_text_from_stream(engine.save_to_stream())
        assert "{--" in raw and "{++" in raw, "an explicitly requested typography change was swallowed:\n" + raw

    def test_curly_quotes_the_caller_does_change_are_honoured(self):
        """
        A real edit INSIDE a smart-quoted phrase still applies, and the quote
        characters around it stay the document's own.
        """
        engine = RedlineEngine(doc_stream(CURLY_BODY), author="Agent")
        engine.apply_edits(
            [
                ModifyText(
                    target_text='designated as "Confidential" under',
                    new_text='designated as "Highly Confidential" under',
                )
            ]
        )
        raw = extract_text_from_stream(engine.save_to_stream())
        clean = extract_text_from_stream(engine.save_to_stream(), clean_view=True)
        assert "Highly Confidential" in clean, clean
        assert "\u201cHighly Confidential\u201d" in clean, (
            f"the document's curly quotes must still wrap the replacement:\n{clean!r}"
        )
        assert '"Highly Confidential"' not in clean, clean
        # Adding a word before "Confidential" is a pure insertion: no quote
        # character is deleted, and the edit does not fragment.
        assert raw.count("{--") == 0, f"a quote character was redlined away:\n{raw}"
        assert raw.count("{++") == 1, f"the edit fragmented into several redlines:\n{raw}"


# ---------------------------------------------------------------------------
# Structural invariants behind the fixes
# ---------------------------------------------------------------------------


class TestModernCommentPartInvariants:
    """
    Package-level invariants that hold for ANY document Adeu writes comments
    into, independent of the specific defects above. Each of B1/B3 was a
    violation of one of them, and each is cheap to check on a saved package —
    which is what makes them worth stating separately from the repros.
    """

    @staticmethod
    def _thread_map(final: io.BytesIO):
        comments_xml = part_bytes(final, CT_COMMENTS).decode("utf-8")
        ext_xml = part_bytes(final, CT_EXTENDED).decode("utf-8")
        ids_xml = part_bytes(final, CT_IDS).decode("utf-8")
        return (
            set(re.findall(r'w14:paraId="([0-9A-Fa-f]{8})"', comments_xml)),
            re.findall(
                r'<w15:commentEx[^>]*w15:paraId="([0-9A-Fa-f]{8})"'
                r'(?:[^>]*w15:paraIdParent="([0-9A-Fa-f]{8})")?[^>]*/?>',
                ext_xml,
            ),
            set(re.findall(r'w16cid:paraId="([0-9A-Fa-f]{8})"', ids_xml)),
        )

    def _threaded_document(self) -> io.BytesIO:
        engine = RedlineEngine(doc_stream("Clause one text. Clause two text."), author="Sarah Chen")
        engine.apply_edits([ModifyText(target_text="Clause one", new_text="Clause one", comment="Root topic")])
        root = next(iter(engine.comments_manager.extract_comments_data()))
        engine.author = "Agent"
        engine.apply_review_actions([ReplyComment(target_id=f"Com:{root}", text="Reply one")])
        engine.apply_review_actions([ReplyComment(target_id=f"Com:{root}", text="Reply two")])
        engine.apply_edits([ModifyText(target_text="Clause two", new_text="Clause two", comment="Second root")])
        return engine.save_to_stream()

    def test_every_comment_has_exactly_one_registration_per_part(self):
        para_ids, ex_entries, id_para_ids = self._thread_map(self._threaded_document())
        ex_para_ids = [pid for pid, _parent in ex_entries]

        assert len(para_ids) == 4, para_ids
        assert sorted(ex_para_ids) == sorted(para_ids), (
            f"commentsExtended registrations {ex_para_ids} do not match the comments {para_ids}"
        )
        assert len(ex_para_ids) == len(set(ex_para_ids)), f"duplicate commentEx entries: {ex_para_ids}"
        assert id_para_ids == para_ids, f"commentsIds registrations {id_para_ids} do not match the comments {para_ids}"

    def test_no_thread_points_at_a_paragraph_that_does_not_exist(self):
        para_ids, ex_entries, _ = self._thread_map(self._threaded_document())
        parents = {parent for _pid, parent in ex_entries if parent}
        assert parents, "the document must contain at least one threaded reply"
        dangling = parents - para_ids
        assert not dangling, (
            f"w15:paraIdParent references {dangling}, which no comment paragraph carries — "
            "Word drops such a reply out of its thread"
        )

    def test_no_thread_is_its_own_parent(self):
        _, ex_entries, _ = self._thread_map(self._threaded_document())
        for pid, parent in ex_entries:
            assert parent != pid, f"commentEx {pid} is its own parent"

    def test_durable_ids_are_unique_and_word_readable(self):
        ids_xml = part_bytes(self._threaded_document(), CT_IDS).decode("utf-8")
        durable_ids = re.findall(r'w16cid:durableId="([0-9A-Fa-f]+)"', ids_xml)
        assert len(durable_ids) == 4, durable_ids
        assert len(set(durable_ids)) == len(durable_ids), f"durableId collision: {durable_ids}"
        assert all(int(d, 16) <= 0x7FFFFFFF for d in durable_ids), durable_ids

    def test_accept_all_books_match_the_document_in_both_modes(self):
        """
        `removed_comments` must equal the number of comment bodies that really
        disappeared — in BOTH modes. Under-reporting (0 while a human's comment
        is gone) and over-reporting (claiming removals that never happened) are
        the same class of defect: books that do not match the document.
        """
        for remove_comments in (False, True):
            engine = RedlineEngine(
                annotated_stream(
                    "Producing party may designate material Attorney's Eyes Only. A second clause stands alone.",
                    "A second clause",
                    "Standalone reviewer note.",
                    "Sarah Chen",
                ),
                author="Agent",
            )
            engine.apply_edits([ModifyText(target_text="Attorney's Eyes Only", new_text="")])
            before = set(engine._existing_comment_ids())

            counts = engine.accept_all_revisions(remove_comments=remove_comments)

            after = set(engine._existing_comment_ids())
            assert counts["removed_comments"] == len(before - after), (
                f"remove_comments={remove_comments}: reported {counts['removed_comments']} "
                f"removals but {len(before - after)} comment bodies actually disappeared"
            )
            assert len(engine.removed_comment_notes) == counts["removed_comments"]
            for note in engine.removed_comment_notes:
                assert note.startswith("Com:") and "(by " in note, note
