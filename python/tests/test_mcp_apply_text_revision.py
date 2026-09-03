import asyncio
import json
import sys
from unittest.mock import patch

import pytest
from docx import Document
from fastmcp.exceptions import ToolError

from adeu.cli import _main_impl
from adeu.mcp_components.tools.document import apply_text_revision
from adeu.payloads import BATCH_RECOVERY_PROTOCOL
from adeu.text_revision import check_criticmarkup, check_major_deletions


def run_cli(args, capsys):
    """Invoke the CLI in-process; returns (exit_code, stdout, stderr)."""
    from adeu.cli import main

    code = 0
    with patch.object(sys, "argv", ["adeu"] + [str(a) for a in args]):
        try:
            main()
        except SystemExit as e:
            code = e.code or 0
    captured = capsys.readouterr()
    return code, captured.out, captured.err


class MockContext:
    """Mock FastMCP Context to absorb async logging calls during tests."""

    async def info(self, msg, **kwargs):
        pass

    async def debug(self, msg, **kwargs):
        pass

    async def warning(self, msg, **kwargs):
        pass

    async def error(self, msg, **kwargs):
        pass


@pytest.fixture
def sample_docx(tmp_path) -> str:
    """Creates a basic DOCX file for testing."""
    doc = Document()
    doc.add_paragraph("This is the original paragraph one of the document.")
    doc.add_paragraph("This is paragraph two, containing more text for testing purposes.")
    doc.add_paragraph("And paragraph three concludes the baseline document content.")
    path = tmp_path / "sample.docx"
    doc.save(path)
    return str(path)


def test_apply_text_revision_produces_tracked_changes(sample_docx, tmp_path):
    ctx = MockContext()
    out_path = str(tmp_path / "output.docx")
    revised = (
        "This is the revised paragraph one of the document.\n\n"
        "This is paragraph two, containing more text for testing purposes.\n\n"
        "And paragraph three concludes the baseline document content."
    )

    result = asyncio.run(
        apply_text_revision(
            file_path=sample_docx,
            revised_text=revised,
            ctx=ctx,
            output_path=out_path,
            author="TestAuthor",
        )
    )

    assert out_path in result or "Saved" in result or "Batch complete" in result or "Applied" in result
    doc = Document(out_path)
    xml = doc.element.xml
    assert "w:del" in xml or "w:ins" in xml
    assert "revised" in xml


def test_apply_text_revision_refuses_major_deletion_without_flag(sample_docx, tmp_path):
    ctx = MockContext()
    out_path = str(tmp_path / "output.docx")

    # Major character deletion (>75% of characters, this document being under 2000 chars)
    short_revised = "This is short."

    with pytest.raises(ToolError) as exc_info:
        asyncio.run(
            apply_text_revision(
                file_path=sample_docx,
                revised_text=short_revised,
                ctx=ctx,
                output_path=out_path,
                allow_major_deletions=False,
            )
        )
    assert "major" in str(exc_info.value).lower() or "deletion" in str(exc_info.value).lower()

    # Allowed with flag
    result = asyncio.run(
        apply_text_revision(
            file_path=sample_docx,
            revised_text=short_revised,
            ctx=ctx,
            output_path=out_path,
            allow_major_deletions=True,
        )
    )
    assert out_path in result or "Saved" in result or "Applied" in result


def test_apply_text_revision_refuses_criticmarkup_input(sample_docx, tmp_path):
    ctx = MockContext()
    out_path = str(tmp_path / "output.docx")
    critic_revised = "This is original text {++with inserted text++} and {--deleted text--}."

    with pytest.raises(ToolError) as exc_info:
        asyncio.run(
            apply_text_revision(
                file_path=sample_docx,
                revised_text=critic_revised,
                ctx=ctx,
                output_path=out_path,
            )
        )
    assert "criticmarkup" in str(exc_info.value).lower()


# ---------------------------------------------------------------------------
# Only the OPEN CriticMarkup tokens mark markup-view text. Matching bare
# closing tokens refused ordinary prose: "~>" and "-->" are arrows people
# write (verifier finding, Task 15 attempt 3).
# ---------------------------------------------------------------------------


def test_criticmarkup_guard_ignores_arrows_and_bare_closing_tokens():
    for prose in (
        "Payment flows A ~> B.",
        "Escalation -> resolution within 5 days.",
        "The rate++} is stated below.",
        "Ends with <<} and --} and ==} and ~~}.",
    ):
        check_criticmarkup(prose)  # must not raise

    for markup in ("a {++b++} c", "a {--b--} c", "a {~~b~>c~~} d", "a {==b==} c", "a {>>b<<} c"):
        with pytest.raises(ValueError, match="CriticMarkup"):
            check_criticmarkup(markup)


def test_apply_text_revision_accepts_arrows_in_plain_text(sample_docx, tmp_path):
    ctx = MockContext()
    out_path = str(tmp_path / "arrows.docx")
    revised = (
        "This is the original paragraph one of the document. Payment flows A ~> B.\n\n"
        "This is paragraph two, containing more text for testing purposes. Escalation -> resolution.\n\n"
        "And paragraph three concludes the baseline document content."
    )

    result = asyncio.run(
        apply_text_revision(
            file_path=sample_docx,
            revised_text=revised,
            ctx=ctx,
            output_path=out_path,
            author="TestAuthor",
        )
    )

    assert "criticmarkup" not in result.lower()
    doc = Document(out_path)
    xml = doc.element.xml
    assert "w:ins" in xml
    assert "Payment flows A ~" in xml


def test_apply_text_revision_verification_failure_writes_unverified_sibling(sample_docx, tmp_path):
    ctx = MockContext()
    out_path = tmp_path / "output.docx"
    unverified_sibling = tmp_path / "output.unverified.docx"

    revised = (
        "This is the revised paragraph one of the document.\n\n"
        "This is paragraph two, containing more text for testing purposes.\n\n"
        "And paragraph three concludes the baseline document content."
    )

    # Force verification failure by patching text extraction on output verification
    from adeu import text_revision

    def bad_extract(doc):
        return "Mismatch text that does not match revised"

    with patch.object(text_revision, "_extract_clean_text_from_doc", side_effect=bad_extract):
        with pytest.raises(ToolError) as exc_info:
            asyncio.run(
                apply_text_revision(
                    file_path=sample_docx,
                    revised_text=revised,
                    ctx=ctx,
                    output_path=str(out_path),
                )
            )

    assert not out_path.exists()
    assert unverified_sibling.exists()
    assert "verification" in str(exc_info.value).lower() or "unverified" in str(exc_info.value).lower()


def test_cli_text_apply_still_behaves_identically(sample_docx, tmp_path):
    revised_file = tmp_path / "revised.txt"
    revised_text = (
        "This is the revised paragraph one of the document.\n\n"
        "This is paragraph two, containing more text for testing purposes.\n\n"
        "And paragraph three concludes the baseline document content."
    )
    revised_file.write_text(revised_text, encoding="utf-8")

    out_docx = tmp_path / "cli_out.docx"

    import sys

    test_args = ["adeu", "apply", sample_docx, str(revised_file), "-o", str(out_docx)]
    with patch.object(sys, "argv", test_args):
        try:
            _main_impl()
        except SystemExit as e:
            assert e.code in (0, None)

    assert out_docx.exists()
    doc = Document(str(out_docx))
    xml = doc.element.xml
    assert "revised" in xml


# ---------------------------------------------------------------------------
# The deletion guard is a CHARACTER budget only: a document made of many short
# paragraphs legitimately loses dozens of them in an ordinary edit, and a
# paragraph-count guard refused those (verifier finding, Task 15 attempt 2).
# ---------------------------------------------------------------------------


def test_major_deletion_guard_counts_characters_not_paragraphs():
    original = "\n\n".join(f"Clause {i}: this paragraph states an obligation." for i in range(200))
    # Drops 60 paragraphs — ~30% of the characters, well inside the budget.
    revised = "\n\n".join(f"Clause {i}: this paragraph states an obligation." for i in range(140))

    assert len(original) >= 2000
    check_major_deletions(original, revised)  # must not raise


def test_major_deletion_refusal_message_carries_the_full_recovery_advice():
    original = "Clause text that is quite long and detailed. " * 60
    revised = "Clause text."

    with pytest.raises(ValueError) as exc_info:
        check_major_deletions(original, revised, source_name="revised.txt")

    msg = str(exc_info.value)
    assert "'revised.txt'" in msg
    assert f"{len(revised):,} vs {len(original):,} characters" in msg
    assert "--page all --clean-view" in msg
    assert "--allow-major-deletions" in msg
    assert "allow_major_deletions=True" in msg


def test_major_deletion_refusal_message_without_a_filename():
    with pytest.raises(ValueError) as exc_info:
        check_major_deletions("Clause text that is quite long and detailed. " * 60, "Clause text.")

    assert "--allow-major-deletions" in str(exc_info.value)


def test_mcp_schema_states_the_character_deletion_thresholds():
    from adeu.server import mcp

    tools = asyncio.run(mcp.list_tools())
    tool = next(t for t in tools if t.name == "apply_text_revision")
    description = tool.parameters["properties"]["allow_major_deletions"]["description"]

    assert ">50% of characters (>75% for documents under 2000 characters)" in description
    assert "paragraph" not in description.lower()


# ---------------------------------------------------------------------------
# CLI text apply keeps the base CLI's reporting and error contract: a progress
# line, the overwrite warning, clean write failures, and a clean batch
# rejection instead of a raw traceback (verifier finding, Task 15 attempt 2).
# ---------------------------------------------------------------------------

REVISED_TEXT = (
    "This is the revised paragraph one of the document.\n\n"
    "This is paragraph two, containing more text for testing purposes.\n\n"
    "And paragraph three concludes the baseline document content."
)


def test_cli_text_apply_announces_progress_and_overwrite(sample_docx, tmp_path, capsys):
    revised_file = tmp_path / "revised.txt"
    revised_file.write_text(REVISED_TEXT, encoding="utf-8")
    out_docx = tmp_path / "cli_out.docx"
    out_docx.write_bytes(b"stale output")

    code, _out, err = run_cli(["apply", sample_docx, revised_file, "-o", out_docx], capsys)

    assert code == 0, err
    assert "changes to sample.docx..." in err, f"missing the 'Applying N changes to ...' progress line: {err}"
    assert f"Overwriting existing '{out_docx}'." in err


def test_cli_text_apply_write_failure_is_a_clean_error(sample_docx, tmp_path, capsys):
    revised_file = tmp_path / "revised.txt"
    revised_file.write_text(REVISED_TEXT, encoding="utf-8")
    # A regular file where the output's parent directory must be: the write
    # cannot succeed, and must not surface as a raw OSError traceback.
    blocker = tmp_path / "blocker"
    blocker.write_text("not a directory", encoding="utf-8")

    code, _out, err = run_cli(["apply", sample_docx, revised_file, "-o", blocker / "out.docx"], capsys)

    assert code == 1
    assert "Could not write output file" in err


def _control_char_fixture(tmp_path):
    doc = Document()
    doc.add_paragraph("This is a simple contract paragraph for testing purposes.")
    doc.add_paragraph("The second paragraph keeps the document plausible enough.")
    src = tmp_path / "ctrl.docx"
    doc.save(str(src))

    txt = tmp_path / "ctrl.txt"
    txt.write_text(
        "This is a bad\x01value contract paragraph for testing purposes.\n\n"
        "The second paragraph keeps the document plausible enough.",
        encoding="utf-8",
    )
    return src, txt


def test_cli_text_apply_batch_rejection_is_a_clean_error(tmp_path, capsys):
    src, txt = _control_char_fixture(tmp_path)
    out = tmp_path / "ctrl_out.docx"

    code, _out, err = run_cli(["apply", src, txt, "-o", out], capsys)

    assert code == 1
    assert "Batch rejected" in err
    assert "control character" in err
    assert BATCH_RECOVERY_PROTOCOL in err
    assert not out.exists()


def test_cli_text_apply_batch_rejection_json_is_a_failure_envelope(tmp_path, capsys):
    src, txt = _control_char_fixture(tmp_path)
    out = tmp_path / "ctrl_out.docx"

    code, stdout, _err = run_cli(["apply", src, txt, "-o", out, "--json"], capsys)

    assert code == 1
    envelope = json.loads(stdout)
    assert envelope["error"] == "batch_validation_failed"
    assert not out.exists()


def test_strip_cell_anchors():
    from adeu.text_revision import strip_cell_anchors

    assert strip_cell_anchors("Widget B | {#cell:05856B27} | $250.00") == "Widget B |  | $250.00"
    assert strip_cell_anchors("Widget B | Audit {#cell:05856B27} | $250.00") == "Widget B | Audit | $250.00"
