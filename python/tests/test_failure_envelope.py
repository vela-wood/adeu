import json
import re
from io import BytesIO
from typing import List

import pytest
from docx import Document

from adeu.cli import _load_batch_from_json, _set_json_mode
from adeu.mcp_components.tools.document import process_document_batch
from adeu.models import DocumentChange, ModifyText, ReplyComment
from adeu.payloads import failure_envelope
from adeu.redline.engine import BatchValidationError, RedlineEngine


def _create_simple_docx() -> BytesIO:
    doc = Document()
    doc.add_paragraph("Paragraph zero.")
    doc.add_paragraph("Paragraph one.")
    doc.add_paragraph("Paragraph two.")
    doc.add_paragraph("Paragraph three.")
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def test_schema_failure_envelope_indices(tmp_path, capsys):
    # Missing required 'target_text' at index 1
    batch_json = [
        {"type": "modify", "target_text": "Paragraph zero.", "new_text": "Zero updated."},
        {"type": "modify"},  # index 1 missing fields
    ]
    p = tmp_path / "changes.json"
    p.write_text(json.dumps(batch_json), encoding="utf-8")

    _set_json_mode(True)
    try:
        with pytest.raises(SystemExit) as exc_info:
            _load_batch_from_json(p)
        assert exc_info.value.code == 1
        captured = capsys.readouterr()
        data = json.loads(captured.out)
        assert data["error"] == "invalid_changes_file"
        assert "\n" not in data["message"]
        assert len(data["failed"]) == 1
        assert data["failed"][0]["index"] == 1
    finally:
        _set_json_mode(False)


def test_engine_failure_envelope_indices():
    stream = _create_simple_docx()
    engine = RedlineEngine(stream, author="TestAuthor")

    changes: List[DocumentChange] = [
        ModifyText(type="modify", target_text="Paragraph zero.", new_text="Zero updated."),
        ModifyText(type="modify", target_text="Paragraph one.", new_text="One updated."),
        ModifyText(type="modify", target_text="Non-existent target text", new_text="Mismatch."),  # index 2
    ]

    with pytest.raises(BatchValidationError) as exc_info:
        engine.process_batch(changes)

    err = exc_info.value
    assert hasattr(err, "failed")
    assert [f[0] for f in err.failed] == [2]


def test_multi_failure_indices():
    stream = _create_simple_docx()
    engine = RedlineEngine(stream, author="TestAuthor")

    changes: List[DocumentChange] = [
        ModifyText(type="modify", target_text="Non-existent 0", new_text="A"),  # index 0 fails
        ModifyText(type="modify", target_text="Paragraph one.", new_text="One updated."),
        ModifyText(type="modify", target_text="Paragraph two.", new_text="Two updated."),
        ModifyText(type="modify", target_text="Non-existent 3", new_text="B"),  # index 3 fails
    ]

    with pytest.raises(BatchValidationError) as exc_info:
        engine.process_batch(changes)

    err = exc_info.value
    assert [f[0] for f in err.failed] == [0, 3]


def test_action_failure_index_is_batch_relative():
    stream = _create_simple_docx()
    engine = RedlineEngine(stream, author="TestAuthor")

    changes: List[DocumentChange] = [
        ModifyText(type="modify", target_text="Paragraph zero.", new_text="Zero updated."),  # index 0
        ReplyComment(type="reply", target_id="Com:1", text="   "),  # index 1 (empty reply fails)
    ]

    with pytest.raises(BatchValidationError) as exc_info:
        engine.process_batch(changes)

    err = exc_info.value
    assert [f[0] for f in err.failed] == [1]


def test_action_not_found_failure_index_is_batch_relative():
    stream = _create_simple_docx()
    engine = RedlineEngine(stream, author="TestAuthor")

    changes: List[DocumentChange] = [
        ModifyText(type="modify", target_text="Paragraph zero.", new_text="Zero updated."),  # index 0
        ReplyComment(type="reply", target_id="Com:999", text="Reply to non-existent comment"),  # index 1
    ]

    with pytest.raises(BatchValidationError) as exc_info:
        engine.process_batch(changes)

    err = exc_info.value
    assert [f[0] for f in err.failed] == [1]


def test_message_field_always_present_and_one_line():
    env = failure_envelope("test_error", [(0, "Reason 0")], "Line 1\nLine 2\n\nLine 3")
    assert "message" in env
    assert "\n" not in env["message"]
    assert env["error"] == "test_error"
    assert env["failed"] == [{"index": 0, "reason": "Reason 0"}]


@pytest.mark.anyio
async def test_mcp_failure_carries_envelope(tmp_path):
    doc_path = tmp_path / "test.docx"
    doc_path.write_bytes(_create_simple_docx().getvalue())

    changes = [
        {"type": "modify", "target_text": "Paragraph zero.", "new_text": "Zero updated."},
        {"type": "modify", "target_text": "Missing target", "new_text": "Fail."},  # index 1
    ]

    class FakeContext:
        async def info(self, *a, **kw):
            pass

        async def warning(self, *a, **kw):
            pass

        async def debug(self, *a, **kw):
            pass

        async def error(self, *a, **kw):
            pass

    # partial=False: the envelope is the REJECTED-batch payload. Under the
    # default salvage mode this batch would apply edit 1 and save, and a saved
    # output must not be reported with a "nothing was written" envelope.
    result = await process_document_batch(
        reasoning="Testing failure envelope",
        original_docx_path=str(doc_path),
        author_name="Tester",
        ctx=FakeContext(),  # type: ignore
        changes=changes,
        partial=False,
    )

    assert "```json" in result
    m = re.search(r"```json\s*(\{.*?\})\s*```", result, re.DOTALL)
    assert m is not None
    data = json.loads(m.group(1))
    assert data["error"] == "batch_validation_failed"
    assert [f["index"] for f in data["failed"]] == [1]


@pytest.mark.anyio
async def test_mcp_mixed_schema_and_engine_failure_indices(tmp_path):
    doc_path = tmp_path / "test.docx"
    doc_path.write_bytes(_create_simple_docx().getvalue())

    # Index 0: Schema-invalid item (missing target_text and new_text)
    # Index 1: Engine-failing edit (non-existent target)
    changes = [
        {"type": "modify"},  # index 0 invalid schema
        {"type": "modify", "target_text": "Missing target", "new_text": "Fail."},  # index 1 engine fail
    ]

    class FakeContext:
        async def info(self, *a, **kw):
            pass

        async def warning(self, *a, **kw):
            pass

        async def debug(self, *a, **kw):
            pass

        async def error(self, *a, **kw):
            pass

    result = await process_document_batch(
        reasoning="Testing mixed failure indices",
        original_docx_path=str(doc_path),
        author_name="Tester",
        ctx=FakeContext(),  # type: ignore
        changes=changes,
    )

    assert "```json" in result
    m = re.search(r"```json\s*(\{.*?\})\s*```", result, re.DOTALL)
    assert m is not None
    data = json.loads(m.group(1))
    assert data["error"] == "batch_validation_failed"
    assert [f["index"] for f in data["failed"]] == [0, 1]


def test_prose_still_present_for_humans():
    stream = _create_simple_docx()
    engine = RedlineEngine(stream, author="TestAuthor")

    changes: List[DocumentChange] = [
        ModifyText(type="modify", target_text="Non-existent target", new_text="Fail."),
    ]

    with pytest.raises(BatchValidationError) as exc_info:
        engine.process_batch(changes)

    err = exc_info.value
    assert hasattr(err, "errors")
    assert len(err.errors) > 0
    assert "Target text not found" in err.errors[0]
