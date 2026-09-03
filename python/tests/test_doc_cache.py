"""
Server-layer projection cache tests (adeu.mcp_components.doc_cache).

The cache contract, ported from the Node server (docs/Performance.md §5.1):
responses served from a warm cache must be BYTE-IDENTICAL to a cache-less
server's, invalidation is purely stat-based (any rewrite changes the key),
and the cache never holds more than MAX_ENTRIES document versions.
"""

import asyncio
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from adeu.ingest import _extract_text_from_doc
from adeu.mcp_components._response_builders import (
    build_appendix_response,
    build_full_document_response,
    build_outline_response,
    build_paginated_response,
    build_search_response,
)
from adeu.mcp_components.doc_cache import MAX_ENTRIES, DocProjectionCache, doc_cache, get_doc_cache_capacity
from adeu.mcp_components.tools.document import _as_tool_result, _read_docx_disk
from adeu.utils.docx import strip_bom_from_docx_bytes


class MockContext:
    async def info(self, msg, **kwargs):
        pass

    async def debug(self, msg, **kwargs):
        pass

    async def warning(self, msg, **kwargs):
        pass

    async def error(self, msg, **kwargs):
        pass


@pytest.fixture(autouse=True)
def _fresh_cache(monkeypatch):
    monkeypatch.delenv("ADEU_DOC_CACHE_ENTRIES", raising=False)
    doc_cache._max_entries = MAX_ENTRIES
    doc_cache.clear()
    yield
    doc_cache.clear()


@pytest.fixture
def structured_docx(tmp_path) -> str:
    doc = Document()
    doc.add_heading("Alpha Section", level=1)
    doc.add_paragraph("Alpha body text mentioning the Agreement.")
    doc.add_heading("Beta Subsection", level=2)
    doc.add_paragraph("Beta body text, second occurrence of Agreement.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "cell one"
    table.cell(1, 1).text = "cell two"
    doc.add_heading("Gamma Section", level=1)
    doc.add_paragraph("Gamma closing paragraph.")
    path = tmp_path / "structured.docx"
    doc.save(path)
    return str(path)


def _read(path, **kwargs):
    ctx = MockContext()
    defaults = dict(clean_view=False, mode="full", page=None)
    defaults.update(kwargs)
    return asyncio.run(_read_docx_disk(path, ctx, **defaults))


def _payload(result):
    return (result.content, result.structured_content)


def _page_num(page):
    """document.py's page coercion, so the reference targets the same page."""
    if page is None:
        return 1
    s = str(page).strip()
    if s.isdigit() or (s.startswith(("-", "+")) and s[1:].isdigit()):
        return int(s)
    return 1


def _cacheless_payload(
    path,
    *,
    clean_view=False,
    mode="full",
    page=None,
    search_query=None,
    outline_max_level=2,
    outline_verbose=False,
):
    """Reference response built with doc_cache completely out of the picture.

    doc_cache hands the builders precomputed artifacts (pagination_result, and
    for outline the extracted nodes). Here the builders receive none of them and
    must derive everything from a freshly parsed document — so this is the only
    comparison that can catch the cache's artifacts drifting from what the
    builders compute for themselves. Comparing a cleared cache against a warm
    one cannot: both sides run the same cache code.
    """
    data = strip_bom_from_docx_bytes(Path(path).read_bytes())
    doc = Document(BytesIO(data))
    text = _extract_text_from_doc(doc, clean_view=clean_view, include_appendix=(mode == "appendix"))

    if search_query is not None:
        res = build_search_response(text, search_query, False, False, page, path)
    elif mode == "appendix":
        res = build_appendix_response(text, _page_num(page), path)
    elif mode == "outline":
        res = build_outline_response(
            doc,
            text,
            path,
            outline_max_level=outline_max_level,
            outline_verbose=outline_verbose,
        )
    elif page is not None and str(page).strip().lower() == "all":
        res = build_full_document_response(text, path)
    else:
        res = build_paginated_response(text, _page_num(page), path)
    # Through the same ToolResult wrapper the tool uses, so the comparison is
    # about content rather than container shape.
    return _payload(_as_tool_result(res))


@pytest.mark.parametrize(
    "kwargs",
    [
        dict(mode="full", page=1),
        dict(mode="full", page="all"),
        dict(mode="outline"),
        dict(mode="appendix"),
        dict(search_query="Agreement"),
        dict(clean_view=True, mode="full", page=1),
    ],
)
def test_warm_responses_byte_identical_to_cold(structured_docx, kwargs):
    cold = _payload(_read(structured_docx, **kwargs))

    # Second call is served from the warm cache.
    warm = _payload(_read(structured_docx, **kwargs))
    assert warm == cold

    # A cleared cache agrees too (this re-runs the cache's cold path).
    doc_cache.clear()
    fresh = _payload(_read(structured_docx, **kwargs))
    assert fresh == cold

    # And so does a response the builders compute with no cache artifacts at
    # all — the actual "byte-identical to a cache-less server" contract.
    assert _cacheless_payload(structured_docx, **kwargs) == cold


def test_modes_share_one_projection(structured_docx):
    """A full read warms the entry; outline/search on the same version must
    match what a cache-less server would produce."""
    _read(structured_docx, mode="full", page=1)
    warm_outline = _payload(_read(structured_docx, mode="outline"))
    warm_search = _payload(_read(structured_docx, search_query="Agreement"))

    doc_cache.clear()
    cold_outline = _payload(_read(structured_docx, mode="outline"))
    doc_cache.clear()
    cold_search = _payload(_read(structured_docx, search_query="Agreement"))

    assert warm_outline == cold_outline
    assert warm_search == cold_search

    # Both comparisons above run cache code on both sides, so neither can see
    # the cache's artifacts drifting from the builders' own computation. Anchor
    # them to a cache-less reference.
    assert warm_outline == _cacheless_payload(structured_docx, mode="outline")
    assert warm_search == _cacheless_payload(structured_docx, search_query="Agreement")


@pytest.mark.parametrize("max_level", [1, 2, 6])
@pytest.mark.parametrize("verbose", [False, True])
def test_outline_artifacts_match_a_cacheless_build(structured_docx, max_level, verbose):
    """The outline is the one artifact the cache extracts itself (it needs the
    paragraph-offset map from its own parse), so it is the most likely to drift
    from what build_outline_response derives on its own."""
    warm = _payload(_read(structured_docx, mode="outline", outline_max_level=max_level, outline_verbose=verbose))
    assert warm == _cacheless_payload(
        structured_docx, mode="outline", outline_max_level=max_level, outline_verbose=verbose
    )


def test_rewrite_invalidates(structured_docx):
    first = _read(structured_docx, mode="full", page=1)
    assert "Alpha body text" in first.structured_content["markdown"]

    doc = Document()
    doc.add_paragraph("Completely different content now.")
    doc.save(structured_docx)

    second = _read(structured_docx, mode="full", page=1)
    assert "Completely different content now." in second.structured_content["markdown"]
    assert "Alpha body text" not in second.structured_content["markdown"]


def test_lru_bound(tmp_path):
    paths = []
    for i in range(MAX_ENTRIES + 2):
        doc = Document()
        doc.add_paragraph(f"Document number {i}.")
        p = tmp_path / f"doc{i}.docx"
        doc.save(p)
        paths.append(str(p))

    for p in paths:
        _read(p, mode="full", page=1)

    assert len(doc_cache._entries) <= MAX_ENTRIES

    # Every document still reads correctly after its entry was evicted.
    for i, p in enumerate(paths):
        res = _read(p, mode="full", page=1)
        assert f"Document number {i}." in res.structured_content["markdown"]


def test_constructor_max_entries_is_honored():
    """entry() used to bound the cache by the module constant, so a
    caller-supplied cap was accepted and then silently ignored —
    DocProjectionCache(max_entries=1) still held MAX_ENTRIES versions.

    entry() is pure dict bookkeeping over an opaque key tuple, so this needs
    no files on disk.
    """
    cache = DocProjectionCache(max_entries=1)
    k1 = ("/doc-a", 1, 10)
    k2 = ("/doc-b", 2, 20)

    e1 = cache.entry(k1)
    assert cache.entry(k1) is e1, "same version must reuse its entry"

    cache.entry(k2)
    assert len(cache._entries) == 1
    assert k2 in cache._entries
    assert k1 not in cache._entries, "LRU eviction must respect max_entries=1"

    # The evicted version comes back as a fresh (cold) entry, not the old one.
    assert cache.entry(k1) is not e1


def test_default_max_entries_matches_the_module_constant():
    cache = DocProjectionCache()
    for i in range(MAX_ENTRIES + 2):
        cache.entry((f"/doc-{i}", i, i))
    assert len(cache._entries) == MAX_ENTRIES


def test_max_entries_is_clamped_to_at_least_one():
    """A zero/negative cap would evict the entry the caller just asked for,
    turning every lookup into a miss."""
    cache = DocProjectionCache(max_entries=0)
    key = ("/doc-a", 1, 10)
    entry = cache.entry(key)
    assert cache.entry(key) is entry


def test_clean_and_raw_views_are_independent(structured_docx):
    raw = _read(structured_docx, mode="full", page=1)
    clean = _read(structured_docx, clean_view=True, mode="full", page=1)
    # No tracked changes in this fixture, so the bodies agree — but both
    # calls must succeed and be individually cached without cross-talk.
    assert raw.structured_content["markdown"]
    assert clean.structured_content["markdown"]

    key = doc_cache.stat_key(structured_docx)
    entry = doc_cache.entry(key)
    assert entry.raw.base_text is not None
    assert entry.clean.base_text is not None


def test_lru_size_is_env_tunable(monkeypatch):
    monkeypatch.setenv("ADEU_DOC_CACHE_ENTRIES", "7")
    assert get_doc_cache_capacity() == 7
    cache = DocProjectionCache()
    assert cache._max_entries == 7
    for i in range(8):
        cache.entry((f"/doc-{i}", i, i))
    assert len(cache._entries) == 7


def test_lru_size_falls_back_on_garbage_and_zero(monkeypatch):
    for bad_val in ["abc", "0", "-2"]:
        monkeypatch.setenv("ADEU_DOC_CACHE_ENTRIES", bad_val)
        assert get_doc_cache_capacity() == 3
