import json
import sys
from unittest.mock import patch

import docx

from adeu.models import ModifyText


def run_cli(args, capsys):
    """Invoke the CLI in-process; returns (exit_code, stdout, stderr)."""
    from adeu.cli import main

    code = 0
    with patch.object(sys, "argv", ["adeu"] + [str(a) for a in args]):
        try:
            main()
        except SystemExit as e:
            code = e.code or 0
    captured = capsys.readouterr()
    return code, captured.out, captured.err


def test_diff_json_is_unindented_and_omits_defaults(tmp_path, capsys):
    orig_path = tmp_path / "orig.docx"
    doc1 = docx.Document()
    doc1.add_paragraph("The quick brown fox jumps over the lazy dog.")
    doc1.save(str(orig_path))

    mod_path = tmp_path / "mod.docx"
    doc2 = docx.Document()
    doc2.add_paragraph("The fast brown fox jumps over the sleeping dog.")
    doc2.save(str(mod_path))

    code, stdout, stderr = run_cli(["diff", str(orig_path), str(mod_path), "--json"], capsys)
    assert code == 0, stderr

    json_str = stdout.strip()
    # Unindented JSON: no newlines inside the JSON output
    assert "\n" not in json_str, f"JSON output should be unindented on one line, got:\n{json_str}"

    data = json.loads(json_str)
    assert isinstance(data, list)
    assert len(data) > 0

    for edit in data:
        # Default fields should be omitted
        assert "match_mode" not in edit, f"Default match_mode should be omitted in {edit}"
        assert "regex" not in edit, f"Default regex should be omitted in {edit}"
        # Boilerplate 'Diff:' comments should be omitted
        if "comment" in edit and edit["comment"] is not None:
            assert not edit["comment"].startswith("Diff:"), f"Boilerplate comment should be omitted in {edit}"


def test_diff_json_at_least_25_percent_smaller(tmp_path, capsys):
    orig_path = tmp_path / "orig.docx"
    doc1 = docx.Document()
    doc1.add_paragraph("First paragraph in the original document.")
    doc1.add_paragraph("Second paragraph in the original document.")
    doc1.add_paragraph("Third paragraph in the original document.")
    doc1.save(str(orig_path))

    mod_path = tmp_path / "mod.docx"
    doc2 = docx.Document()
    doc2.add_paragraph("First modified paragraph in the document.")
    doc2.add_paragraph("Second paragraph in the original document.")
    doc2.add_paragraph("Third modified paragraph in the document.")
    doc2.save(str(mod_path))

    code, stdout, stderr = run_cli(["diff", str(orig_path), str(mod_path), "--json"], capsys)
    assert code == 0, stderr

    compact_len = len(stdout.strip())

    # Calculate old indented size with default fields and boilerplate comments
    from adeu.diff import generate_structured_edits
    from adeu.ingest import _extract_text_from_doc

    t1_ret = _extract_text_from_doc(doc1, clean_view=True, include_appendix=False, return_structure=True)
    t2_ret = _extract_text_from_doc(doc2, clean_view=True, include_appendix=False, return_structure=True)
    assert isinstance(t1_ret, tuple) and isinstance(t2_ret, tuple)
    t1, s1 = t1_ret
    t2, s2 = t2_ret
    assert s1 is not None and s2 is not None
    old_edits, _ = generate_structured_edits(t1, s1, t2, s2)  # type: ignore[arg-type]
    old_data = [edit.model_dump(exclude={"_match_start_index"}) for edit in old_edits]
    old_json = json.dumps(old_data, indent=2, ensure_ascii=False)
    old_len = len(old_json)

    assert compact_len <= 0.75 * old_len, (
        f"Compact size {compact_len} is not <= 75% of old size {old_len} ({compact_len / old_len:.2%})"
    )


def test_diff_json_round_trips_through_apply(tmp_path, capsys):
    orig_path = tmp_path / "orig.docx"
    doc1 = docx.Document()
    doc1.add_paragraph("Section 1: Initial terms and conditions.")
    doc1.add_paragraph("Section 2: Payment schedule is monthly.")
    doc1.save(str(orig_path))

    mod_path = tmp_path / "mod.docx"
    doc2 = docx.Document()
    doc2.add_paragraph("Section 1: Updated terms and conditions.")
    doc2.add_paragraph("Section 2: Payment schedule is quarterly.")
    doc2.save(str(mod_path))

    edits_json_path = tmp_path / "edits.json"
    code_diff, stdout_diff, stderr_diff = run_cli(
        ["diff", str(orig_path), str(mod_path), "--json", "-o", str(edits_json_path)], capsys
    )
    assert code_diff == 0, stderr_diff

    # Assert compaction properties before applying
    json_str = edits_json_path.read_text(encoding="utf-8").strip()
    assert "\n" not in json_str, f"JSON output should be unindented on one line, got:\n{json_str}"
    assert "Diff:" not in json_str, f"Boilerplate 'Diff:' comment should be omitted, got:\n{json_str}"

    data = json.loads(json_str)
    assert isinstance(data, list)
    assert len(data) > 0
    for edit in data:
        assert "match_mode" not in edit, f"Default match_mode should be omitted in {edit}"
        assert "regex" not in edit, f"Default regex should be omitted in {edit}"

    out_docx_path = tmp_path / "applied.docx"
    code_apply, stdout_apply, stderr_apply = run_cli(
        ["apply", str(orig_path), str(edits_json_path), "-o", str(out_docx_path)], capsys
    )
    assert code_apply == 0, stderr_apply

    # Extract text from applied document (clean view) and verify it matches doc2
    from adeu.ingest import _extract_text_from_doc

    applied_doc = docx.Document(str(out_docx_path))
    applied_text = _extract_text_from_doc(applied_doc, clean_view=True, include_appendix=False)
    expected_text = _extract_text_from_doc(doc2, clean_view=True, include_appendix=False)
    assert applied_text == expected_text


def test_diff_json_preserves_a_meaningful_comment(tmp_path, capsys):
    orig_path = tmp_path / "orig.docx"
    doc1 = docx.Document()
    doc1.add_paragraph("Original paragraph.")
    doc1.add_paragraph("Second paragraph.")
    doc1.save(str(orig_path))

    mod_path = tmp_path / "mod.docx"
    doc2 = docx.Document()
    doc2.add_paragraph("Modified paragraph.")
    doc2.add_paragraph("Second modified paragraph.")
    doc2.save(str(mod_path))

    custom_edit = ModifyText(
        type="modify",
        target_text="Original paragraph.",
        new_text="Modified paragraph.",
        comment="Reviewed by Legal Team",
    )
    boilerplate_edit = ModifyText(
        type="modify",
        target_text="Second paragraph.",
        new_text="Second modified paragraph.",
        comment="Diff: replaced 'Second paragraph.' with 'Second modified paragraph.'",
    )

    with patch("adeu.diff.generate_structured_edits", return_value=([custom_edit, boilerplate_edit], [])):
        code, stdout, stderr = run_cli(["diff", str(orig_path), str(mod_path), "--json"], capsys)
        assert code == 0, stderr

        json_str = stdout.strip()
        assert "\n" not in json_str, f"JSON output should be unindented on one line, got:\n{json_str}"
        assert "Diff:" not in json_str, f"Boilerplate 'Diff:' comment should be stripped, got:\n{json_str}"

        data = json.loads(json_str)
        assert len(data) == 2
        assert data[0].get("comment") == "Reviewed by Legal Team"
        assert "comment" not in data[1]
