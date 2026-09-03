import io

from docx import Document

from adeu.ingest import extract_text_from_stream
from adeu.models import ModifyText
from adeu.redline.engine import RedlineEngine


def test_surgical_interior_word_diff():
    doc = Document()
    doc.add_paragraph("The quick brown fox jumped.")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="Test AI")
    engine.process_batch(
        [ModifyText(type="modify", target_text="The quick brown fox jumped.", new_text="The slow brown fox leapt.")]
    )

    result_text = extract_text_from_stream(engine.save_to_stream(), clean_view=False)

    # Assertions:
    # 1. "brown fox" should NOT be inside a deletion or insertion tag.
    assert "{--The quick brown fox jumped.--}" not in result_text
    # 2. It should surgically strike "quick" and "jumped"
    assert "{--quick--}{++slow++}" in result_text
    assert " brown fox " in result_text
    assert "{--jumped--}{++leapt++}" in result_text


def _apply_modify(target_text: str, new_text: str) -> io.BytesIO:
    doc = Document()
    doc.add_paragraph(target_text)
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    engine = RedlineEngine(stream, author="Test AI")
    engine.process_batch([ModifyText(type="modify", target_text=target_text, new_text=new_text)])
    return engine.save_to_stream()


def test_single_line_hyphen_insertion_is_not_converted_to_list_paragraph():
    """
    Regression: a single-line word-diff insertion fragment beginning with "- "
    (hyphen + space) must stay an inline tracked insertion.

    ModifyText("Product" -> "Product - Draft") trims the common prefix and mints
    an INSERTION sub-edit whose new_text is " - Draft". `_parse_markdown_style`
    lstrips that to "- ", reads it as a bullet marker, and because the anchor
    supplies a paragraph context (`current_p is not None`) the block-mode gate
    fires: the fragment becomes a brand-new numbered ListParagraph, the "- " is
    eaten as a fabricated list marker, and the edit still reports
    status="applied". Silent structural corruption, valid OOXML, no warning.

    A fragment with no line break is by construction part of an existing
    paragraph and must never enter heading/list-style conversion.
    """
    out_stream = _apply_modify("Product", "Product - Draft")

    raw_text = extract_text_from_stream(out_stream, clean_view=False)
    out_stream.seek(0)
    clean_text = extract_text_from_stream(out_stream, clean_view=True)
    out_stream.seek(0)
    doc_result = Document(out_stream)

    # 1. The insertion is inline and the literal hyphen survives verbatim.
    assert "Product{++ - Draft++}" in raw_text, f"Expected inline insertion, got: {raw_text!r}"

    # 2. No paragraph split and no fabricated bullet marker.
    assert "\n\n" not in clean_text, f"Insertion split the paragraph: {clean_text!r}"
    assert "* {++Draft++}" not in raw_text, f"Hyphen was eaten as a list marker: {raw_text!r}"
    assert clean_text.strip() == "Product - Draft"

    # 3. Structurally: still one paragraph, no list style, no numbering.
    assert len(doc_result.paragraphs) == 1, (
        f"Expected 1 paragraph, got {len(doc_result.paragraphs)}: {[p.style.name for p in doc_result.paragraphs]}"
    )
    body_xml = doc_result.element.body.xml
    assert "w:numPr" not in body_xml, "Spurious numbering applied to an inline insertion"
    assert "ListParagraph" not in body_xml, "Spurious ListParagraph style applied to an inline insertion"


def test_single_line_dash_insertion_matches_em_dash_control():
    """
    The defect is specific to the ASCII hyphen-minus. The em-dash spelling of the
    same edit already behaves correctly, so both must produce the same shape.
    """
    hyphen = extract_text_from_stream(_apply_modify("Product", "Product - Draft"), clean_view=False)
    em_dash = extract_text_from_stream(_apply_modify("Product", "Product \u2014 Draft"), clean_view=False)

    assert "Product{++ \u2014 Draft++}" in em_dash, f"Control case changed behaviour: {em_dash!r}"
    assert hyphen.replace(" - ", " \u2014 ") == em_dash, (
        f"Hyphen and em-dash insertions diverge:\n  hyphen : {hyphen!r}\n  em-dash: {em_dash!r}"
    )


def test_multiline_bullet_insertion_still_creates_a_list_paragraph():
    """
    Guard for the fix: requiring a real line break before block mode must not
    disable genuine Markdown block insertions.
    """
    out_stream = _apply_modify("Intro", "Intro\n\n- Bullet one")
    doc_result = Document(out_stream)

    assert len(doc_result.paragraphs) == 2, "Multi-line bullet insert should still split the paragraph"
    p_new = doc_result.paragraphs[1]
    assert p_new.style is not None and "List" in p_new.style.name
    assert p_new._element.xpath(".//w:numPr"), "Genuine bullet insert lost its numbering"
